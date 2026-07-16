from __future__ import annotations

import env_local  # noqa: F401 — 加载 .env / .env.local
import aippt_auth
import base64
import hashlib
import json
import os
import re
import secrets
import sqlite3
import sys
import threading
import time
import urllib.error
import urllib.parse
import urllib.request
import uuid
from datetime import datetime, timedelta, timezone
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from urllib.parse import parse_qs, urlparse

import export_utils
import institution_store as org_inst_store
from learning_workflow import normalize_error_type, practice_tier, transition
from paper_workflow import normalize_answer_state, normalize_eight_steps, paper_summary, split_numbered_questions
from gaokao_core import SEED_MOTHER_QUESTIONS, build_gaokao_card, match_mother_question
import gaokao_import
import gaokao_rag
import metaso_client
import parent_report
import teacher_portal
from vector_search import rerank_hits
from handwriting_ocr import (
    HANDWRITING_OCR_HINT,
    HANDWRITING_PHOTO_TIPS,
    merge_handwriting_prompt,
    merge_ocr_passes,
    normalize_ocr_confidence,
    preprocess_for_handwriting,
    preprocess_red_marks,
    score_single_ocr_result,
)
from speed_pipeline import (
    INSTANT_HIT_THRESHOLD,
    QUICK_DIAGNOSE_PROMPT,
    SEARCH_HIT_THRESHOLD,
    build_fast_diagnosis_from_hits,
    build_instant_diagnosis,
    build_skeleton_diagnosis,
    fallback_eliminate_variants,
    parallel_run,
)
from question_search import (
    ZONE_GAOKAO_MATH,
    ensure_gaokao_question_tables,
    list_documents,
    search_questions,
    zone_stats,
)
from latex_pipeline import convert_bytes, detect_paste_format, text_to_latex


ROOT = Path(__file__).resolve().parent
PUBLIC_DIR = ROOT / "web"
DATA_DIR = ROOT / "data"
UPLOAD_DIR = PUBLIC_DIR / "uploads"
CARD_DIR = PUBLIC_DIR / "cards"
EXPORT_DIR = PUBLIC_DIR / "exports"


def deduplicate_repeated_ocr_text(value: str) -> str:
    """Remove OCR's common full-block duplication without deleting real repeated terms."""
    text = str(value or "").replace("\r\n", "\n").strip()
    if not text:
        return ""
    compact = re.sub(r"\s+", "", text)
    if len(compact) >= 40 and len(compact) % 2 == 0 and compact[: len(compact) // 2] == compact[len(compact) // 2 :]:
        target_length = len(compact) // 2
        consumed = 0
        kept = []
        for char in text:
            kept.append(char)
            if not char.isspace():
                consumed += 1
            if consumed >= target_length:
                break
        return "".join(kept).strip()
    lines: list[str] = []
    for line in text.splitlines():
        normalized = re.sub(r"\s+", "", line)
        if normalized and lines and normalized == re.sub(r"\s+", "", lines[-1]) and len(normalized) >= 12:
            continue
        lines.append(line)
    return "\n".join(lines).strip()
DB_PATH = DATA_DIR / "gaokao.db"

MINIMAX_ENDPOINT = "https://api.minimax.chat/v1/chat/completions"
DEEPSEEK_ENDPOINT = "https://api.deepseek.com/chat/completions"
OPENAI_IMAGE_ENDPOINT = "https://api.openai.com/v1/images/generations"
FENNO_BASE_URL = os.environ.get("FENNO_BASE_URL", "https://api.fenno.ai")
DEFAULT_OPENAI_IMAGE_MODEL = os.environ.get("OPENAI_IMAGE_MODEL", "gpt-image-1.5")
DEFAULT_OPENAI_IMAGE_SIZE = os.environ.get("OPENAI_IMAGE_SIZE", "1024x1536")
DEFAULT_OPENAI_IMAGE_QUALITY = os.environ.get("OPENAI_IMAGE_QUALITY", "high")
ADMIN_EMAIL = os.environ.get("ADMIN_EMAIL", "admin@edu.ms1001.com").strip().lower()
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "GaoKao2026-edu-ms1001-admin").strip()
PASS_REQUIRED_CORRECT = 2
AUTH_BASE_URL = os.environ.get("AUTH_BASE_URL", "https://ai.ms1001.com").strip().rstrip("/")
PUBLIC_AUTH_BASE_URL = os.environ.get("PUBLIC_AUTH_BASE_URL", AUTH_BASE_URL or "https://ai.ms1001.com").strip().rstrip("/")
USE_UNIFIED_AUTH = os.environ.get("USE_UNIFIED_AUTH", "1").strip().lower() not in {"0", "false", "no", "off"}
UNIFIED_PLATFORM_ID = os.environ.get("UNIFIED_PLATFORM_ID", "edu.ms1001.com").strip()


PORTAL_TOOLS = [
    {"id": "wrong-transfer", "number": "01", "label": "错题拆解核心", "tagline": "核心闭环 · 归因拆题、变式训练、过关移除", "category": "错题突破", "route": "diagnose", "mode": "wrong", "delivery": "diagnose", "featured": True},
    {"id": "paper-analysis", "number": "02", "label": "卷面学情分析", "tagline": "上传试卷生成诊断报告与 Word", "category": "学情诊断", "route": "paper", "mode": "analysis", "delivery": "report"},
    {"id": "paper-variant", "number": "03", "label": "试题变式生成", "tagline": "生成可打印变式卷 Word", "category": "命题训练", "route": "tool", "mode": "variant", "delivery": "docx"},
    {"id": "ai-paper", "number": "04", "label": "智能批量命题", "tagline": "分层命题并导出 Word 卷", "category": "命题训练", "route": "tool", "mode": "question", "delivery": "docx"},
    {"id": "paper-word", "number": "05", "label": "试题文档整理", "tagline": "导出 Word / Markdown 文档", "category": "文档整理", "route": "tool", "mode": "document", "delivery": "docx"},
    {"id": "image-teacher", "number": "06", "label": "教案配图生成", "tagline": "直接生成教学配图 PNG", "category": "配图生成", "route": "tool", "mode": "image", "delivery": "image"},
    {"id": "ppt-review", "number": "07", "label": "讲评课件生成", "tagline": "导出可编辑 PowerPoint 课件", "category": "课件辅助", "route": "tool", "mode": "ppt", "delivery": "pptx"},
    {"id": "aippt-online3", "number": "08", "label": "在线生成3（测试）", "tagline": "AiPPT 预装版 iframe · PC 在线生成课件", "category": "课件辅助", "route": "aippt", "mode": "aippt", "delivery": "iframe"},
    {"id": "review-skill", "number": "09", "label": "读题破题训练", "tagline": "审题训练方案 Word 导出", "category": "审题训练", "route": "tool", "mode": "review", "delivery": "docx"},
    {"id": "big-question", "number": "10", "label": "主观题采分拆解", "tagline": "采分点拆解 Word 导出", "category": "解题拆解", "route": "tool", "mode": "decompose", "delivery": "docx"},
    {"id": "word-paper", "number": "11", "label": "词汇练习组卷", "tagline": "词汇卷 Word 导出", "category": "语言训练", "route": "tool", "mode": "english", "delivery": "docx"},
    {"id": "coverage-check", "number": "12", "label": "备考覆盖扫描", "tagline": "考点覆盖报告 Word 导出", "category": "备考规划", "route": "tool", "mode": "coverage", "delivery": "docx"},
    {"id": "sprint-plan", "number": "13", "label": "临考冲刺规划", "tagline": "冲刺方案 Word 导出", "category": "备考规划", "route": "tool", "mode": "plan", "delivery": "docx"},
    {"id": "class-notes", "number": "14", "label": "授课纪要整理", "tagline": "课堂笔记 Word 导出", "category": "课件辅助", "route": "tool", "mode": "notes", "delivery": "docx"},
    {"id": "preview-sheet", "number": "15", "label": "课前预习助手", "tagline": "预习单 Word 导出", "category": "预习辅助", "route": "tool", "mode": "preview", "delivery": "docx"},
    {"id": "loss-analysis", "number": "16", "label": "失分原因诊断", "tagline": "失分诊断报告 Word 导出", "category": "学情诊断", "route": "tool", "mode": "loss", "delivery": "docx"},
    {"id": "question-sense", "number": "17", "label": "题型直觉练习", "tagline": "题感训练 Word 导出", "category": "审题训练", "route": "tool", "mode": "sense", "delivery": "docx"},
    {"id": "knowledge-map", "number": "18", "label": "知识图谱绘制", "tagline": "导出 HTML 图谱与配图 PNG", "category": "配图生成", "route": "tool", "mode": "map", "delivery": "map"},
    {"id": "ten-solutions", "number": "19", "label": "多路径解法探索", "tagline": "多解法方案 Word 导出", "category": "解题拆解", "route": "tool", "mode": "multi", "delivery": "docx"},
    {"id": "knowledge-explain", "number": "20", "label": "概念精讲助手", "tagline": "讲解稿 Word 导出", "category": "概念讲解", "route": "tool", "mode": "explain", "delivery": "docx"},
    {"id": "essay-polish", "number": "21", "label": "表达润色教练", "tagline": "润色稿 Word 导出", "category": "语言训练", "route": "tool", "mode": "writing", "delivery": "docx"},
    {"id": "score-action", "number": "22", "label": "提分行动路线图", "tagline": "行动方案 Word 导出", "category": "备考规划", "route": "tool", "mode": "score", "delivery": "docx"},
]

# 历史版本曾以错误编码写入中文常量。这里使用规范中文作为唯一线上工具目录，
# 保持既有 id/mode/route，避免只在前端掩盖乱码。
_TOOL_ZH = {
    "wrong-transfer": ("错题八步拆解", "作答还原、错因诊断、母题模型与巩固过关", "错题突破"),
    "paper-analysis": ("全卷学情分析", "上传试卷生成逐题诊断与分析报告", "学情诊断"),
    "paper-variant": ("试卷变式机", "基于原卷生成同型变式训练卷", "命题训练"),
    "ai-paper": ("AI分层出题", "按知识点、难度和题量智能命题", "命题训练"),
    "paper-word": ("题卷重排Word", "整理题卷并导出可编辑Word", "文档整理"),
    "image-teacher": ("教学配图生成", "生成课堂、讲义和课件配图", "配图生成"),
    "ppt-review": ("试卷讲评课件", "生成可编辑的讲评PowerPoint", "课件辅助"),
    "aippt-online3": ("AI课件在线生成", "在线生成并编辑教学课件", "课件辅助"),
    "review-skill": ("审题破题训练", "训练题干识别、条件提取和切入点", "审题训练"),
    "big-question": ("主观题采分拆解", "按步骤拆出规范答案与评分点", "解题拆解"),
    "word-paper": ("英语词汇组卷", "生成词汇练习与答案Word", "语言训练"),
    "coverage-check": ("考点覆盖扫描", "检查试卷考点、难度与遗漏", "备考规划"),
    "sprint-plan": ("临考冲刺规划", "依据薄弱项生成冲刺计划", "备考规划"),
    "class-notes": ("课堂纪要整理", "把课堂材料整理为结构化笔记", "课件辅助"),
    "preview-sheet": ("课前预习助手", "生成目标、问题链和预习单", "预习辅助"),
    "loss-analysis": ("失分原因诊断", "归类知识、方法、计算与表达失分", "学情诊断"),
    "question-sense": ("题型直觉训练", "训练题型辨识和母题调用", "审题训练"),
    "knowledge-map": ("知识图谱绘制", "生成知识关系图谱与配图", "配图生成"),
    "ten-solutions": ("一题多解探索", "比较多种路径、条件与适用边界", "解题拆解"),
    "knowledge-explain": ("概念精讲助手", "生成例题、反例与通俗讲解", "概念讲解"),
    "essay-polish": ("表达润色教练", "改进作文表达并说明修改依据", "语言训练"),
    "score-action": ("提分行动路线图", "把学情诊断转成每日行动清单", "备考规划"),
}
PORTAL_TOOLS = [
    {**tool, "label": _TOOL_ZH[tool["id"]][0], "tagline": _TOOL_ZH[tool["id"]][1], "category": _TOOL_ZH[tool["id"]][2]}
    for tool in PORTAL_TOOLS
]


def get_portal_tool(tool_id: str) -> dict | None:
    return next((tool for tool in PORTAL_TOOLS if tool["id"] == tool_id), None)


OCR_PROMPT = """你是一个严谨的中高考全科 OCR 助手。

任务：只从图片中识别题目文字、公式、图表、材料、选项、学生手写/批注信息，不要解题。

要求：
1. 保留题号、已知条件、求解目标。
2. 数学/物理/化学公式尽量转写成 LaTeX 或清晰纯文本。
3. 语文/英语/政治/历史/地理等材料题要完整保留材料、设问和选项。
4. 如果图片中有学生答案、红叉、圈画、老师批注，请单独列出。
5. 无法确定的字符用 [?] 标记，不要编造。
6. 输出 JSON，不要输出 Markdown。

JSON 格式：
{
  "ocr_text": "完整识别文本",
  "printed_question": "印刷题干",
  "student_work": "学生作答或空字符串",
  "teacher_marks": "批改痕迹或空字符串",
  "uncertain_parts": ["不确定内容1"]
}
"""


DIAGNOSIS_PROMPT = """你是一个中高考全科拆题教练，不是拍照搜答案工具。

产品目标：
把一道题拆得非常细，让学生知道“题目如何被拆开、应该套哪个答题/解题模型、有没有对应题型原型/母题雏形、还能怎么解、最后用有趣小诗复盘”。

当前阶段说明：
- 暂时不做 RAG。
- 暂时不做正式母题库检索。
- 数理学科可以输出“母题雏形/母题归纳”；文科和语言学科输出“题型原型/答题模型雏形”。
- 该字段必须标记为 prompt_reserved，表示后续会接入母题/题型模型接口。
- 不要假装查到了题库。

必须体现三大亮点：
1. 拆题、答题/解题模型、题型原型/母题雏形。
2. 一题多解/多视角：至少给出 2 种解法、答题路径或思考视角；如果题目不适合多解，说明原因并给出 1 个替代视角。
3. 解完来点趣味：用生活化比喻拆解全过程，并写一首小诗/口诀，再逐句复盘。

科目适配：
- 数学：强调判型、拆公式、解题模型、母题雏形、一题多解。
- 物理/化学：强调情境建模、已知量/未知量、公式选择、实验/守恒/反应模型。
- 语文：强调材料拆读、设问类型、答题模板、采分点、语言组织。
- 英语：强调题型、定位句、语法/语义线索、选项排除、表达模板。
- 政史地生：强调材料信息、概念调用、因果链、答题角度和规范表述。

输出必须是合法 JSON，不要 Markdown，不要代码块，不要额外解释。

JSON 格式：
{
  "cleaned_question": "修正后的题干",
  "subject": "识别或用户指定的学科",
  "topic": "专题名称",
  "difficulty": 1,
  "confidence": 0.85,
  "core_pattern": "标准题型/母题雏形/答题模型名称",
  "knowledge_points": ["知识点1"],
  "problem_goal": "这题最终要求什么",
  "student_answer_analysis": {
    "answer_presence": "未提供/只给结论/有过程/有批改痕迹",
    "extracted_work": "从学生作答或批注中整理出的关键作答内容",
    "answer_status": "空白不会/思路卡住/步骤错误/计算错误/概念误用/表达不规范/基本正确但不完整",
    "likely_issue": "最可能的问题诊断",
    "evidence": ["依据1", "依据2"],
    "next_action": "下一步最应该补的动作"
  },
  "learning_strategy": {
    "decomposition_answer": "直接回答：这道题到底怎么拆解，按什么顺序拆",
    "make_it_easier": "直接回答：用什么方法可以让这题更容易学",
    "entry_point": "学生第一眼应该先抓哪个入口",
    "cognitive_ladder": ["先会什么", "再会什么", "最后迁移什么"],
    "micro_drills": ["1分钟小练习1", "1分钟小练习2"],
    "teacher_hint": "老师/产品引导学生时最该问的一句话"
  },
  "decomposition": {
    "total_formula": "识别题型→拆条件→选模型→计算→验证→总结",
    "step_formulas": [
      {
        "name": "判型公式",
        "formula": "看到...→判定为...",
        "operation": "这一步具体做什么",
        "student_trap": "学生容易错在哪里"
      }
    ]
  },
  "fun_analogy": {
    "theme": "生活化比喻主题，例如拆快递/造房子/破案/做菜",
    "overview": "一句话说明这个比喻如何对应题目",
    "steps": [
      {
        "step": "步骤名",
        "analogy": "有趣比喻",
        "math_action": "对应数学动作"
      }
    ]
  },
  "solution_models": [
    {
      "model_name": "标准化通用解题/答题模型名称",
      "applies_when": "适用条件",
      "steps": ["步骤1", "步骤2"],
      "checkpoints": ["检查点1"],
      "common_mistakes": ["常见错误1"]
    }
  ],
  "mother_question_reserved": {
    "status": "prompt_reserved",
    "name": "母题雏形/题型原型名称",
    "abstract_pattern": "去掉数字、材料背景后的抽象题型或答题模型",
    "recognition_signals": ["识别信号1"],
    "future_interface_hint": "后续可接入 /api/mother-questions 做正式沉淀"
  },
  "multiple_solutions": [
    {
      "method_name": "方法一名称",
      "idea": "核心思路",
      "steps": ["步骤1", "步骤2"],
      "pros_cons": "优缺点"
    },
    {
      "method_name": "方法二名称",
      "idea": "核心思路",
      "steps": ["步骤1", "步骤2"],
      "pros_cons": "优缺点"
    }
  ],
  "standard_answer": {
    "final_answer": "最终答案",
    "concise_solution": "标准简洁解答"
  },
  "poem": {
    "title": "小诗标题",
    "lines": ["诗句1", "诗句2", "诗句3"],
    "line_reviews": [
      {
        "line": "诗句1",
        "review": "这句对应哪一步解题操作"
      }
    ]
  },
  "practice_variants": [
    {
      "level": 1,
      "title": "同结构巩固",
      "stem": "变式题题干",
      "answer": "答案",
      "analysis": "解析"
    },
    {
      "level": 2,
      "title": "条件替换",
      "stem": "变式题题干",
      "answer": "答案",
      "analysis": "解析"
    },
    {
      "level": 3,
      "title": "迁移挑战",
      "stem": "变式题题干",
      "answer": "答案",
      "analysis": "解析"
    }
  ]
}

质量要求：
- 拆题必须细，不能只写泛泛步骤。
- 必须识别学生作答情况：如果用户提供了学生答案、错解、手写过程、批注、红叉或口头卡点，要单独归纳 student_answer_analysis；如果没有提供，answer_presence 写“未提供”，不要编造。
- 必须单独回答两个产品问题：“怎么拆解？”和“用什么方法，可以让这个题更容易学？”，写入 learning_strategy，不要只散落在解析里。
- learning_strategy 要像老师指导学生一样具体：入口、认知台阶、微练习、引导问题都要能直接执行。
- 比喻要贴合题目结构，不要硬搞笑。
- 小诗/口诀必须能反过来复盘解题或答题模型。
- 不要轻易拒答。若题干是常见中高考表达但略有省略，请按最常见考试语义合理补全，并在 cleaned_question 中说明“按常规理解为...”。
- 只有当题干严重缺失、完全无法判断要求时，才返回 needs_review=true。
"""



AGENT_LAYERS = [
    {
        "key": "input_recognition",
        "name": "01 输入识别层",
        "role": "接收图片、文件、粘贴题干与学生作答，区分题目、作答、批注和用户诉求。",
        "quality_gate": "题干与作答信息分栏清楚，缺失处明确标注。",
    },
    {
        "key": "question_structuring",
        "name": "02 题目结构化层",
        "role": "把题目拆成已知条件、设问目标、材料/图表/选项、隐含限制。",
        "quality_gate": "结构字段可直接服务后续推理，不能只复述题干。",
    },
    {
        "key": "subject_routing",
        "name": "03 学科路由层",
        "role": "自动判断学科、题型和知识板块，选择对应的解题语言与评分标准。",
        "quality_gate": "给出学科判断依据，允许用户指定学科覆盖自动判断。",
    },
    {
        "key": "exam_translation",
        "name": "04 审题翻译层",
        "role": "把考试语言翻译成学生听得懂的任务语言，指出第一眼抓什么入口。",
        "quality_gate": "回答“这题到底让我做什么”和“怎么拆解”。",
    },
    {
        "key": "solution_planning",
        "name": "05 解题规划层",
        "role": "选择最稳的解题/答题模型，安排步骤、公式、材料依据和检查点。",
        "quality_gate": "路径可执行，包含至少一个让题目更容易学的降阶方法。",
    },
    {
        "key": "step_solving",
        "name": "06 分步求解层",
        "role": "按计划完成规范解答，关键步骤给出理由，不跳步。",
        "quality_gate": "结论、过程和评分点互相一致。",
    },
    {
        "key": "answer_verification",
        "name": "07 答案校验层",
        "role": "检查计算、逻辑、单位、选项、材料引用和边界条件。",
        "quality_gate": "至少给出一个反查或代入验证动作。",
    },
    {
        "key": "teaching_explanation",
        "name": "08 教学讲解层",
        "role": "用生活化比喻、拆解公式和通用模型讲给学生听。",
        "quality_gate": "讲法要能降低认知负荷，而不是只换一种说法。",
    },
    {
        "key": "error_diagnosis",
        "name": "09 错因诊断层",
        "role": "对照学生作答定位错因、断点、证据和下一步补救动作。",
        "quality_gate": "没有作答时不编造错因，有作答时必须引用证据。",
    },
    {
        "key": "practice_generation",
        "name": "10 训练生成层",
        "role": "生成同结构巩固、条件替换、迁移挑战三类训练。",
        "quality_gate": "训练题覆盖核心能力，不只是换数字。",
    },
    {
        "key": "archive_update",
        "name": "11 档案沉淀层",
        "role": "沉淀题目、答案、解析、详细思路、同类题、错因和复盘小诗。",
        "quality_gate": "结果可进入分科学习档案和历史记录。",
    },
]


AGENT_SOLVE_PROMPT = """你是“AI错题拆博士”的高考全科解题智能体总控。

你的任务不是简单给答案，而是把一道题按产品化智能体层次跑完：识别、结构化、路由、审题、规划、求解、校验、讲解、诊断、训练、归档。

请严格输出合法 JSON，不要 Markdown，不要代码块，不要额外解释。

必须覆盖全部学科：数学、物理、化学、生物、语文、英语、政治、历史、地理及其他考试型题目。若用户指定学科，以用户指定为准；否则自动判断。

核心要求：
1. 直接回答“怎么拆解？”和“用什么方法，可以让这个题更容易学？”
2. 如果有学生作答，必须识别作答情况、错因证据和补救动作；如果没有作答，不要编造。
3. 至少给出 2 种方法/视角；若题目不适合多解，要说明并给替代视角。
4. 必须有标准答案、规范解析、评分点、易错点、同类训练、小诗/口诀复盘。
5. 暂不做正式 RAG 和母题库检索，但要预留 mother_question_reserved 字段，标记 status 为 prompt_reserved。

固定层次必须全部返回，key 必须逐一对应：
input_recognition, question_structuring, subject_routing, exam_translation, solution_planning,
step_solving, answer_verification, teaching_explanation, error_diagnosis, practice_generation, archive_update。

JSON 格式：
{
  "title": "题目短标题",
  "subject": "学科",
  "question_type": "题型/专题",
  "difficulty": 3,
  "confidence": 0.86,
  "quick_answer": {
    "how_to_decompose": "这道题怎么拆解",
    "make_it_easier": "用什么方法让这题更容易学",
    "first_entry": "第一眼入口"
  },
  "layers": [
    {
      "key": "input_recognition",
      "name": "01 输入识别层",
      "status": "done",
      "summary": "本层结论",
      "input": "本层读取的信息",
      "output": "本层产出",
      "quality_gate": "本层质检门",
      "next_action": "下一步动作"
    }
  ],
  "structured_question": {
    "cleaned_question": "修正后的题干",
    "known_conditions": ["条件1"],
    "target": "求什么/答什么",
    "hidden_constraints": ["隐含条件1"],
    "student_work": "学生作答/空字符串"
  },
  "student_answer_analysis": {
    "answer_presence": "未提供/只给结论/有过程/有批注",
    "answer_status": "空白不会/思路卡住/步骤错误/计算错误/概念误用/表达不规范/基本正确但不完整",
    "likely_issue": "最可能错因",
    "evidence": ["证据1"],
    "next_action": "补救动作"
  },
  "solution_model": {
    "model_name": "通用解题/答题模型名",
    "applies_when": "适用条件",
    "step_formula": "识别题型→拆条件→选模型→执行→校验→复盘",
    "steps": ["步骤1", "步骤2"],
    "checkpoints": ["检查点1"]
  },
  "multiple_solutions": [
    {"method_name": "方法一", "idea": "核心思路", "steps": ["步骤1"], "pros_cons": "优缺点"},
    {"method_name": "方法二", "idea": "核心思路", "steps": ["步骤1"], "pros_cons": "优缺点"}
  ],
  "standard_solution": "规范解析正文",
  "final_answer": "最终答案",
  "score_points": ["采分点1"],
  "common_mistakes": ["易错点1"],
  "mother_question_reserved": {
    "status": "prompt_reserved",
    "name": "母题雏形/题型原型",
    "abstract_pattern": "抽象题型",
    "future_interface_hint": "后续接入 /api/mother-questions 或 RAG"
  },
  "fun_analogy": {
    "theme": "比喻主题",
    "overview": "比喻说明",
    "steps": [{"step": "步骤名", "analogy": "比喻", "action": "对应操作"}]
  },
  "poem": {
    "title": "口诀/小诗标题",
    "lines": ["诗句1", "诗句2"],
    "line_reviews": [{"line": "诗句1", "review": "对应哪一步"}]
  },
  "training_tasks": [
    {"level": 1, "title": "同结构巩固", "stem": "题干", "answer": "答案", "analysis": "解析"},
    {"level": 2, "title": "条件替换", "stem": "题干", "answer": "答案", "analysis": "解析"},
    {"level": 3, "title": "迁移挑战", "stem": "题干", "answer": "答案", "analysis": "解析"}
  ],
  "archive_payload": {
    "subject": "学科",
    "question": "题目",
    "answer": "答案",
    "analysis": "解析",
    "detailed_thinking": "详细思路",
    "similar_questions": ["同类题简述"],
    "tags": ["标签1"]
  }
}
"""

GRADING_PROMPT = """你是中高考全科批改老师。

任务：批改学生对一道巩固题的答案，判断是否掌握对应解题/答题模型。

输出必须是合法 JSON：
{
  "is_correct": true,
  "score": 90,
  "comment": "总体评价",
  "detected_issue": "主要问题或空字符串",
  "reference_answer": "标准答案",
  "analysis": "关键步骤讲解",
  "next_advice": "下一步训练建议",
  "poem_review": "用一句轻松小诗或口诀帮学生记住本题"
}
"""


def now_iso() -> str:
    return datetime.now().isoformat(timespec="seconds")


def ensure_dirs() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    CARD_DIR.mkdir(parents=True, exist_ok=True)
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)


def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def hash_password(password: str, salt: str | None = None) -> str:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120_000).hex()
    return f"pbkdf2_sha256${salt}${digest}"


def verify_password(password: str, stored: str) -> bool:
    try:
        method, salt, digest = stored.split("$", 2)
    except ValueError:
        return False
    if method != "pbkdf2_sha256":
        return False
    return secrets.compare_digest(hash_password(password, salt).split("$", 2)[2], digest)


def bearer_token(headers) -> str:
    auth = headers.get("Authorization", "")
    if auth.lower().startswith("bearer "):
        return auth.split(" ", 1)[1].strip()
    return ""


def is_loopback_url(url: str) -> bool:
    try:
        host = (urllib.parse.urlparse(str(url or "")).hostname or "").lower()
        return host in {"127.0.0.1", "localhost", "::1"}
    except Exception:
        return False


def public_auth_base_url() -> str:
    if PUBLIC_AUTH_BASE_URL and not is_loopback_url(PUBLIC_AUTH_BASE_URL):
        return PUBLIC_AUTH_BASE_URL
    if AUTH_BASE_URL and not is_loopback_url(AUTH_BASE_URL):
        return AUTH_BASE_URL
    return "https://ai.ms1001.com"


def auth_base_url() -> str:
    return AUTH_BASE_URL or public_auth_base_url()


def unified_auth_enabled() -> bool:
    return bool(USE_UNIFIED_AUTH and auth_base_url())


def verify_unified_token_remote(token: str) -> dict | None:
    if not token or not unified_auth_enabled():
        return None
    url = f"{auth_base_url()}/api/auth/verify"
    req = urllib.request.Request(
        url,
        headers={"Authorization": f"Bearer {token}", "Accept": "application/json"},
        method="GET",
    )
    try:
        with urllib.request.urlopen(req, timeout=8) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except Exception:
        return None
    if not data.get("valid"):
        return None
    user = data.get("user") or {}
    claims = data.get("claims") or {}
    return {
        "username": user.get("username") or claims.get("sub") or claims.get("username") or "",
        "email": user.get("email") or claims.get("email") or "",
        "name": user.get("name") or claims.get("name") or "",
        "phone": user.get("phone") or claims.get("phone") or "",
        "role": user.get("role") or claims.get("role") or "USER",
        "tier": user.get("tier") or claims.get("tier") or "standard",
        "isOwner": user.get("isOwner") if user.get("isOwner") is not None else claims.get("isOwner"),
        "platformPermissions": user.get("platformPermissions") or claims.get("platformPermissions") or {},
        "platformMemberships": user.get("platformMemberships") or claims.get("platformMemberships") or {},
    }


def unified_org_user(headers) -> dict | None:
    token = bearer_token(headers)
    if not token:
        return None
    claims = verify_unified_token_remote(token)
    if not claims:
        return None
    return {
        "username": str(claims.get("username") or claims.get("sub") or "").strip(),
        "name": str(claims.get("name") or "").strip(),
        "role": str(claims.get("role") or "USER").upper(),
    }


org_inst_store.set_unified_auth_resolver(unified_org_user)


def unified_email_from_claims(claims: dict) -> str:
    email = str(claims.get("email") or "").strip().lower()
    if re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email):
        return email
    username = str(claims.get("username") or "").strip()
    phone = re.sub(r"\D", "", str(claims.get("phone") or username or ""))
    if re.fullmatch(r"1\d{10}", phone):
        return f"{phone}@unified.ms1001.com"
    safe = re.sub(r"[^a-zA-Z0-9_.-]+", "-", username or str(claims.get("name") or "user")).strip("-").lower()
    return f"{safe or 'user'}@unified.ms1001.com"


def public_user(row: sqlite3.Row | dict | None) -> dict | None:
    if not row:
        return None
    item = dict(row)
    return {
        "id": item.get("id"),
        "email": item.get("email"),
        "display_name": item.get("display_name") or "",
        "membership_tier": item.get("membership_tier") or "standard",
        "membership_expires_at": item.get("membership_expires_at"),
        "credits": int(item.get("credits") or 0),
        "created_at": item.get("created_at"),
        "is_admin": bool(int(item.get("is_admin") or 0)),
    }


def user_is_admin(user: dict | None) -> bool:
    return bool(user and user.get("is_admin"))


def user_owns_wrong_question(conn: sqlite3.Connection, user: dict | None, wrong_id: str) -> bool:
    row = conn.execute("select user_id from wrong_questions where id = ?", (wrong_id,)).fetchone()
    if not row:
        return False
    owner = row["user_id"]
    if not owner or not str(owner).strip():
        return True
    if not user:
        return False
    if user_is_admin(user):
        return True
    return owner == user["id"]


def user_owns_tool_run(user: dict | None, run: dict | None) -> bool:
    if not user or not run:
        return False
    if user_is_admin(user):
        return True
    return run.get("user_id") == user["id"]


def current_user_from_token(conn: sqlite3.Connection, token: str) -> dict | None:
    if not token:
        return None
    row = conn.execute(
        """
        select u.* from app_sessions s
        join app_users u on u.id = s.user_id
        where s.token = ?
        """,
        (token,),
    ).fetchone()
    return public_user(row)


def ensure_user_from_unified_auth(conn: sqlite3.Connection, claims: dict) -> dict | None:
    if not claims:
        return None
    email = unified_email_from_claims(claims)
    row = conn.execute("select * from app_users where email = ?", (email,)).fetchone()
    if not row:
        user_id = str(uuid.uuid4())
        display_name = str(claims.get("name") or claims.get("username") or "").strip()
        membership_tier = "special" if str(claims.get("tier") or "").lower() in {"vip", "admin"} else "standard"
        conn.execute(
            "insert into app_users (id, email, password_hash, credits, display_name, membership_tier, created_at) values (?, ?, ?, ?, ?, ?, ?)",
            (user_id, email, hash_password(secrets.token_urlsafe(24)), 9999 if membership_tier == "special" else 9, display_name, membership_tier, now_iso()),
        )
        row = conn.execute("select * from app_users where id = ?", (user_id,)).fetchone()
    display_name = str(claims.get("name") or claims.get("username") or "").strip()
    if display_name and row and not row["display_name"]:
        conn.execute("update app_users set display_name = ? where id = ?", (display_name, row["id"]))
        row = conn.execute("select * from app_users where id = ?", (row["id"],)).fetchone()
    return public_user(row)


def current_user_from_request(conn: sqlite3.Connection, headers) -> dict | None:
    token = bearer_token(headers)
    if token and unified_auth_enabled():
        claims = verify_unified_token_remote(token)
        if claims:
            return ensure_user_from_unified_auth(conn, claims)
    return current_user_from_token(conn, token)


_MOJIBAKE_RE = re.compile(r"(?:Ã.|Â.|â.|ã€|ï¼|å[^\u4e00-\u9fff]|æ[^\u4e00-\u9fff]|ç[^\u4e00-\u9fff]|è[^\u4e00-\u9fff]|ä[^\u4e00-\u9fff]|é[^\u4e00-\u9fff]|�)")


def _mojibake_score(text: str) -> int:
    return len(_MOJIBAKE_RE.findall(text or ""))


def repair_mojibake_text(value: object) -> str:
    text = str(value or "")
    if not text or _mojibake_score(text) == 0:
        return text
    current = text
    for _ in range(2):
        candidates = []
        try:
            mixed = bytearray()
            for char in current:
                if ord(char) <= 255:
                    mixed.append(ord(char))
                else:
                    mixed.extend(char.encode("cp1252"))
            candidates.append(bytes(mixed).decode("utf-8"))
        except (UnicodeEncodeError, UnicodeDecodeError, ValueError):
            pass
        for encoding in ("cp1252", "latin1"):
            try:
                candidates.append(current.encode(encoding).decode("utf-8"))
            except (UnicodeEncodeError, UnicodeDecodeError):
                continue
        if not candidates:
            break
        candidate = min(candidates, key=lambda item: (_mojibake_score(item), -len(re.findall(r"[\u4e00-\u9fff]", item))))
        if _mojibake_score(candidate) >= _mojibake_score(current):
            break
        current = candidate
    return current


def repair_text_tree(value: object) -> object:
    if isinstance(value, str):
        return repair_mojibake_text(value)
    if isinstance(value, list):
        return [repair_text_tree(item) for item in value]
    if isinstance(value, tuple):
        return tuple(repair_text_tree(item) for item in value)
    if isinstance(value, dict):
        return {key: repair_text_tree(item) for key, item in value.items()}
    return value


def _replace_unhelpful_placeholder(text: str) -> str:
    if "未返回" not in text:
        return text
    if "总拆解公式" in text or "公式" in text:
        return "读题 → 提取条件 → 确定题型 → 分步作答 → 检查结论"
    if "最终答案" in text or "答案" in text:
        return "请依据完整题干与已知条件完成最终作答；当前结果需结合原题核对。"
    if "题目目标" in text or "目标" in text:
        return "明确题目设问，提取关键条件，并给出规范结论。"
    return "当前信息不足，请补充完整题干或作答过程后继续分析。"


def _replace_placeholder_tree(value: object) -> object:
    if isinstance(value, str):
        return _replace_unhelpful_placeholder(value)
    if isinstance(value, list):
        return [_replace_placeholder_tree(item) for item in value]
    if isinstance(value, dict):
        return {key: _replace_placeholder_tree(item) for key, item in value.items()}
    return value


def normalize_diagnosis_payload(payload: object, question_text: str = "", student_answer: str = "") -> dict:
    result = repair_text_tree(payload if isinstance(payload, dict) else {})
    result = dict(result)
    result["core_pattern"] = str(result.get("core_pattern") or result.get("topic") or "待归纳题型")
    result["subject"] = str(result.get("subject") or "自动识别")
    result["problem_goal"] = str(result.get("problem_goal") or "明确题目设问，提取关键条件，并给出规范结论。")
    points = result.get("knowledge_points")
    result["knowledge_points"] = points if isinstance(points, list) and points else ["题干信息提取与规范作答"]

    decomposition = result.get("decomposition") if isinstance(result.get("decomposition"), dict) else {}
    decomposition["total_formula"] = str(decomposition.get("total_formula") or "读题 → 提取条件 → 确定题型 → 分步作答 → 检查结论")
    steps = decomposition.get("step_formulas") if isinstance(decomposition.get("step_formulas"), list) else []
    if not steps:
        steps = [
            {"name": "读题定位", "formula": "题型 + 已知条件 + 设问", "operation": "圈出关键词，明确最终要回答什么"},
            {"name": "选择方法", "formula": "知识点 → 解题模型", "operation": "根据触发特征选择公式、依据或答题模板"},
            {"name": "规范作答", "formula": "依据 → 步骤 → 结论", "operation": "逐步写清推理、计算与评分点"},
            {"name": "检查迁移", "formula": "条件 + 单位 + 边界 + 结论", "operation": "核对结果，并总结同类题的识别入口"},
        ]
    decomposition["step_formulas"] = steps
    result["decomposition"] = decomposition

    standard = result.get("standard_answer") if isinstance(result.get("standard_answer"), dict) else {}
    standard["final_answer"] = str(standard.get("final_answer") or "请依据完整题干与已知条件完成最终作答；当前结果需结合原题核对。")
    standard["concise_solution"] = str(standard.get("concise_solution") or "先提取已知条件和设问，再选择对应知识点与方法，分步完成并检查结论。")
    scoring = standard.get("scoring_points") if isinstance(standard.get("scoring_points"), list) else []
    standard["scoring_points"] = scoring or ["识别条件与设问", "写出关键依据或公式", "完成推导并给出结论"]
    result["standard_answer"] = standard

    answer = result.get("student_answer_analysis") if isinstance(result.get("student_answer_analysis"), dict) else {}
    answer["answer_presence"] = str(answer.get("answer_presence") or ("已提供" if student_answer else "未提供"))
    answer["extracted_work"] = str(answer.get("extracted_work") or student_answer or "")
    answer["answer_status"] = str(answer.get("answer_status") or ("需要结合题干核对" if student_answer else "尚未提供作答"))
    answer["likely_issue"] = str(answer.get("likely_issue") or ("请对照规范步骤定位最先偏离的位置" if student_answer else "补充学生答案后可进一步判断错因"))
    answer["evidence"] = answer.get("evidence") if isinstance(answer.get("evidence"), list) else []
    answer["next_action"] = str(answer.get("next_action") or ("对照分步路径复盘并完成一道同类题" if student_answer else "补充作答过程或批改痕迹"))
    result["student_answer_analysis"] = answer

    strategy = result.get("learning_strategy") if isinstance(result.get("learning_strategy"), dict) else {}
    strategy["decomposition_answer"] = str(strategy.get("decomposition_answer") or decomposition["total_formula"])
    strategy["make_it_easier"] = str(strategy.get("make_it_easier") or "先完成一个更小的同型题，掌握关键步骤后再回到原题。")
    strategy["entry_point"] = str(strategy.get("entry_point") or result["core_pattern"])
    strategy["teacher_hint"] = str(strategy.get("teacher_hint") or "这道题的设问要求什么？哪一个条件最先决定解题方法？")
    strategy["cognitive_ladder"] = strategy.get("cognitive_ladder") if isinstance(strategy.get("cognitive_ladder"), list) and strategy.get("cognitive_ladder") else ["读懂设问", "识别题型", "完成分步作答"]
    strategy["micro_drills"] = strategy.get("micro_drills") if isinstance(strategy.get("micro_drills"), list) and strategy.get("micro_drills") else ["用一句话概括题型", "只写第一步依据或公式"]
    result["learning_strategy"] = strategy

    result.setdefault("cleaned_question", question_text)
    result.setdefault("practice_variants", [])
    return _replace_placeholder_tree(repair_text_tree(result))


def text_quality_issues(value: object) -> list[str]:
    text = json.dumps(value, ensure_ascii=False) if not isinstance(value, str) else value
    issues = []
    if _mojibake_score(text) or "�" in text:
        issues.append("mojibake")
    if "未返回" in text:
        issues.append("unhelpful_placeholder")
    if any(ord(char) < 32 and char not in "\n\r\t" for char in text):
        issues.append("control_character")
    return issues


def row_to_dict(row: sqlite3.Row | None) -> dict | None:
    if row is None:
        return None
    item = dict(row)
    json_fields = {
        "diagnosis",
        "grading_result",
        "metadata",
        "artifacts",
        "report",
        "result",
        "recognition_signals",
        "knowledge_points",
        "solution_steps",
        "common_error_causes",
    }
    for key in list(item.keys()):
        if key in json_fields and isinstance(item[key], str) and item[key]:
            try:
                item[key] = json.loads(item[key])
            except json.JSONDecodeError:
                pass
    item = repair_text_tree(item)
    if isinstance(item.get("diagnosis"), dict):
        item["diagnosis"] = normalize_diagnosis_payload(
            item["diagnosis"], str(item.get("corrected_text") or item.get("ocr_text") or ""), str(item.get("student_wrong_answer") or "")
        )
    return item


def archive_legacy_table(conn: sqlite3.Connection, table: str, required_columns: set[str]) -> None:
    exists = conn.execute(
        "select name from sqlite_master where type='table' and name = ?",
        (table,),
    ).fetchone()
    if not exists:
        return
    cols = {row["name"] for row in conn.execute(f"pragma table_info({table})").fetchall()}
    if required_columns.issubset(cols):
        return
    archived = f"{table}_legacy_{datetime.now().strftime('%Y%m%d%H%M%S')}"
    conn.execute(f"alter table {table} rename to {archived}")


def init_db() -> None:
    ensure_dirs()
    with db() as conn:
        archive_legacy_table(conn, "wrong_questions", {"id", "ocr_text", "diagnosis", "model_id"})
        archive_legacy_table(conn, "exercise_variants", {"id", "wrong_question_id", "title", "stem", "answer"})
        archive_legacy_table(conn, "student_answers", {"id", "exercise_variant_id", "grading_result"})
        conn.executescript(
            """
            create table if not exists model_configs (
              id text primary key,
              name text not null,
              provider text not null,
              endpoint text not null,
              model text not null,
              api_key text,
              supports_vision integer not null default 1,
              temperature real not null default 0.6,
              max_tokens integer not null default 6000,
              is_default integer not null default 0,
              created_at text not null,
              updated_at text not null
            );

            create table if not exists prompt_templates (
              key text primary key,
              name text not null,
              content text not null,
              updated_at text not null
            );

            create table if not exists image_model_configs (
              id text primary key,
              name text not null,
              provider text not null,
              endpoint text not null,
              model text not null,
              api_key text,
              size text not null default '1024x1536',
              quality text not null default 'high',
              is_default integer not null default 0,
              created_at text not null,
              updated_at text not null
            );

            create table if not exists exam_papers (
              id text primary key,
              user_id text,
              title text not null,
              subject text,
              status text not null,
              source_name text,
              summary text,
              progress integer not null default 0,
              error text,
              created_at text not null,
              updated_at text not null
            );

            create table if not exists paper_pages (
              id text primary key,
              paper_id text not null,
              page_no integer not null,
              source_url text,
              source_text text,
              ocr_result text,
              confidence real not null default 0,
              foreign key(paper_id) references exam_papers(id)
            );

            create table if not exists paper_questions (
              id text primary key,
              paper_id text not null,
              page_id text,
              question_no text not null,
              printed_text text not null,
              student_work text,
              teacher_marks text,
              answer_state text not null,
              score real,
              max_score real,
              confidence real not null default 0,
              bbox text,
              eight_steps text,
              diagnosis text,
              wrong_question_id text,
              review_required integer not null default 0,
              created_at text not null,
              foreign key(paper_id) references exam_papers(id)
            );

            create table if not exists paper_jobs (
              id text primary key,
              paper_id text not null,
              status text not null,
              progress integer not null default 0,
              message text,
              attempts integer not null default 0,
              created_at text not null,
              updated_at text not null,
              foreign key(paper_id) references exam_papers(id)
            );

            create table if not exists wrong_questions (
              id text primary key,
              image_url text,
              ocr_text text,
              corrected_text text not null,
              student_wrong_answer text,
              model_id text,
              diagnosis text not null,
              status text not null,
              confidence real not null,
              created_at text not null
            );

            create table if not exists exercise_variants (
              id text primary key,
              wrong_question_id text not null,
              level integer not null,
              title text not null,
              stem text not null,
              answer text not null,
              analysis text not null,
              created_at text not null,
              foreign key(wrong_question_id) references wrong_questions(id)
            );

            create table if not exists student_answers (
              id text primary key,
              exercise_variant_id text not null,
              answer_text text not null,
              grading_result text not null,
              is_correct integer not null,
              submitted_at text not null,
              foreign key(exercise_variant_id) references exercise_variants(id)
            );

            create table if not exists study_cards (
              id text primary key,
              wrong_question_id text not null,
              image_model_id text,
              image_url text not null,
              prompt text not null,
              model text not null,
              size text not null,
              quality text not null,
              status text not null,
              created_at text not null,
              foreign key(wrong_question_id) references wrong_questions(id)
            );

            create table if not exists mother_questions (
              id text primary key,
              code text unique not null,
              name text not null,
              status text not null,
              metadata text,
              created_at text not null
            );

            create table if not exists app_users (
              id text primary key,
              email text unique not null,
              password_hash text not null,
              credits integer not null default 9,
              created_at text not null
            );

            create table if not exists app_sessions (
              token text primary key,
              user_id text not null,
              created_at text not null,
              foreign key(user_id) references app_users(id)
            );

            create table if not exists redeem_codes (
              code text primary key,
              credits integer not null,
              is_used integer not null default 0,
              used_by text,
              used_at text,
              created_at text not null
            );

            create table if not exists tool_runs (
              id text primary key,
              tool_id text not null,
              tool_label text not null,
              subject text,
              input_text text not null,
              output_text text not null,
              model_id text,
              user_id text,
              created_at text not null
            );

            create table if not exists profile_exports (
              id text primary key,
              filename text not null,
              subject text not null,
              markdown text not null,
              count integer not null,
              created_at text not null
            );

            create table if not exists profile_shares (
              id text primary key,
              token text unique not null,
              user_id text not null,
              title text not null,
              audience text,
              note text,
              status text not null,
              permissions text,
              created_at text not null,
              last_viewed_at text,
              expires_at text
            );

            create table if not exists agent_runs (
              id text primary key,
              user_id text,
              subject text,
              question_text text not null,
              student_answer text,
              model_id text,
              result text not null,
              status text not null,
              created_at text not null
            );

            create table if not exists rag_documents (
              id text primary key,
              user_id text,
              title text not null,
              filename text,
              subject text,
              source_type text not null,
              chars integer not null,
              created_at text not null
            );

            create table if not exists rag_chunks (
              id text primary key,
              document_id text not null,
              user_id text,
              chunk_index integer not null,
              title text not null,
              subject text,
              content text not null,
              chars integer not null,
              created_at text not null,
              foreign key(document_id) references rag_documents(id)
            );

            create virtual table if not exists rag_chunks_fts using fts5(
              chunk_id unindexed,
              document_id unindexed,
              user_id unindexed,
              title,
              subject,
              content,
              tokenize = 'unicode61'
            );

            create table if not exists question_method_packages (
              id text primary key,
              question_id text unique not null,
              title text,
              stem text not null,
              conditions text,
              target text,
              knowledge_points text,
              method_steps text,
              formulas text,
              pitfalls text,
              transfer_pattern text,
              source text,
              created_at text not null
            );

            create virtual table if not exists question_method_packages_fts using fts5(
              package_id unindexed,
              question_id unindexed,
              title,
              stem,
              knowledge_points,
              method_steps,
              formulas,
              pitfalls,
              transfer_pattern,
              tokenize = 'unicode61'
            );

            create table if not exists learning_profiles (
              user_id text primary key,
              role text not null,
              display_name text,
              class_id text,
              updated_at text not null
            );

            create table if not exists checkin_tasks (
              id text primary key,
              creator_user_id text not null,
              assignee_user_id text,
              class_id text,
              title text not null,
              description text,
              task_type text not null default 'study',
              due_at text,
              status text not null default 'active',
              created_at text not null
            );

            create table if not exists checkin_records (
              id text primary key,
              task_id text not null,
              user_id text not null,
              note text,
              evidence_url text,
              status text not null default 'submitted',
              created_at text not null,
              unique(task_id, user_id)
            );
            """
        )
        seed_defaults(conn)
        ensure_table_column(conn, "tool_runs", "artifacts", "text")
        ensure_table_column(conn, "tool_runs", "report", "text")
        ensure_table_column(conn, "app_users", "is_admin", "integer not null default 0")
        ensure_table_column(conn, "app_users", "display_name", "text")
        ensure_table_column(conn, "app_users", "membership_tier", "text not null default 'standard'")
        ensure_table_column(conn, "app_users", "membership_expires_at", "text")
        ensure_table_column(conn, "wrong_questions", "user_id", "text")
        ensure_table_column(conn, "wrong_questions", "institution_id", "text")
        ensure_table_column(conn, "wrong_questions", "institution_name", "text")
        ensure_table_column(conn, "wrong_questions", "institution_badge", "text")
        ensure_table_column(conn, "wrong_questions", "error_type", "text")
        ensure_table_column(conn, "wrong_questions", "mastery_score", "integer not null default 0")
        ensure_table_column(conn, "wrong_questions", "review_stage", "integer not null default 0")
        ensure_table_column(conn, "wrong_questions", "next_review_at", "text")
        ensure_table_column(conn, "wrong_questions", "last_reviewed_at", "text")
        ensure_table_column(conn, "wrong_questions", "workflow_state", "text not null default 'diagnosed'")
        ensure_table_column(conn, "student_answers", "hint_count", "integer not null default 0")
        ensure_table_column(conn, "profile_exports", "user_id", "text")
        ensure_table_column(conn, "profile_shares", "audience", "text")
        ensure_table_column(conn, "profile_shares", "note", "text")
        ensure_table_column(conn, "profile_shares", "permissions", "text")
        ensure_table_column(conn, "profile_shares", "last_viewed_at", "text")
        ensure_table_column(conn, "profile_shares", "expires_at", "text")
        ensure_table_column(conn, "mother_questions", "code", "text")
        ensure_table_column(conn, "mother_questions", "name", "text not null default '未命名母题'")
        ensure_table_column(conn, "mother_questions", "status", "text not null default 'review_required'")
        ensure_table_column(conn, "mother_questions", "metadata", "text")
        ensure_table_column(conn, "mother_questions", "created_at", "text not null default ''")
        ensure_gaokao_question_tables(conn)
        gaokao_rag.ensure_rag_schema(conn)
        backfill_question_method_packages(conn)
        ensure_table_column(conn, "exam_papers", "stage", "text not null default ''")
        ensure_admin_user(conn)


def ensure_table_column(conn: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    cols = {row["name"] for row in conn.execute(f"pragma table_info({table})").fetchall()}
    if column not in cols:
        conn.execute(f"alter table {table} add column {column} {definition}")


METHOD_KEYWORDS = ["函数", "导数", "数列", "向量", "概率", "排列", "组合", "三角", "圆锥曲线", "立体几何", "不等式", "集合", "复数", "解析几何", "统计"]


def derive_question_method_package(row: sqlite3.Row | dict) -> dict:
    item = dict(row)
    stem = compact_text(item.get("stem") or "", 2200)
    answer = compact_text(item.get("answer") or "", 1200)
    analysis = compact_text(item.get("analysis") or "", 2200)
    knowledge = [word for word in METHOD_KEYWORDS if word in f"{stem}{analysis}"]
    conditions = re.split(r"求|证明|判断|计算|若", stem, maxsplit=1)[0].strip(" ，。；")
    target_match = re.search(r"(?:求|证明|判断|计算|写出|确定)([^。；？?]{2,180})", stem)
    target = target_match.group(0).strip() if target_match else compact_text(stem[-180:], 180)
    formulas = "；".join(re.findall(r"(?:[A-Za-z][A-Za-z0-9_]*\s*=\s*[^，。；\n]{1,100}|\\(?:frac|sum|int|sqrt|sin|cos|tan)[^，。；\n]{0,100})", f"{analysis}\n{answer}")[:8])
    method_steps = analysis or answer or "先识别条件与设问，再匹配同型母题入口，逐步推导并回代检验。"
    pitfalls = "检查隐含条件、定义域、符号方向、分类边界和最终作答格式。"
    transfer = f"遇到包含“{'、'.join(knowledge) or '同类条件'}”的新题，先抽取已知量与目标，再复用本题的入口和步骤；改变数值、问法或情境后仍需重新校验边界。"
    return {
        "id": f"pkg-{item.get('id')}", "question_id": item.get("id"), "title": item.get("title") or item.get("question_no") or "题型方法包",
        "stem": stem, "conditions": conditions, "target": target, "knowledge_points": json.dumps(knowledge, ensure_ascii=False),
        "method_steps": method_steps, "formulas": formulas, "pitfalls": pitfalls, "transfer_pattern": transfer,
        "source": "既有题库自动拆包", "created_at": item.get("created_at") or now_iso(),
    }


def backfill_question_method_packages(conn: sqlite3.Connection, limit: int | None = None) -> dict:
    exists = conn.execute("select 1 from sqlite_master where type='table' and name='gaokao_questions'").fetchone()
    if not exists:
        return {"total": 0, "created": 0}
    sql = "select q.* from gaokao_questions q left join question_method_packages p on p.question_id=q.id where p.question_id is null order by q.created_at"
    params: list = []
    if limit:
        sql += " limit ?"; params.append(int(limit))
    rows = conn.execute(sql, params).fetchall()
    for row in rows:
        item = derive_question_method_package(row)
        conn.execute("""insert or ignore into question_method_packages
          (id,question_id,title,stem,conditions,target,knowledge_points,method_steps,formulas,pitfalls,transfer_pattern,source,created_at)
          values (?,?,?,?,?,?,?,?,?,?,?,?,?)""", tuple(item[key] for key in ("id","question_id","title","stem","conditions","target","knowledge_points","method_steps","formulas","pitfalls","transfer_pattern","source","created_at")))
        conn.execute("""insert into question_method_packages_fts
          (package_id,question_id,title,stem,knowledge_points,method_steps,formulas,pitfalls,transfer_pattern)
          values (?,?,?,?,?,?,?,?,?)""", (item["id"], item["question_id"], item["title"], item["stem"], item["knowledge_points"], item["method_steps"], item["formulas"], item["pitfalls"], item["transfer_pattern"]))
    total = conn.execute("select count(*) c from question_method_packages").fetchone()["c"]
    return {"total": int(total or 0), "created": len(rows)}


def search_question_method_packages(conn: sqlite3.Connection, query: str, limit: int = 6) -> list[dict]:
    tokens = re.findall(r"[0-9A-Za-z_]+|[\u4e00-\u9fff]{2,}", query or "")
    if not tokens:
        return []
    fts = " OR ".join(sorted(set(tokens), key=len, reverse=True)[:14])
    try:
        rows = conn.execute("""select p.*, bm25(question_method_packages_fts) rank
          from question_method_packages_fts join question_method_packages p on p.id=question_method_packages_fts.package_id
          where question_method_packages_fts match ? order by rank limit ?""", (fts, max(1, min(limit, 12)))).fetchall()
    except sqlite3.OperationalError:
        rows = []
    return [{**dict(row), "score": round(100 / (1 + abs(float(row["rank"] or 0))), 2)} for row in rows]


def method_packages_context(packages: list[dict]) -> str:
    return "\n\n".join(f"【方法包{i}｜{p.get('title')}】\n条件：{p.get('conditions')}\n目标：{p.get('target')}\n方法：{compact_text(p.get('method_steps'), 800)}\n公式：{p.get('formulas') or '按题型推导'}\n易错：{p.get('pitfalls')}\n迁移：{p.get('transfer_pattern')}" for i, p in enumerate(packages, 1))


LEARNING_ROLES = {"teacher": "老师", "parent": "家长", "student": "同学"}


def checkin_dashboard(conn: sqlite3.Connection, user: dict) -> dict:
    profile_row = conn.execute("select * from learning_profiles where user_id=?", (user["id"],)).fetchone()
    profile = dict(profile_row) if profile_row else None
    rows = conn.execute("""select t.*, r.id submission_id, r.note submission_note, r.evidence_url, r.status submission_status, r.created_at submitted_at
      from checkin_tasks t left join checkin_records r on r.task_id=t.id and r.user_id=?
      where t.status='active' and (t.assignee_user_id=? or t.assignee_user_id is null or trim(t.assignee_user_id)='')
      order by case when t.due_at is null then 1 else 0 end, t.due_at, t.created_at desc""", (user["id"], user["id"])).fetchall()
    task_rows = [dict(row) for row in rows]
    dates = [str(row["d"])[:10] for row in conn.execute("select distinct substr(created_at,1,10) d from checkin_records where user_id=? order by d desc", (user["id"],)).fetchall()]
    streak = 0
    cursor = datetime.now(timezone.utc).date()
    for value in dates:
        try: day = datetime.fromisoformat(value).date()
        except ValueError: continue
        if day == cursor or (streak == 0 and day == cursor - timedelta(days=1)):
            streak += 1; cursor = day - timedelta(days=1)
        elif day < cursor: break
    return {"profile": profile, "roles": [{"value": key, "label": label} for key, label in LEARNING_ROLES.items()], "tasks": task_rows, "streak": streak, "user": public_user(user)}


def save_learning_profile(conn: sqlite3.Connection, user: dict, data: dict) -> dict:
    role = str(data.get("role") or "").strip().lower()
    if role not in LEARNING_ROLES:
        raise ValueError("请选择老师、家长或同学身份")
    display_name = compact_text(data.get("display_name") or user.get("display_name") or user.get("email") or LEARNING_ROLES[role], 80)
    conn.execute("""insert into learning_profiles(user_id,role,display_name,class_id,updated_at) values(?,?,?,?,?)
      on conflict(user_id) do update set role=excluded.role,display_name=excluded.display_name,class_id=excluded.class_id,updated_at=excluded.updated_at""", (user["id"], role, display_name, compact_text(data.get("class_id"), 80), now_iso()))
    return checkin_dashboard(conn, user)


def create_checkin_task(conn: sqlite3.Connection, user: dict, data: dict) -> dict:
    profile = conn.execute("select * from learning_profiles where user_id=?", (user["id"],)).fetchone()
    if not profile or profile["role"] not in {"teacher", "parent"}:
        raise PermissionError("只有老师或家长可以发布指定任务")
    title = compact_text(data.get("title"), 100)
    if not title:
        raise ValueError("请填写任务名称")
    task_id = str(uuid.uuid4())
    conn.execute("""insert into checkin_tasks(id,creator_user_id,assignee_user_id,class_id,title,description,task_type,due_at,status,created_at)
      values(?,?,?,?,?,?,?,?,?,?)""", (task_id, user["id"], compact_text(data.get("assignee_user_id"), 100) or None, compact_text(data.get("class_id") or profile["class_id"], 100) or None, title, compact_text(data.get("description"), 500), compact_text(data.get("task_type") or "study", 40), compact_text(data.get("due_at"), 40) or None, "active", now_iso()))
    return dict(conn.execute("select * from checkin_tasks where id=?", (task_id,)).fetchone())


def submit_checkin_task(conn: sqlite3.Connection, user: dict, data: dict) -> dict:
    task_id = str(data.get("task_id") or "").strip()
    task = conn.execute("select * from checkin_tasks where id=? and status='active'", (task_id,)).fetchone()
    if not task:
        raise ValueError("任务不存在或已截止")
    if task["assignee_user_id"] and task["assignee_user_id"] != user["id"]:
        raise PermissionError("该任务没有指派给当前账号")
    record_id = str(uuid.uuid4())
    conn.execute("""insert into checkin_records(id,task_id,user_id,note,evidence_url,status,created_at) values(?,?,?,?,?,?,?)
      on conflict(task_id,user_id) do update set note=excluded.note,evidence_url=excluded.evidence_url,status=excluded.status,created_at=excluded.created_at""", (record_id, task_id, user["id"], compact_text(data.get("note"), 1000), compact_text(data.get("evidence_url"), 500), "submitted", now_iso()))
    return checkin_dashboard(conn, user)


def upsert_chat_model(
    conn: sqlite3.Connection,
    model_id: str,
    name: str,
    provider: str,
    endpoint: str,
    model: str,
    api_key: str,
    *,
    supports_vision: int = 0,
    temperature: float = 0.45,
    max_tokens: int = 7000,
    is_default: int = 0,
) -> None:
    existing = conn.execute("select * from model_configs where id = ?", (model_id,)).fetchone()
    ts = now_iso()
    if existing:
        old = dict(existing)
        key = api_key or old.get("api_key") or ""
        conn.execute(
            """
            update model_configs
            set name=?, provider=?, endpoint=?, model=?, api_key=?,
                supports_vision=?, temperature=?, max_tokens=?, updated_at=?
            where id=?
            """,
            (name, provider, endpoint, model, key, supports_vision, temperature, max_tokens, ts, model_id),
        )
        if is_default:
            conn.execute("update model_configs set is_default = 0")
            conn.execute("update model_configs set is_default = 1, updated_at = ? where id = ?", (ts, model_id))
        return
    conn.execute(
        """
        insert into model_configs
        (id, name, provider, endpoint, model, api_key, supports_vision, temperature,
         max_tokens, is_default, created_at, updated_at)
        values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            model_id,
            name,
            provider,
            endpoint,
            model,
            api_key,
            supports_vision,
            temperature,
            max_tokens,
            is_default,
            ts,
            ts,
        ),
    )
    if is_default:
        conn.execute("update model_configs set is_default = 0 where id != ?", (model_id,))
        conn.execute("update model_configs set is_default = 1, updated_at = ? where id = ?", (ts, model_id))


def upsert_image_model(
    conn: sqlite3.Connection,
    model_id: str,
    name: str,
    provider: str,
    endpoint: str,
    model: str,
    api_key: str,
    *,
    size: str = DEFAULT_OPENAI_IMAGE_SIZE,
    quality: str = DEFAULT_OPENAI_IMAGE_QUALITY,
    is_default: int = 0,
) -> None:
    existing = conn.execute("select * from image_model_configs where id = ?", (model_id,)).fetchone()
    ts = now_iso()
    if existing:
        old = dict(existing)
        key = api_key or old.get("api_key") or ""
        conn.execute(
            """
            update image_model_configs
            set name=?, provider=?, endpoint=?, model=?, api_key=?,
                size=?, quality=?, updated_at=?
            where id=?
            """,
            (name, provider, endpoint, model, key, size, quality, ts, model_id),
        )
        if is_default:
            conn.execute("update image_model_configs set is_default = 0")
            conn.execute("update image_model_configs set is_default = 1, updated_at = ? where id = ?", (ts, model_id))
        return
    conn.execute(
        """
        insert into image_model_configs
        (id, name, provider, endpoint, model, api_key, size, quality,
         is_default, created_at, updated_at)
        values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (model_id, name, provider, endpoint, model, api_key, size, quality, is_default, ts, ts),
    )
    if is_default:
        conn.execute("update image_model_configs set is_default = 0 where id != ?", (model_id,))
        conn.execute("update image_model_configs set is_default = 1, updated_at = ? where id = ?", (ts, model_id))


def seed_defaults(conn: sqlite3.Connection) -> None:
    minimax_key = os.environ.get("MINIMAX_API_KEY") or os.environ.get("AI_API_KEY") or ""
    deepseek_key = os.environ.get("DEEPSEEK_API_KEY") or ""
    fenno_key = os.environ.get("FENNO_API_KEY") or ""
    metaso_key = os.environ.get("METASO_API_KEY") or ""
    openai_image_key = os.environ.get("OPENAI_IMAGE_API_KEY") or os.environ.get("OPENAI_API_KEY") or ""

    upsert_chat_model(
        conn,
        "minimax-m3-default",
        "MiniMax M3 视觉 OCR",
        "minimax",
        MINIMAX_ENDPOINT,
        "MiniMax-M3",
        minimax_key,
        supports_vision=1,
        temperature=0.45,
        max_tokens=7000,
        is_default=1,
    )
    conn.execute("delete from model_configs where provider = 'deepseek' or id = 'deepseek-reasoner-default'")
    upsert_chat_model(
        conn,
        "fenno-gpt-default",
        "Fenno GPT 深度拆题",
        "fenno",
        FENNO_BASE_URL,
        "gpt-5.4",
        fenno_key,
        supports_vision=0,
        temperature=0.35,
        max_tokens=8000,
        is_default=0,
    )
    upsert_chat_model(
        conn,
        "metaso-qa-default",
        "秘塔联网拆题",
        "metaso",
        "https://metaso.cn/api/v1/chat/completions",
        os.environ.get("METASO_CHAT_MODEL") or "fast",
        metaso_key,
        supports_vision=0,
        temperature=0.2,
        max_tokens=7000,
        is_default=0,
    )

    upsert_image_model(
        conn,
        "openai-gpt-image-default",
        "OpenAI GPT Image 卡片生成",
        "openai",
        OPENAI_IMAGE_ENDPOINT,
        DEFAULT_OPENAI_IMAGE_MODEL,
        openai_image_key,
        is_default=1 if openai_image_key else 0,
    )
    upsert_image_model(
        conn,
        "fenno-gpt-image-default",
        "Fenno GPT-Image2 卡片生成",
        "fenno",
        FENNO_BASE_URL,
        "gpt-image-2",
        fenno_key,
        is_default=0 if openai_image_key else 1,
    )
    prompts = {
        "ocr": ("OCR 提示词", OCR_PROMPT),
        "diagnosis": ("拆题诊断提示词", DIAGNOSIS_PROMPT),
        "grading": ("批改提示词", GRADING_PROMPT),
    }
    for key, (name, content) in prompts.items():
        conn.execute(
            """
            insert or ignore into prompt_templates (key, name, content, updated_at)
            values (?, ?, ?, ?)
            """,
            (key, name, content, now_iso()),
        )
    diagnosis_prompt = conn.execute(
        "select content from prompt_templates where key = ?",
        ("diagnosis",),
    ).fetchone()
    if diagnosis_prompt and (
        "student_answer_analysis" not in diagnosis_prompt["content"]
        or "learning_strategy" not in diagnosis_prompt["content"]
    ):
        conn.execute(
            "update prompt_templates set content = ?, updated_at = ? where key = ?",
            (DIAGNOSIS_PROMPT, now_iso(), "diagnosis"),
        )
    conn.execute(
        """
        insert or ignore into redeem_codes (code, credits, is_used, created_at)
        values (?, ?, 0, ?)
        """,
        ("DEMO2026", 30, now_iso()),
    )


def ensure_admin_user(conn: sqlite3.Connection) -> str:
    row = conn.execute(
        "select id from app_users where is_admin = 1 order by created_at asc limit 1",
    ).fetchone()
    if row:
        admin_id = row["id"]
    else:
        existing = conn.execute("select id from app_users where email = ?", (ADMIN_EMAIL,)).fetchone()
        if existing:
            admin_id = existing["id"]
            conn.execute("update app_users set is_admin = 1 where id = ?", (admin_id,))
        else:
            admin_id = str(uuid.uuid4())
            conn.execute(
                """
                insert into app_users (id, email, password_hash, credits, is_admin, created_at)
                values (?, ?, ?, ?, 1, ?)
                """,
                (admin_id, ADMIN_EMAIL, hash_password(ADMIN_PASSWORD), 9999, now_iso()),
            )
    conn.execute(
        "update wrong_questions set user_id = ? where user_id is null or trim(user_id) = ''",
        (admin_id,),
    )
    conn.execute(
        "update tool_runs set user_id = ? where user_id is null or trim(user_id) = ''",
        (admin_id,),
    )
    conn.execute(
        "update profile_exports set user_id = ? where user_id is null or trim(user_id) = ''",
        (admin_id,),
    )
    return admin_id


def mask_key(value: str | None) -> str:
    if not value:
        return ""
    if len(value) <= 10:
        return "*" * len(value)
    return f"{value[:5]}...{value[-4:]}"


def public_model(row: sqlite3.Row | dict) -> dict:
    item = dict(row)
    item["api_key_masked"] = mask_key(item.get("api_key"))
    item.pop("api_key", None)
    return item


def public_image_model(row: sqlite3.Row | dict) -> dict:
    item = dict(row)
    item["api_key_masked"] = mask_key(item.get("api_key"))
    item.pop("api_key", None)
    return item


def get_model(conn: sqlite3.Connection, model_id: str | None = None) -> dict:
    row = None
    if model_id:
        row = conn.execute("select * from model_configs where id = ?", (model_id,)).fetchone()
    if not row:
        row = conn.execute("select * from model_configs where is_default = 1 limit 1").fetchone()
    if not row:
        row = conn.execute("select * from model_configs limit 1").fetchone()
    if not row:
        raise ValueError("请先在后台配置模型")
    return dict(row)


def get_image_model(conn: sqlite3.Connection, image_model_id: str | None = None) -> dict:
    row = None
    if image_model_id:
        row = conn.execute("select * from image_model_configs where id = ?", (image_model_id,)).fetchone()
    if not row:
        row = conn.execute("select * from image_model_configs where is_default = 1 limit 1").fetchone()
    if not row:
        row = conn.execute("select * from image_model_configs limit 1").fetchone()
    if not row:
        raise ValueError("请先在后台配置图片生成模型")
    return dict(row)


def get_vision_model(conn: sqlite3.Connection, model_id: str | None = None) -> dict:
    model = get_model(conn, model_id)
    if model.get("supports_vision"):
        return model
    row = conn.execute(
        "select * from model_configs where supports_vision = 1 order by is_default desc, updated_at desc limit 1"
    ).fetchone()
    if row:
        return dict(row)
    raise ValueError("当前没有支持图片/OCR的模型，请到后台配置一个视觉模型")


def get_prompt(conn: sqlite3.Connection, key: str) -> str:
    row = conn.execute("select content from prompt_templates where key = ?", (key,)).fetchone()
    if not row:
        raise ValueError(f"缺少提示词模板：{key}")
    return row["content"]


def save_data_url(data_url: str | None) -> str:
    if not data_url or "," not in data_url:
        return ""
    header, payload = data_url.split(",", 1)
    ext = "png"
    if "jpeg" in header or "jpg" in header:
        ext = "jpg"
    elif "webp" in header:
        ext = "webp"
    filename = f"{uuid.uuid4()}.{ext}"
    path = UPLOAD_DIR / filename
    path.write_bytes(base64.b64decode(payload))
    return f"/uploads/{filename}"


def decode_data_url(data_url: str) -> tuple[str, bytes]:
    if not data_url or "," not in data_url:
        raise ValueError("无效的文件数据")
    header, payload = data_url.split(",", 1)
    return header, base64.b64decode(payload)


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = Path(filename or "").suffix.lower()
    if suffix in {".txt", ".md", ".csv", ".json", ".log"}:
        for encoding in ("utf-8", "utf-8-sig", "gb18030", "gbk"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("PDF 解析依赖未安装，请执行 pip install pypdf") from exc
        reader = PdfReader(BytesIO(content))
        pages = []
        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append(f"【第 {idx} 页】\n{text}")
        if not pages:
            raise ValueError("PDF 未提取到可读文本，可能是扫描版，请改用图片 OCR 上传")
        return "\n\n".join(pages)
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("Word 解析依赖未安装，请执行 pip install python-docx") from exc
        document = Document(BytesIO(content))
        lines = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append(" | ".join(cells))
        if not lines:
            raise ValueError("Word 文档未提取到可读文本")
        return "\n".join(lines)
    if suffix == ".doc":
        raise ValueError("暂不支持旧版 .doc，请另存为 .docx 或 PDF 后上传")
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}:
        raise ValueError("图片文件请使用图片 OCR 上传，或在错题拆解工作台选择图片输入")
    raise ValueError(f"暂不支持该文件类型：{suffix or '未知'}，可上传 PDF、Word(.docx)、TXT、Markdown")


def extract_document_from_data_url(filename: str, data_url: str) -> dict:
    _, content = decode_data_url(data_url)
    if not content:
        raise ValueError("文件内容为空")
    if len(content) > 25 * 1024 * 1024:
        raise ValueError("文件过大，请控制在 25MB 以内")
    text = extract_document_text(filename or "upload.txt", content).strip()
    if not text:
        raise ValueError("未能从文件中提取文本")
    return {
        "filename": filename or "upload.txt",
        "text": text,
        "chars": len(text),
    }


def validate_paper_source_text(value: object) -> str:
    """Reject archive/binary payloads before they can become OCR question text."""
    text = str(value or "").strip()
    if not text:
        raise ValueError("试卷文本为空")
    prefix = text[:32]
    if prefix.startswith("PK\x03\x04") or "\x00" in text or "word/embeddings/" in text or "word/document.xml" in text:
        raise ValueError("检测到二进制文档内容；Word/PDF必须通过文档上传方式解析，不能按纯文本读取")
    replacement_ratio = text.count("�") / max(1, len(text))
    control_count = sum(1 for char in text if ord(char) < 32 and char not in "\n\r\t")
    if replacement_ratio > 0.02 or control_count / max(1, len(text)) > 0.01:
        raise ValueError("试卷文本包含大量乱码或控制字符，请重新上传原始图片、DOCX、PDF或UTF-8文本")
    return text


def prepare_paper_page(page: dict) -> tuple[str | None, str]:
    if page.get("image_data_url"):
        return save_data_url(page.get("image_data_url")), ""
    if page.get("file_data_url"):
        filename = str(page.get("name") or "upload.txt")
        extracted = extract_document_from_data_url(filename, str(page.get("file_data_url") or ""))
        return None, validate_paper_source_text(extracted.get("text"))
    return None, validate_paper_source_text(page.get("text"))


def normalize_endpoint(value: str | None, kind: str) -> str:
    endpoint = (value or "").strip().rstrip("/")
    if not endpoint:
        endpoint = FENNO_BASE_URL if kind in {"fenno-chat", "image"} else ""
    if not endpoint:
        return ""
    lower = endpoint.lower()
    if lower.endswith(("/chat/completions", "/responses", "/images/generations")):
        return endpoint
    if lower.endswith("/v1"):
        base = endpoint
    else:
        base = f"{endpoint}/v1"
    if kind == "image":
        return f"{base}/images/generations"
    if kind == "fenno-chat":
        return f"{base}/chat/completions"
    return f"{base}/chat/completions"


def response_api_text(data: dict) -> str:
    if isinstance(data.get("output_text"), str):
        return data["output_text"]
    parts: list[str] = []
    for output in data.get("output") or []:
        for content in output.get("content") or []:
            if isinstance(content, dict) and isinstance(content.get("text"), str):
                parts.append(content["text"])
    return "\n".join(parts).strip()


def messages_to_response_input(messages: list[dict]) -> str:
    chunks: list[str] = []
    for message in messages:
        role = str(message.get("role") or "user").upper()
        content = message.get("content")
        if isinstance(content, str):
            text = content
        else:
            text = json.dumps(content, ensure_ascii=False)
        chunks.append(f"{role}:\n{text}")
    return "\n\n".join(chunks)


def chat_completion(model_config: dict, messages: list[dict], max_tokens: int | None = None, temperature: float | None = None, *, timeout: int = 120) -> str:
    provider = (model_config.get("provider") or "").lower()
    provider_label = provider or "model"
    if provider == "metaso":
        result = metaso_client.answer(
            messages_to_response_input(messages),
            model=model_config.get("model") or os.environ.get("METASO_CHAT_MODEL") or "fast",
            api_key=model_config.get("api_key") or os.environ.get("METASO_API_KEY") or "",
        )
        content = str(result.get("content") or "").strip()
        if not content:
            raise RuntimeError("秘塔问答没有返回可用的拆题内容")
        return content
    endpoint = model_config.get("endpoint") or MINIMAX_ENDPOINT
    if provider in {"fenno", "openai-compatible", "oneapi", "newapi"}:
        kind = "fenno-chat" if provider == "fenno" or "api.fenno.ai" in endpoint else "chat"
        endpoint = normalize_endpoint(endpoint, kind)
    api_key = (
        model_config.get("api_key")
        or (os.environ.get("DEEPSEEK_API_KEY") if provider == "deepseek" else "")
        or (os.environ.get("MINIMAX_API_KEY") if provider == "minimax" else "")
        or (os.environ.get("FENNO_API_KEY") if provider in {"fenno", "openai-compatible"} else "")
        or os.environ.get("AI_API_KEY")
    )
    if not api_key:
        raise RuntimeError(f"{provider_label} API Key 未配置。请到后台模型设置中保存 API Key。")
    use_responses_api = endpoint.lower().endswith("/responses")
    if use_responses_api:
        payload = {
            "model": model_config["model"],
            "input": messages_to_response_input(messages),
            "temperature": float(model_config.get("temperature", 0.45) if temperature is None else temperature),
            "max_output_tokens": int(max_tokens or model_config.get("max_tokens") or 6000),
        }
    else:
        payload = {
            "model": model_config["model"],
            "messages": messages,
            "temperature": float(model_config.get("temperature", 0.45) if temperature is None else temperature),
        }
        if provider == "minimax":
            payload["max_completion_tokens"] = int(max_tokens or model_config.get("max_tokens") or 6000)
            payload["thinking"] = {"type": "disabled"}
        else:
            payload["max_tokens"] = int(max_tokens or model_config.get("max_tokens") or 6000)
        if provider == "deepseek":
            payload["response_format"] = {"type": "json_object"}
    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    last_error: Exception | None = None
    for attempt in range(3):
        req = urllib.request.Request(
            endpoint,
            data=body,
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
                "Connection": "close",
            },
            method="POST",
        )
        try:
            with urllib.request.urlopen(req, timeout=timeout) as resp:
                data = json.loads(resp.read().decode("utf-8"))
            break
        except urllib.error.HTTPError as exc:
            detail = exc.read().decode("utf-8", errors="replace")
            raise RuntimeError(f"{provider_label} 调用失败：HTTP {exc.code} {detail[:600]}") from exc
        except (urllib.error.URLError, OSError) as exc:
            last_error = exc
            if attempt < 2:
                time.sleep(1.2 * (attempt + 1))
                continue
            raise RuntimeError(f"{provider_label} 网络错误，已自动重连3次仍失败：{last_error}") from exc
    if use_responses_api:
        text = response_api_text(data)
        if not text:
            raise RuntimeError(f"{provider_label} 返回格式异常：{json.dumps(data, ensure_ascii=False)[:600]}")
        return text
    if "choices" not in data or not data["choices"]:
        raise RuntimeError(f"{provider_label} 返回格式异常：{json.dumps(data, ensure_ascii=False)[:600]}")
    return data["choices"][0]["message"].get("content", "")


def minimax_chat(model_config: dict, messages: list[dict], max_tokens: int | None = None, temperature: float | None = None, *, timeout: int = 120) -> str:
    return chat_completion(model_config, messages, max_tokens=max_tokens, temperature=temperature, timeout=timeout)


def strip_thinking(text: str) -> str:
    return re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL).strip()


def repair_common_json(text: str) -> str:
    """Repair deterministic formatting slips commonly produced by vision models."""
    repaired = text.lstrip("\ufeff").strip()
    repaired = re.sub(r",\s*([}\]])", r"\1", repaired)
    value = r'(?:"(?:\\.|[^"\\])*"|true|false|null|-?\d+(?:\.\d+)?(?:[eE][+-]?\d+)?|[}\]])'
    # A completed value followed by the next object key without a comma.
    repaired = re.sub(rf'({value})(\s*)(?="(?:\\.|[^"\\])+"\s*:)', r"\1,\2", repaired)
    # Adjacent objects/arrays inside a list without a comma.
    repaired = re.sub(r'([}\]])(\s*)(?=[{\[])', r"\1,\2", repaired)
    return repaired


def extract_json(text: str) -> dict:
    clean = strip_thinking(text)
    clean = re.sub(r"^```(?:json)?", "", clean.strip(), flags=re.IGNORECASE).strip()
    clean = re.sub(r"```$", "", clean.strip()).strip()
    candidates = [clean]
    start = clean.find("{")
    end = clean.rfind("}")
    if start >= 0 and end > start:
        candidates.append(clean[start : end + 1])
    for candidate in list(candidates):
        repaired = repair_common_json(candidate)
        if repaired != candidate:
            candidates.append(repaired)
    last_error: json.JSONDecodeError | None = None
    for candidate in candidates:
        try:
            value = json.loads(candidate, strict=False)
            if isinstance(value, dict):
                return value
        except json.JSONDecodeError as exc:
            last_error = exc
    if last_error:
        raise last_error
    raise json.JSONDecodeError("model response is not a JSON object", clean, 0)


def recover_single_ocr_payload(raw: str) -> dict:
    """Keep usable OCR text even when the model's surrounding JSON is malformed."""
    clean = strip_thinking(raw)

    def field(name: str) -> str:
        match = re.search(rf'"{re.escape(name)}"\s*:\s*"((?:\\.|[^"\\])*)"', clean, re.DOTALL)
        if not match:
            return ""
        try:
            return json.loads(f'"{match.group(1)}"', strict=False).strip()
        except json.JSONDecodeError:
            return match.group(1).replace(r"\n", "\n").replace(r'\"', '"').strip()

    result = {
        "ocr_text": field("ocr_text"),
        "printed_question": field("printed_question"),
        "student_work": field("student_work"),
        "teacher_marks": field("teacher_marks"),
        "uncertain_parts": [],
        "parse_mode": "recovered_fields",
        "needs_review": True,
    }
    if not result["ocr_text"]:
        result["ocr_text"] = result["printed_question"]
    if not result["printed_question"]:
        result["printed_question"] = result["ocr_text"]
    if not result["ocr_text"]:
        plain = re.sub(r"[`{}\[\]]", " ", clean)
        plain = re.sub(r'"(?:ocr_text|printed_question|student_work|teacher_marks|uncertain_parts)"\s*:', " ", plain)
        result["ocr_text"] = re.sub(r"\s+", " ", plain).strip(' ,"')[:6000]
        result["printed_question"] = result["ocr_text"]
    return result


def _clamp_confidence(value: object, default: float = 0.0) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        number = default
    return round(max(0.0, min(1.0, number)), 4)


def _normalize_bbox(value: object) -> list[float] | None:
    if not isinstance(value, list) or len(value) != 4:
        return None
    normalized = []
    for item in value:
        try:
            normalized.append(max(0.0, min(1.0, float(item))))
        except (TypeError, ValueError):
            return None
    return normalized


def normalize_paper_ocr_payload(payload: dict, *, parse_mode: str = "json", repaired: bool = False) -> dict:
    page_text = str(payload.get("page_text") or payload.get("ocr_text") or "").strip()
    raw_questions = payload.get("questions") if isinstance(payload.get("questions"), list) else []
    questions = []
    for index, raw in enumerate(raw_questions, start=1):
        if not isinstance(raw, dict):
            continue
        printed_text = str(raw.get("printed_text") or raw.get("question_text") or "").strip()
        student_work = str(raw.get("student_work") or "").strip()
        teacher_marks = str(raw.get("teacher_marks") or "").strip()
        if not printed_text and not student_work and not teacher_marks:
            continue
        score, max_score = raw.get("score"), raw.get("max_score")
        questions.append({
            **raw,
            "question_no": str(raw.get("question_no") or index),
            "printed_text": printed_text,
            "student_work": student_work,
            "teacher_marks": teacher_marks,
            "answer_state": normalize_answer_state(raw.get("answer_state"), score, max_score),
            "confidence": _clamp_confidence(raw.get("confidence"), 0.45),
            "bbox": _normalize_bbox(raw.get("bbox")),
            "continuation": bool(raw.get("continuation", False)),
        })
    if not page_text and questions:
        page_text = "\n\n".join(q["printed_text"] for q in questions if q["printed_text"])
    return {
        **payload,
        "page_text": page_text,
        "page_confidence": _clamp_confidence(payload.get("page_confidence"), 0.0),
        "questions": questions,
        "parse_mode": parse_mode,
        "response_repaired": bool(repaired),
    }


def _extract_truncated_page_text(raw: str) -> str:
    match = re.search(r'"page_text"\s*:\s*"((?:\\.|[^"\\])*)"', raw or "", flags=re.DOTALL)
    if not match:
        return ""
    try:
        return json.loads(f'"{match.group(1)}"', strict=False)
    except json.JSONDecodeError:
        return match.group(1).replace("\\n", "\n").replace('\\"', '"')


def parse_paper_ocr_response(raw: str) -> dict:
    """Decode vision output without discarding readable OCR when JSON is malformed."""
    try:
        payload = extract_json(raw)
        direct = strip_thinking(raw).strip().lstrip("`").removeprefix("json").strip()
        repaired = re.sub(r",\s*([}\]])", r"\1", direct) != direct
        return normalize_paper_ocr_payload(payload, repaired=repaired)
    except json.JSONDecodeError:
        page_text = _extract_truncated_page_text(raw)
        parse_mode = "truncated_json_fallback" if page_text else "text_fallback"
        page_text = page_text or strip_thinking(raw).strip()
        questions = split_numbered_questions(page_text)
        for question in questions:
            question.setdefault("answer_state", "review_required")
            question["confidence"] = min(float(question.get("confidence") or 0.45), 0.55)
        return normalize_paper_ocr_payload(
            {"page_text": page_text, "page_confidence": 0.35 if page_text else 0.0, "questions": questions},
            parse_mode=parse_mode,
            repaired=True,
        )


PAPER_OCR_PROMPT = """你是全卷分析 OCR 与阅卷助手。只根据试卷页面识别，不编造。
识别印刷题干、题号、分值、选项、学生手写笔迹、演算过程、红叉/勾/圈画、教师批语和得分。
把页面切分为独立题目；跨页题要标记 continuation=true。无法确认的字符使用 [?]。
判断每题状态：correct、wrong、partial、blank、review_required。低置信度必须使用 review_required。
严格输出 JSON：
{"page_text":"", "page_confidence":0.0, "questions":[{
 "question_no":"1", "printed_text":"", "student_work":"", "teacher_marks":"",
 "answer_state":"review_required", "score":null, "max_score":null,
 "confidence":0.0, "bbox":[0,0,1,1], "continuation":false
}]}
"""


def _vision_ocr_once(
    model: dict,
    prompt: str,
    image_data_url: str,
    *,
    max_long_side: int = 2200,
    max_tokens: int = 2800,
) -> dict:
    content = [
        {"type": "text", "text": prompt},
        {
            "type": "image_url",
            "image_url": {
                "url": image_data_url,
                "detail": "high",
                "max_long_side_pixel": max_long_side,
            },
        },
    ]
    raw = minimax_chat(model, [{"role": "user", "content": content}], max_tokens=max_tokens, temperature=0.08)
    try:
        result = extract_json(raw)
        result.setdefault("parse_mode", "json")
    except json.JSONDecodeError:
        result = recover_single_ocr_payload(raw)
    result.setdefault("ocr_text", "")
    result.setdefault("printed_question", result.get("ocr_text") or "")
    result.setdefault("student_work", "")
    result.setdefault("teacher_marks", "")
    result.setdefault("uncertain_parts", [])
    result["confidence"] = normalize_ocr_confidence(result)
    return result


def run_ocr(image_data_url: str, model_id: str | None = None, *, handwriting: bool = False) -> dict:
    if not image_data_url:
        raise ValueError("请先上传题目图片")
    with db() as conn:
        model = get_vision_model(conn, model_id)
        prompt = merge_handwriting_prompt(get_prompt(conn, "ocr"), handwriting)
    base_image, base_meta = normalize_ocr_image_data_url(image_data_url, max_long_side=2400 if handwriting else 2000)
    if not handwriting:
        result = _vision_ocr_once(model, prompt, base_image, max_long_side=2000)
        result["handwriting_mode"] = False
        result["image_preprocessing"] = base_meta
        result["ocr_confidence"] = result.get("confidence")
        return result

    enhanced_image, enhanced_meta = preprocess_for_handwriting(base_image, max_long_side=2600)
    red_image, red_meta = preprocess_red_marks(base_image, max_long_side=2600)
    pass_specs = [
        ("原图", base_image, 2400, 3000),
        ("笔迹增强", enhanced_image, 2600, 3000),
        ("红笔增强", red_image, 2600, 2800),
    ]

    def recognize_pass(spec: tuple[str, str, int, int]) -> dict:
        label, image, long_side, tokens = spec
        try:
            return {"label": label, "result": _vision_ocr_once(model, prompt, image, max_long_side=long_side, max_tokens=tokens)}
        except Exception as exc:
            return {"label": label, "failure": f"{label}识别未完成：{str(exc)[:100]}"}

    # 三路识别互不依赖，并行后取最佳结果；耗时由三次相加降为最慢一路。
    outcomes = parallel_run(pass_specs, recognize_pass, max_workers=3)
    results = [item["result"] for item in outcomes if item.get("result")]
    failures = [item["failure"] for item in outcomes if item.get("failure")]
    if not results:
        raise RuntimeError("图片识别服务暂时没有返回可用文字，图片已经保留，请稍后从错题本继续处理")
    chosen = merge_ocr_passes(*results)
    chosen["handwriting_mode"] = True
    chosen["recognition_passes"] = len(results)
    chosen["recognition_warnings"] = failures
    chosen["image_preprocessing"] = {"base": base_meta, "enhanced": enhanced_meta, "red_channel": red_meta}
    chosen["ocr_confidence"] = chosen.get("confidence")
    chosen["photo_tips"] = HANDWRITING_PHOTO_TIPS
    return chosen


def diagnose_with_llm(
    question_text: str,
    wrong_answer: str = "",
    image_text: str = "",
    model_id: str | None = None,
    subject: str = "自动识别",
    rag_context: str = "",
) -> dict:
    with db() as conn:
        model = get_model(conn, model_id)
        prompt = get_prompt(conn, "diagnosis")
    selected_provider = str(model.get("provider") or "").lower()
    metaso_context = ""
    try:
        metaso_context = metaso_client.build_answer_context(
            f"请分析并解答下面这道题，指出关键方法、易错点，并给出可核验依据：\n学科：{subject}\n题目：{question_text}\n学生作答：{wrong_answer or '未提供'}"
        )
    except Exception:
        metaso_context = ""
    combined_rag_context = "\n\n".join(part for part in (rag_context, metaso_context) if part)
    user_text = f"""用户选择的学科/场景：
{subject or "自动识别"}

题目 OCR 文本：
{image_text or "(无)"}

用户修正后的题干：
{question_text}

学生作答/错解/批注/卡点：
{wrong_answer or "(未提供)"}

高考母题库 RAG 检索资料（优先依据；若与题目无关须说明）：
{combined_rag_context or "(未检索到母题资料，请按通用数学推理作答)"}

请按提示词固定 JSON 结构输出；若使用了 RAG 资料，请在结果中体现依据，且不得捏造资料中不存在的内容。"""
    messages = [
        {"role": "system", "content": prompt},
        {"role": "user", "content": user_text},
    ]
    primary_model_succeeded = False
    primary_model_error = ""
    fallback_provider = ""
    try:
        raw = minimax_chat(
            model,
            messages,
            max_tokens=int(model.get("max_tokens") or 7000),
            temperature=float(model.get("temperature") or 0.45),
            timeout=90,
        )
        result = extract_json(raw)
        primary_model_succeeded = True
    except Exception as exc:
        primary_model_error = str(exc)[:160]
        result = None
        if selected_provider == "metaso":
            try:
                with db() as conn:
                    fallback_row = conn.execute(
                        "select * from model_configs where provider != 'metaso' order by is_default desc, updated_at desc limit 1"
                    ).fetchone()
                if fallback_row:
                    fallback_model = dict(fallback_row)
                    fallback_provider = str(fallback_model.get("provider") or "model")
                    fallback_raw = minimax_chat(
                        fallback_model,
                        messages,
                        max_tokens=int(fallback_model.get("max_tokens") or 7000),
                        temperature=float(fallback_model.get("temperature") or 0.45),
                        timeout=90,
                    )
                    result = extract_json(fallback_raw)
            except Exception as fallback_exc:
                primary_model_error = f"{primary_model_error}；备用模型也未完成：{str(fallback_exc)[:120]}"
        if result is None:
            result = build_skeleton_diagnosis(question_text, wrong_answer, subject)
            result["needs_review"] = True
            result.setdefault("student_answer_analysis", {})["evidence"] = [
                f"深拆模型暂不可用，已回退骨架拆题：{primary_model_error}"
            ]
            result["gaokao_rag"] = {"used": bool(combined_rag_context), "mode": "deep_fallback_skeleton", "metaso_used": bool(metaso_context)}
    result.setdefault("subject", subject or "自动识别")
    result.setdefault("confidence", 0.75)
    result.setdefault("external_qa", {
        "provider": "metaso",
        "attempted": selected_provider == "metaso" or bool(metaso_context),
        "used": bool(metaso_context) or (selected_provider == "metaso" and primary_model_succeeded),
        "primary_model": selected_provider == "metaso" and primary_model_succeeded,
        "fallback_provider": fallback_provider,
        "error": primary_model_error if selected_provider == "metaso" and not primary_model_succeeded else "",
        "kept_minimax": selected_provider != "metaso" or fallback_provider == "minimax",
    })
    result.setdefault("core_pattern", "待归纳题型")
    result.setdefault("practice_variants", [])
    result.setdefault("knowledge_points", [])
    result.setdefault(
        "student_answer_analysis",
        {
            "answer_presence": "未提供" if not wrong_answer else "已提供",
            "extracted_work": wrong_answer,
            "answer_status": "待模型进一步判断" if wrong_answer else "未提供作答",
            "likely_issue": "" if not wrong_answer else "已记录学生作答/卡点，需结合拆题结果判断",
            "evidence": [],
            "next_action": "先补充学生作答或错解过程" if not wrong_answer else "对照模型步骤定位断点",
        },
    )
    result.setdefault(
        "learning_strategy",
        {
            "decomposition_answer": "先识别题型和题目目标，再拆条件、选模型、执行计算或作答，最后用小题验证。",
            "make_it_easier": "先把题目降阶成一个更小的原型题，再从原型题迁移回原题。",
            "entry_point": result.get("core_pattern") or result.get("topic") or "题型入口",
            "cognitive_ladder": ["看懂题目目标", "找到可套用的模型", "完成一步一验算"],
            "micro_drills": ["用一句话说出题目类型", "写出第一步拆解公式或答题角度"],
            "teacher_hint": "这道题第一眼最像哪一类你已经做过的题？",
        },
    )
    if selected_provider == "metaso" and not primary_model_succeeded and fallback_provider:
        evidence = result.setdefault("student_answer_analysis", {}).setdefault("evidence", [])
        evidence.append(f"秘塔联网拆题未完成（{primary_model_error}），本次已自动由 {fallback_provider} 完成结构化拆题。")
    return normalize_diagnosis_payload(result, question_text, wrong_answer)


def grade_with_llm(variant: dict, answer_text: str, diagnosis: dict | None, model_id: str | None = None) -> dict:
    with db() as conn:
        model = get_model(conn, model_id)
        prompt = get_prompt(conn, "grading")
    user_text = f"""原题诊断摘要：
{json.dumps(diagnosis or {}, ensure_ascii=False)}

变式题：
{variant["stem"]}

参考答案：
{variant["answer"]}

参考解析：
{variant["analysis"]}

学生答案：
{answer_text}

请按 JSON 格式批改。"""
    raw = minimax_chat(
        model,
        [{"role": "system", "content": prompt}, {"role": "user", "content": user_text}],
        max_tokens=2200,
        temperature=0.25,
    )
    result = extract_json(raw)
    result["is_correct"] = bool(result.get("is_correct"))
    result["score"] = int(result.get("score", 0))
    return result


def save_variants(conn: sqlite3.Connection, wrong_id: str, diagnosis: dict) -> list[dict]:
    existing = conn.execute(
        "select * from exercise_variants where wrong_question_id = ? order by level",
        (wrong_id,),
    ).fetchall()
    if existing:
        return [row_to_dict(row) for row in existing]
    variants = diagnosis.get("practice_variants") or []
    saved: list[dict] = []
    for idx, item in enumerate(variants[:3], start=1):
        variant_id = str(uuid.uuid4())
        level = int(item.get("level") or idx)
        title = item.get("title") or ("同结构巩固" if level == 1 else "条件替换" if level == 2 else "迁移挑战")
        conn.execute(
            """
            insert into exercise_variants
            (id, wrong_question_id, level, title, stem, answer, analysis, created_at)
            values (?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                variant_id,
                wrong_id,
                level,
                title,
                item.get("stem", ""),
                item.get("answer", ""),
                item.get("analysis", ""),
                now_iso(),
            ),
        )
        saved.append(
            {
                "id": variant_id,
                "wrong_question_id": wrong_id,
                "level": level,
                "title": title,
                "stem": item.get("stem", ""),
                "answer": item.get("answer", ""),
                "analysis": item.get("analysis", ""),
                "created_at": now_iso(),
            }
        )
    return saved


def compact_text(value: object, limit: int = 720) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        text = json.dumps(value, ensure_ascii=False)
    else:
        text = str(value)
    text = re.sub(r"\s+", " ", text).strip()
    return text[:limit] + ("..." if len(text) > limit else "")


def build_study_card_prompt(item: dict, style: str = "") -> str:
    diagnosis = item.get("diagnosis") or {}
    student = diagnosis.get("student_answer_analysis") or {}
    strategy = diagnosis.get("learning_strategy") or {}
    answer = diagnosis.get("standard_answer") or {}
    decomposition = diagnosis.get("decomposition") or {}
    poem = diagnosis.get("poem") or {}
    variants = item.get("variants") or []
    first_variant = variants[0] if variants else {}
    solution_models = diagnosis.get("solution_models") or []
    first_model = solution_models[0] if solution_models else {}

    traps: list[str] = []
    for model in solution_models:
        traps.extend(model.get("common_mistakes") or [])
    if student.get("likely_issue"):
        traps.insert(0, student.get("likely_issue"))

    step_lines = []
    for idx, step in enumerate((decomposition.get("step_formulas") or [])[:5], start=1):
        step_lines.append(
            f"{idx}. {compact_text(step.get('name'), 30)}：{compact_text(step.get('formula') or step.get('operation'), 120)}"
        )

    self_checks = [
        "我是否先识别题型入口？",
        "我是否写清关键公式/采分点？",
        "我是否区分了易混概念和最终目标？",
    ]

    payload = {
        "标题": f"{diagnosis.get('subject') or '全科题目'}｜错题指导学习训练卡片",
        "主题": diagnosis.get("topic") or diagnosis.get("core_pattern") or "题目拆解训练",
        "错因标签": (diagnosis.get("knowledge_points") or [])[:3] + traps[:3],
        "题目": compact_text(diagnosis.get("cleaned_question") or item.get("corrected_text"), 520),
        "为什么容易错": compact_text(student.get("likely_issue") or "入口、公式、条件或表达容易混在一起，需要先拆题再作答。", 260),
        "纠错思路": step_lines or [compact_text(strategy.get("decomposition_answer"), 260)],
        "规范解答": compact_text(answer.get("concise_solution") or answer.get("final_answer"), 620),
        "关键提醒": compact_text(strategy.get("make_it_easier") or first_model.get("applies_when"), 220),
        "变式训练": compact_text(first_variant.get("stem") or "把条件替换成同结构新题，先说入口，再写关键步骤。", 280),
        "自我复盘": self_checks,
        "总结小诗": compact_text("；".join(poem.get("lines") or []), 180),
    }

    return f"""请生成一张竖版中文教辅学习卡片，比例 2:3，适合手机保存和打印。

视觉风格：
- 参考中国教辅产品的清爽蓝白底色，辅以少量橙金提醒色；线条干净，层级清楚。
- 不要做营销海报，不要真实机构 logo，不要二维码，不要水印。
- 版式必须像一张完整的“错题指导学习训练卡片”，包含编号 01 到 07 的七块内容。
- 尽量保持中文文字清晰，公式用可读的数学排版；内容过长时可以压缩摘要，但不要改题意。
- 每个区块用圆角 8px 以内的清晰边框或分隔线，避免文字重叠。

七块固定结构：
01 题目
02 这道题为什么容易错？
03 纠错思路
04 规范解答
05 关键提醒
06 变式训练
07 自我复盘

卡片数据：
{json.dumps(payload, ensure_ascii=False, indent=2)}

额外风格要求：
{style or "做成专业、漂亮、可信的高考/中考错题训练卡。"}"""


def call_image_generation(model_config: dict, prompt: str) -> bytes:
    provider = (model_config.get("provider") or "openai").lower()
    compatible_providers = {"openai", "fenno", "openai-compatible", "oneapi", "newapi"}
    if provider not in compatible_providers:
        raise RuntimeError("当前图片生成支持 OpenAI 兼容 Images API。请把中转站 provider 填为 openai、fenno 或 openai-compatible。")
    provider_label = model_config.get("provider") or "openai"
    api_key = (
        model_config.get("api_key")
        or (os.environ.get("FENNO_API_KEY") if provider == "fenno" else "")
        or os.environ.get("OPENAI_IMAGE_API_KEY")
        or os.environ.get("OPENAI_API_KEY")
    )
    if not api_key:
        raise RuntimeError(f"{provider_label} 图片 API Key 未配置，请在后台图片模型配置中保存 Key，或在服务器环境变量中设置 FENNO_API_KEY / OPENAI_API_KEY。")

    payload = {
        "model": model_config.get("model") or DEFAULT_OPENAI_IMAGE_MODEL,
        "prompt": prompt,
        "n": 1,
        "size": model_config.get("size") or DEFAULT_OPENAI_IMAGE_SIZE,
    }
    quality = model_config.get("quality") or DEFAULT_OPENAI_IMAGE_QUALITY
    if quality:
        payload["quality"] = quality

    body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    endpoint = normalize_endpoint(model_config.get("endpoint") or OPENAI_IMAGE_ENDPOINT, "image")
    req = urllib.request.Request(
        endpoint,
        data=body,
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "Connection": "close",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=240) as resp:
            data = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise RuntimeError(f"{provider_label} 图片生成失败：HTTP {exc.code} {detail[:800]}") from exc
    except (urllib.error.URLError, OSError) as exc:
        raise RuntimeError(f"{provider_label} 图片生成网络错误：{exc}") from exc

    items = data.get("data") or []
    if not items:
        raise RuntimeError(f"{provider_label} 图片返回为空：{json.dumps(data, ensure_ascii=False)[:600]}")
    first = items[0]
    if first.get("b64_json"):
        return base64.b64decode(first["b64_json"])
    if first.get("url"):
        with urllib.request.urlopen(first["url"], timeout=180) as resp:
            return resp.read()
    raise RuntimeError(f"{provider_label} 图片返回格式异常：{json.dumps(data, ensure_ascii=False)[:600]}")


def generate_study_card(wrong_id: str, image_model_id: str | None = None, style: str = "") -> dict:
    with db() as conn:
        item = get_wrong_question(conn, wrong_id)
        if not item:
            raise ValueError("wrong question not found")
        image_model = get_image_model(conn, image_model_id)

    prompt = build_study_card_prompt(item, style)
    image_bytes = call_image_generation(image_model, prompt)
    card_id = str(uuid.uuid4())
    filename = f"{card_id}.png"
    (CARD_DIR / filename).write_bytes(image_bytes)
    image_url = f"/cards/{filename}"

    with db() as conn:
        conn.execute(
            """
            insert into study_cards
            (id, wrong_question_id, image_model_id, image_url, prompt, model,
             size, quality, status, created_at)
            values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                card_id,
                wrong_id,
                image_model.get("id"),
                image_url,
                prompt,
                image_model.get("model") or DEFAULT_OPENAI_IMAGE_MODEL,
                image_model.get("size") or DEFAULT_OPENAI_IMAGE_SIZE,
                image_model.get("quality") or DEFAULT_OPENAI_IMAGE_QUALITY,
                "ready",
                now_iso(),
            ),
        )
        row = conn.execute("select * from study_cards where id = ?", (card_id,)).fetchone()
        return row_to_dict(row)


def get_wrong_question(conn: sqlite3.Connection, wrong_id: str) -> dict | None:
    row = conn.execute("select * from wrong_questions where id = ?", (wrong_id,)).fetchone()
    item = row_to_dict(row)
    if not item:
        return None
    variants = conn.execute(
        "select * from exercise_variants where wrong_question_id = ? order by level",
        (wrong_id,),
    ).fetchall()
    item["variants"] = []
    for row in variants:
        variant = row_to_dict(row)
        answer_rows = conn.execute(
            """
            select * from student_answers
            where exercise_variant_id = ?
            order by submitted_at desc
            """,
            (variant["id"],),
        ).fetchall()
        variant["answers"] = [row_to_dict(answer_row) for answer_row in answer_rows]
        item["variants"].append(variant)
    card_rows = conn.execute(
        """
        select * from study_cards
        where wrong_question_id = ?
        order by created_at desc
        """,
        (wrong_id,),
    ).fetchall()
    item["study_cards"] = [row_to_dict(card_row) for card_row in card_rows]
    return item


def clean_subject_name(value: object) -> str:
    subject = str(value or "").strip()
    if not subject or re.fullmatch(r"\?+", subject):
        return "未识别学科"
    return subject


def profile_items(conn: sqlite3.Connection, user: dict | None) -> dict:
    if not user:
        return {"subjects": [], "items": []}
    wq_where, wq_params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])
    rows = conn.execute(
        f"select id from wrong_questions{wq_where} order by created_at desc",
        wq_params,
    ).fetchall()
    items = []
    subjects: dict[str, int] = {}
    for row in rows:
        item = get_wrong_question(conn, row["id"])
        if not item:
            continue
        diagnosis = item.get("diagnosis") or {}
        subject = clean_subject_name(diagnosis.get("subject"))
        item["subject"] = subject
        subjects[subject] = subjects.get(subject, 0) + 1
        items.append(item)
    return {
        "subjects": [{"name": name, "count": count} for name, count in sorted(subjects.items())],
        "items": items,
    }


def md_escape(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value).replace("\r\n", "\n").strip()


def append_list(lines: list[str], title: str, values: list) -> None:
    if not values:
        return
    lines.append(f"### {title}")
    for value in values:
        lines.append(f"- {md_escape(value)}")
    lines.append("")


def build_profile_markdown(conn: sqlite3.Connection, data: dict, user: dict | None) -> dict:
    selections: dict = data.get("selections") or {}
    subject_filter = data.get("subject") or "全部学科"
    if not selections:
        raise ValueError("请至少选择一道题目")
    if not user:
        raise ValueError("请先登录后再导出档案")

    lines = [
        f"# 个人学习档案 - {subject_filter}",
        "",
        f"> 导出时间：{now_iso()}",
        "> 内容来源：OCR 题目、AI 拆题结果、巩固题与批改记录。",
        "",
    ]
    exported_count = 0

    for wrong_id, options in selections.items():
        if not options or not any(options.values()):
            continue
        if not user_owns_wrong_question(conn, user, wrong_id):
            continue
        item = get_wrong_question(conn, wrong_id)
        if not item:
            continue
        diagnosis = item.get("diagnosis") or {}
        subject = clean_subject_name(diagnosis.get("subject"))
        if subject_filter != "全部学科" and subject != subject_filter:
            continue

        exported_count += 1
        title = diagnosis.get("core_pattern") or diagnosis.get("topic") or "题目档案"
        lines.extend([f"## {exported_count}. {title}", "", f"- 学科：{subject}", f"- 状态：{item.get('status')}", ""])

        if options.get("question"):
            lines.extend(["### 题目", md_escape(item.get("corrected_text")), ""])
        if options.get("ocr") and item.get("ocr_text"):
            lines.extend(["### OCR 原文", md_escape(item.get("ocr_text")), ""])
        if options.get("student"):
            analysis = diagnosis.get("student_answer_analysis") or {}
            lines.extend(["### 作答情况", f"- 作答存在：{md_escape(analysis.get('answer_presence') or ('已提供' if item.get('student_wrong_answer') else '未提供'))}"])
            if item.get("student_wrong_answer"):
                lines.append(f"- 原始作答/卡点：{md_escape(item.get('student_wrong_answer'))}")
            if analysis.get("extracted_work"):
                lines.append(f"- 结构化作答：{md_escape(analysis.get('extracted_work'))}")
            if analysis.get("answer_status"):
                lines.append(f"- 状态判断：{md_escape(analysis.get('answer_status'))}")
            if analysis.get("likely_issue"):
                lines.append(f"- 主要问题：{md_escape(analysis.get('likely_issue'))}")
            if analysis.get("next_action"):
                lines.append(f"- 下一步：{md_escape(analysis.get('next_action'))}")
            lines.append("")
        if options.get("answer"):
            answer = diagnosis.get("standard_answer") or {}
            lines.extend(["### 答案", md_escape(answer.get("final_answer") or "未返回最终答案"), ""])
        if options.get("analysis"):
            answer = diagnosis.get("standard_answer") or {}
            lines.extend(["### 解析", md_escape(answer.get("concise_solution") or "未返回标准解析"), ""])
            decomposition = diagnosis.get("decomposition") or {}
            if decomposition.get("total_formula"):
                lines.extend(["### 总拆解公式", md_escape(decomposition.get("total_formula")), ""])
        if options.get("strategy"):
            strategy = diagnosis.get("learning_strategy") or {}
            lines.extend(["### 学习方法"])
            if strategy.get("decomposition_answer"):
                lines.append(f"- 怎么拆解：{md_escape(strategy.get('decomposition_answer'))}")
            if strategy.get("make_it_easier"):
                lines.append(f"- 怎么更容易学：{md_escape(strategy.get('make_it_easier'))}")
            if strategy.get("entry_point"):
                lines.append(f"- 入口：{md_escape(strategy.get('entry_point'))}")
            if strategy.get("teacher_hint"):
                lines.append(f"- 老师引导：{md_escape(strategy.get('teacher_hint'))}")
            for step in strategy.get("cognitive_ladder", []) or []:
                lines.append(f"- 认知台阶：{md_escape(step)}")
            for drill in strategy.get("micro_drills", []) or []:
                lines.append(f"- 微练习：{md_escape(drill)}")
            lines.append("")
        if options.get("thinking"):
            decomposition = diagnosis.get("decomposition") or {}
            lines.append("### 详细思路")
            for step in decomposition.get("step_formulas", []) or []:
                lines.append(f"- **{md_escape(step.get('name'))}**：{md_escape(step.get('formula'))}")
                if step.get("operation"):
                    lines.append(f"  - 操作：{md_escape(step.get('operation'))}")
                if step.get("student_trap"):
                    lines.append(f"  - 易错点：{md_escape(step.get('student_trap'))}")
            for model in diagnosis.get("solution_models", []) or []:
                lines.append(f"- **模型：{md_escape(model.get('model_name'))}**")
                for step in model.get("steps", []) or []:
                    lines.append(f"  - {md_escape(step)}")
            lines.append("")
        if options.get("multi"):
            lines.append("### 一题多解/多视角")
            for method in diagnosis.get("multiple_solutions", []) or []:
                lines.append(f"- **{md_escape(method.get('method_name'))}**：{md_escape(method.get('idea'))}")
                for step in method.get("steps", []) or []:
                    lines.append(f"  - {md_escape(step)}")
                if method.get("pros_cons"):
                    lines.append(f"  - 优缺点：{md_escape(method.get('pros_cons'))}")
            lines.append("")
        if options.get("poem"):
            poem = diagnosis.get("poem") or {}
            lines.extend([f"### {md_escape(poem.get('title') or '复盘小诗')}"])
            for line in poem.get("lines", []) or []:
                lines.append(f"> {md_escape(line)}")
            for review in poem.get("line_reviews", []) or []:
                lines.append(f"- {md_escape(review.get('line'))}：{md_escape(review.get('review'))}")
            lines.append("")
        if options.get("similar"):
            lines.append("### 同类题/巩固题")
            for variant in item.get("variants", []) or []:
                lines.append(f"#### 第 {variant.get('level')} 题：{md_escape(variant.get('title'))}")
                lines.append(md_escape(variant.get("stem")))
                lines.append("")
                lines.append(f"- 答案：{md_escape(variant.get('answer'))}")
                lines.append(f"- 解析：{md_escape(variant.get('analysis'))}")
                if variant.get("answers"):
                    latest = variant["answers"][0]
                    grading = latest.get("grading_result") or {}
                    lines.append(f"- 最近作答：{md_escape(latest.get('answer_text'))}")
                    lines.append(f"- 批改：{md_escape(grading.get('comment'))}")
                lines.append("")
        lines.append("---")
        lines.append("")

    if exported_count == 0:
        raise ValueError("没有符合条件的导出内容")
    filename_subject = re.sub(r"[^0-9A-Za-z\u4e00-\u9fff_-]+", "_", subject_filter)
    return {
        "filename": f"{filename_subject}_个人学习档案.md",
        "markdown": "\n".join(lines),
        "count": exported_count,
    }


def update_pass_status(conn: sqlite3.Connection, wrong_id: str) -> None:
    wrong = conn.execute(
        "select review_stage from wrong_questions where id = ?", (wrong_id,)
    ).fetchone()
    variants = conn.execute(
        "select id, level from exercise_variants where wrong_question_id = ?", (wrong_id,)
    ).fetchall()
    if len(variants) < 3:
        return
    results = []
    for variant in variants:
        row = conn.execute(
            """
            select is_correct, hint_count from student_answers
            where exercise_variant_id = ?
            order by submitted_at desc limit 1
            """,
            (variant["id"],),
        ).fetchone()
        results.append({
            "level": variant["level"],
            "is_correct": bool(row and row["is_correct"]),
            "hint_count": int(row["hint_count"] or 0) if row else 0,
        })
    state = transition(results, int(wrong["review_stage"] or 0) if wrong else 0)
    public_status = "passed" if state["state"] == "mastered" else state["state"]
    conn.execute(
        """
        update wrong_questions
        set status = ?, workflow_state = ?, mastery_score = ?, review_stage = ?,
            next_review_at = ?, last_reviewed_at = ?
        where id = ?
        """,
        (
            public_status, state["state"], state["mastery_score"], state["review_stage"],
            state["next_review_at"], state["last_reviewed_at"], wrong_id,
        ),
    )


def build_report(conn: sqlite3.Connection, user: dict | None) -> dict:
    empty = {
        "total_wrong_questions": 0,
        "passed_questions": 0,
        "pass_rate": 0,
        "weak_mothers": [],
        "error_causes": [],
    }
    if not user:
        return empty
    wq_where, wq_params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])
    total = conn.execute(f"select count(*) as n from wrong_questions{wq_where}", wq_params).fetchone()["n"]
    passed = conn.execute(
        f"select count(*) as n from wrong_questions{wq_where}{' and' if wq_where else ' where'} status = 'passed'",
        wq_params,
    ).fetchone()["n"]
    pattern_counts: dict[str, int] = {}
    issue_counts: dict[str, int] = {}
    for row in conn.execute(f"select diagnosis from wrong_questions{wq_where}", wq_params).fetchall():
        diagnosis = json.loads(row["diagnosis"])
        pattern = diagnosis.get("core_pattern") or diagnosis.get("topic") or "未归类"
        pattern_counts[pattern] = pattern_counts.get(pattern, 0) + 1
        answer_status = (diagnosis.get("student_answer_analysis") or {}).get("answer_status") or ""
        if answer_status and answer_status not in {"未提供作答", "待模型进一步判断", "空白不会"}:
            issue_counts[answer_status] = issue_counts.get(answer_status, 0) + 1
        for model in diagnosis.get("solution_models", []):
            for mistake in model.get("common_mistakes", []):
                issue_counts[mistake] = issue_counts.get(mistake, 0) + 1
    return {
        "total_wrong_questions": total,
        "passed_questions": passed,
        "pass_rate": round(passed / total * 100, 1) if total else 0,
        "weak_mothers": sorted(
            [{"name": name, "count": count} for name, count in pattern_counts.items()],
            key=lambda item: item["count"],
            reverse=True,
        )[:5],
        "error_causes": sorted(
            [{"label": label, "count": count} for label, count in issue_counts.items()],
            key=lambda item: item["count"],
            reverse=True,
        )[:5],
    }


def structured_tool_schema(delivery: str) -> str:
    schemas = {
        "docx": """
{
  "title": "文档标题",
  "summary": "200字以内摘要",
  "document_markdown": "完整 Markdown 正文",
  "sections": [
    {
      "title": "章节标题",
      "content": "章节正文，可含列表",
      "questions": [
        {"stem": "题干", "answer": "答案", "analysis": "解析"}
      ]
    }
  ]
}""",
        "pptx": """
{
  "title": "课件标题",
  "summary": "讲评摘要",
  "slides": [
    {
      "title": "页标题",
      "bullets": ["要点1", "要点2"],
      "speaker_notes": "讲评备注"
    }
  ],
  "document_markdown": "完整讲评提纲 Markdown"
}""",
        "report": """
{
  "title": "卷面学情分析报告",
  "summary": "总体结论",
  "score_overview": {"total_score": "100", "estimated_score": "78", "pass_rate": "78%"},
  "module_analysis": [{"module": "模块名", "score_rate": "65%", "issue": "问题", "action": "改进动作"}],
  "question_type_analysis": [{"type": "题型", "loss_points": "失分点", "action": "训练建议"}],
  "weak_knowledge_points": ["薄弱点1", "薄弱点2"],
  "error_causes": [{"cause": "原因", "evidence": "依据", "fix": "修正方案"}],
  "layered_suggestions": {"A": "优等生建议", "B": "中等生建议", "C": "待提升建议"},
  "action_plan_7d": [{"day": "第1天", "tasks": ["任务1", "任务2"]}],
  "document_markdown": "完整报告 Markdown，含表格与分节"
}""",
        "image": """
{
  "title": "配图标题",
  "summary": "配图说明",
  "image_prompt": "用于生成教学配图的详细中文提示词，描述布局、元素、标注、颜色、风格",
  "caption": "图注"
}""",
        "map": """
{
  "title": "知识图谱标题",
  "summary": "图谱说明",
  "nodes": [{"label": "概念", "detail": "解释"}],
  "edges": [{"from": "概念A", "to": "概念B", "label": "关系"}],
  "image_prompt": "可选，生成可视化图谱配图的提示词",
  "document_markdown": "文字版图谱说明"
}""",
    }
    return schemas.get(delivery, schemas["docx"])


def structured_to_markdown(data: dict, delivery: str) -> str:
    if data.get("document_markdown"):
        parts = [f"# {data.get('title') or '生成结果'}", "", data.get("summary") or "", "", data["document_markdown"]]
        return "\n".join(part for part in parts if part is not None)
    lines = [f"# {data.get('title') or '生成结果'}", ""]
    if data.get("summary"):
        lines.extend([str(data["summary"]), ""])
    if delivery == "report":
        overview = data.get("score_overview") or {}
        if overview:
            lines.append("## 成绩概览")
            for key, value in overview.items():
                lines.append(f"- {key}: {value}")
            lines.append("")
        for section_key, title in [
            ("module_analysis", "模块诊断"),
            ("question_type_analysis", "题型诊断"),
            ("weak_knowledge_points", "薄弱知识点"),
            ("error_causes", "失分归因"),
            ("action_plan_7d", "7日行动"),
        ]:
            block = data.get(section_key)
            if not block:
                continue
            lines.append(f"## {title}")
            if isinstance(block, list):
                for item in block:
                    lines.append(f"- {json.dumps(item, ensure_ascii=False) if isinstance(item, (dict, list)) else item}")
            elif isinstance(block, dict):
                for key, value in block.items():
                    lines.append(f"- {key}: {value}")
            lines.append("")
    if delivery == "pptx":
        lines.append("## 课件页")
        for idx, slide in enumerate(data.get("slides") or [], start=1):
            lines.append(f"### 第{idx}页 {slide.get('title') or ''}")
            for bullet in slide.get("bullets") or []:
                lines.append(f"- {bullet}")
            if slide.get("speaker_notes"):
                lines.append(f"> 备注：{slide['speaker_notes']}")
            lines.append("")
    if delivery == "map":
        lines.append("## 知识节点")
        for node in data.get("nodes") or []:
            lines.append(f"- **{node.get('label') or ''}**：{node.get('detail') or ''}")
        lines.append("")
    if data.get("caption"):
        lines.append(f"图注：{data['caption']}")
    return "\n".join(lines).strip()


def build_tool_artifacts(conn: sqlite3.Connection, tool: dict, run_id: str, data: dict) -> tuple[list[dict], dict]:
    delivery = tool.get("delivery") or "docx"
    title = data.get("title") or tool.get("label") or "生成结果"
    safe = export_utils.safe_filename(title, run_id[:8])
    prefix = f"{run_id}_{safe}"
    artifacts: list[dict] = []
    report: dict = {}

    body = data.get("document_markdown") or structured_to_markdown(data, delivery)
    sections = data.get("sections") or []

    if delivery in {"docx", "report"}:
        docx_path = EXPORT_DIR / f"{prefix}.docx"
        if delivery == "report":
            report = {
                "title": title,
                "summary": data.get("summary") or "",
                "score_overview": data.get("score_overview") or {},
                "module_analysis": data.get("module_analysis") or [],
                "question_type_analysis": data.get("question_type_analysis") or [],
                "weak_knowledge_points": data.get("weak_knowledge_points") or [],
                "error_causes": data.get("error_causes") or [],
                "layered_suggestions": data.get("layered_suggestions") or {},
                "action_plan_7d": data.get("action_plan_7d") or [],
            }
            report_sections = []
            if report["score_overview"]:
                report_sections.append({"title": "成绩概览", "content": "\n".join(f"- {k}: {v}" for k, v in report["score_overview"].items())})
            if report["module_analysis"]:
                report_sections.append({"title": "模块诊断", "content": "\n".join(json.dumps(item, ensure_ascii=False) for item in report["module_analysis"])})
            if report["question_type_analysis"]:
                report_sections.append({"title": "题型诊断", "content": "\n".join(json.dumps(item, ensure_ascii=False) for item in report["question_type_analysis"])})
            if report["weak_knowledge_points"]:
                report_sections.append({"title": "薄弱知识点", "content": "\n".join(f"- {x}" for x in report["weak_knowledge_points"])})
            if report["error_causes"]:
                report_sections.append({"title": "失分归因", "content": "\n".join(json.dumps(item, ensure_ascii=False) for item in report["error_causes"])})
            if report["layered_suggestions"]:
                report_sections.append({"title": "分层建议", "content": "\n".join(f"- {k}: {v}" for k, v in report["layered_suggestions"].items())})
            if report["action_plan_7d"]:
                report_sections.append({"title": "7日行动方案", "content": "\n".join(json.dumps(item, ensure_ascii=False) for item in report["action_plan_7d"])})
            if body:
                report_sections.append({"title": "完整报告", "content": body})
            export_utils.write_docx(docx_path, title, body, report_sections)
        else:
            export_utils.write_docx(docx_path, title, body, sections or None)
        artifacts.append(export_utils.artifact_record(docx_path.name, f"/exports/{docx_path.name}", "docx", "下载 Word"))

    if delivery == "pptx":
        pptx_path = EXPORT_DIR / f"{prefix}.pptx"
        slides = data.get("slides") or []
        if not slides and body:
            slides = [{"title": "讲评提纲", "bullets": [line[2:] for line in body.splitlines() if line.strip().startswith("- ")][:8]}]
        export_utils.write_pptx(pptx_path, title, slides)
        artifacts.append(export_utils.artifact_record(pptx_path.name, f"/exports/{pptx_path.name}", "pptx", "下载 PowerPoint"))

    if delivery in {"docx", "pptx", "report", "map"}:
        md_path = EXPORT_DIR / f"{prefix}.md"
        export_utils.write_markdown(md_path, title, body)
        artifacts.append(export_utils.artifact_record(md_path.name, f"/exports/{md_path.name}", "markdown", "下载 Markdown"))

    if delivery == "map":
        html_path = EXPORT_DIR / f"{prefix}.html"
        export_utils.write_knowledge_map_html(html_path, title, data.get("nodes") or [], data.get("edges") or [])
        artifacts.append(export_utils.artifact_record(html_path.name, f"/exports/{html_path.name}", "html", "打开知识图谱"))

    if delivery == "image":
        md_path = EXPORT_DIR / f"{prefix}.md"
        export_utils.write_markdown(md_path, title, body or data.get("summary") or data.get("caption") or "")
        artifacts.append(export_utils.artifact_record(md_path.name, f"/exports/{md_path.name}", "markdown", "下载说明 Markdown"))

    image_prompt = (data.get("image_prompt") or "").strip()
    if delivery in {"image", "map"} and image_prompt:
        image_model = get_image_model(conn)
        image_bytes = call_image_generation(image_model, image_prompt)
        image_path = EXPORT_DIR / f"{prefix}.png"
        image_path.write_bytes(image_bytes)
        artifacts.append(export_utils.artifact_record(image_path.name, f"/exports/{image_path.name}", "image", "下载配图 PNG"))

    return artifacts, report


def portal_tool_prompt(tool: dict, subject: str, user_input: str) -> list[dict]:
    delivery = tool.get("delivery") or "docx"
    if delivery == "diagnose":
        delivery = "docx"
    mode_guides = {
        "document": "整理题卷/试卷为规范文档结构，修正 OCR 换行，保留题号、选项、答案区、解析区。",
        "analysis": "基于完整试卷或作答材料，输出符合教研规范的卷面学情诊断：成绩概览、模块/题型诊断、失分归因、分层建议、7日行动。",
        "variant": "生成同结构、条件替换、迁移挑战三类变式题，每题带答案与解析。",
        "question": "围绕知识点批量命题，按基础/提高/压轴分层，每题带答案、解析和易错提醒。",
        "wrong": "对错题做归因、拆题、同类题迁移和复盘建议。",
        "image": "根据教学需求生成可直接用于课堂的配图，并给出详细 image_prompt。",
        "ppt": "生成可直接授课的讲评课件，至少 8 页，含典型题、互动提问、课后练习与讲评备注。",
        "review": "输出审题训练方案：关键词、隐含条件、设问类型、陷阱、第一步动作与练习题。",
        "decompose": "把大题拆成采分点、公式/材料依据、步骤、易错点和答题模板。",
        "english": "生成词汇练习卷：词义、拼写、语境填空、翻译和答案。",
        "coverage": "检测考点覆盖：已覆盖、遗漏、重复、难度比例和补题建议。",
        "plan": "生成冲刺计划：目标、每日任务、错题复盘、输出物和验收标准。",
        "notes": "整理课堂笔记：概念、例题、方法、易错点和课后任务。",
        "preview": "生成预习单：预习目标、关键词、问题链、任务和检查。",
        "loss": "分析失分原因：知识、审题、表达、计算、时间和心理因素，并给修正动作。",
        "sense": "输出题感训练：判型、入口、常用模型和最小验证题。",
        "map": "构建知识点图谱：节点、关系链、例题入口，并给出可视化 image_prompt。",
        "multi": "给出多种解法/视角，比较适用场景、优缺点和迁移建议。",
        "explain": "精讲知识点：定义、比喻、例题、误区和小练习。",
        "writing": "优化作文/表达：改词、句式升级、段落建议和评分理由。",
        "score": "拆解目标分：分数差距、题型收益、每日练习和复盘节点。",
    }
    guide = mode_guides.get(tool.get("mode"), "生成可直接用于教学或学习的结构化结果。")
    schema = structured_tool_schema(delivery)
    system = f"""你是 AI错题拆博士 的专业教研引擎。当前工具：{tool['label']}。

要求：
1. 输出必须是合法 JSON，不要 Markdown，不要代码块，不要额外解释。
2. 内容必须符合 K12 / 中考 / 高考 / 本硕博解题辅导规范，可直接交付教师、学生或家长使用。
3. 结构完整、结论先行、建议可执行；禁止空泛套话。
4. 若材料不足，可合理补全并写入 summary。

工具任务：{guide}

JSON 格式：
{schema}"""
    user = f"""学科/场景：{subject or '自动识别'}

用户材料：
{user_input}

请严格按 JSON 格式输出。"""
    return [{"role": "system", "content": system}, {"role": "user", "content": user}]


def run_portal_tool(conn: sqlite3.Connection, data: dict, user: dict | None) -> dict:
    tool_id = data.get("tool_id") or "knowledge-explain"
    tool = get_portal_tool(tool_id)
    if not tool:
        raise ValueError("工具不存在")
    if tool.get("delivery") == "diagnose":
        raise ValueError("该工具请使用错题拆解工作台")
    input_text = (data.get("input_text") or "").strip()
    if not input_text:
        raise ValueError("请先输入要处理的题目、试卷、知识点或学习材料")
    subject = data.get("subject") or "自动识别"
    model = get_model(conn, data.get("model_id"))
    raw = minimax_chat(
        model,
        portal_tool_prompt(tool, subject, input_text),
        max_tokens=int(model.get("max_tokens") or 7000),
        temperature=float(model.get("temperature") or 0.35),
    )
    structured = extract_json(raw)
    delivery = tool.get("delivery") or "docx"
    output = structured_to_markdown(structured, delivery)
    run_id = str(uuid.uuid4())
    artifacts, report = build_tool_artifacts(conn, tool, run_id, structured)
    if not artifacts:
        raise RuntimeError("未能生成可下载文件，请检查模型返回内容")
    user_id = user.get("id") if user else None
    created_at = now_iso()
    conn.execute(
        """
        insert into tool_runs
        (id, tool_id, tool_label, subject, input_text, output_text, model_id, user_id, artifacts, report, created_at)
        values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            run_id,
            tool["id"],
            tool["label"],
            subject,
            input_text,
            output,
            model.get("id"),
            user_id,
            json.dumps(artifacts, ensure_ascii=False),
            json.dumps(report, ensure_ascii=False) if report else "",
            created_at,
        ),
    )
    if user_id:
        conn.execute("update app_users set credits = max(credits - 1, 0) where id = ?", (user_id,))
    return {
        "id": run_id,
        "tool_id": tool["id"],
        "tool_label": tool["label"],
        "subject": subject,
        "input_text": input_text,
        "output_text": output,
        "artifacts": artifacts,
        "report": report,
        "model_id": model.get("id"),
        "created_at": created_at,
    }


def rag_user_filter(user: dict | None) -> tuple[str, list]:
    if user and user_is_admin(user):
        return "", []
    if user:
        return " where (user_id = ? or user_id is null or trim(user_id) = '')", [user["id"]]
    return " where (user_id is null or trim(user_id) = '')", []


def chunk_text_for_rag(text: str, max_chars: int = 900, overlap: int = 140) -> list[str]:
    clean = re.sub(r"\n{3,}", "\n\n", (text or "").strip())
    if not clean:
        return []
    blocks = [part.strip() for part in re.split(r"\n\s*\n", clean) if part.strip()]
    chunks: list[str] = []
    current = ""
    for block in blocks:
        if len(block) > max_chars:
            if current:
                chunks.append(current.strip())
                current = ""
            start = 0
            while start < len(block):
                chunks.append(block[start:start + max_chars].strip())
                start += max_chars - overlap
            continue
        if current and len(current) + len(block) + 2 > max_chars:
            chunks.append(current.strip())
            current = current[-overlap:] if overlap and len(current) > overlap else ""
        current = f"{current}\n\n{block}".strip() if current else block
    if current.strip():
        chunks.append(current.strip())
    return [chunk for chunk in chunks if chunk]


def fallback_rag_score(query: str, content: str, subject: str = "") -> float:
    q = re.sub(r"\s+", "", query or "").lower()
    c = (content or "").lower()
    if not q or not c:
        return 0.0
    score = 0.0
    for token in re.findall(r"[0-9A-Za-z_]+|[\u4e00-\u9fff]{2,}", query or ""):
        if token.lower() in c:
            score += min(4.0, len(token) / 2)
    for i in range(max(0, len(q) - 1)):
        if q[i:i + 2] in c:
            score += 0.08
    if subject and subject != "自动识别" and subject in content:
        score += 1.5
    return score


def public_rag_document(row: sqlite3.Row | dict) -> dict:
    item = dict(row)
    return {
        "id": item.get("id"),
        "title": item.get("title"),
        "filename": item.get("filename"),
        "subject": clean_subject_name(item.get("subject")),
        "source_type": item.get("source_type"),
        "chars": int(item.get("chars") or 0),
        "created_at": item.get("created_at"),
    }


def create_rag_document(conn: sqlite3.Connection, data: dict, user: dict | None) -> dict:
    filename = data.get("filename") or "paste.txt"
    subject = clean_subject_name(data.get("subject") or "自动识别")
    title = (data.get("title") or Path(filename).stem or "知识资料").strip()[:120]
    text = (data.get("text") or "").strip()
    source_type = "paste"
    if data.get("file_data_url"):
        extracted = extract_document_from_data_url(filename, data.get("file_data_url") or "")
        text = (extracted.get("text") or "").strip()
        filename = extracted.get("filename") or filename
        title = (data.get("title") or Path(filename).stem or title).strip()[:120]
        source_type = "file"
    if len(text) < 20:
        raise ValueError("资料内容太短，无法入库检索")
    doc_id = str(uuid.uuid4())
    user_id = user["id"] if user else None
    created_at = now_iso()
    chunks = chunk_text_for_rag(text)
    if not chunks:
        raise ValueError("未能切分出有效知识片段")
    conn.execute(
        """
        insert into rag_documents (id, user_id, title, filename, subject, source_type, chars, created_at)
        values (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (doc_id, user_id, title, filename, subject, source_type, len(text), created_at),
    )
    saved_chunks = []
    for index, chunk in enumerate(chunks):
        chunk_id = str(uuid.uuid4())
        conn.execute(
            """
            insert into rag_chunks (id, document_id, user_id, chunk_index, title, subject, content, chars, created_at)
            values (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (chunk_id, doc_id, user_id, index, title, subject, chunk, len(chunk), created_at),
        )
        try:
            conn.execute(
                "insert into rag_chunks_fts (chunk_id, document_id, user_id, title, subject, content) values (?, ?, ?, ?, ?, ?)",
                (chunk_id, doc_id, user_id or "", title, subject, chunk),
            )
        except sqlite3.OperationalError:
            pass
        saved_chunks.append({"id": chunk_id, "index": index, "chars": len(chunk)})
    return {
        "id": doc_id,
        "title": title,
        "filename": filename,
        "subject": subject,
        "source_type": source_type,
        "chars": len(text),
        "chunk_count": len(saved_chunks),
        "created_at": created_at,
    }


def search_rag(
    conn: sqlite3.Connection,
    query: str,
    subject: str = "",
    user: dict | None = None,
    limit: int = 5,
    source_type: str | None = None,
) -> list[dict]:
    if source_type == gaokao_rag.GAOKAO_RAG_SOURCE_TYPE:
        return gaokao_rag.search_gaokao_rag(conn, query, limit)
    query = (query or "").strip()
    if not query:
        return []
    limit = max(1, min(int(limit or 5), 12))
    subject = clean_subject_name(subject or "自动识别")
    user_clause, params = rag_user_filter(user)
    base_where = user_clause.replace(" where ", "") if user_clause else "1=1"
    source_sql = ""
    source_params: list = []
    if source_type:
        source_sql = " and d.source_type = ?"
        source_params.append(source_type)
    subject_sql = ""
    subject_params: list = []
    if subject and subject != "自动识别":
        subject_sql = " and (c.subject = ? or c.subject = '自动识别' or c.subject is null or trim(c.subject) = '')"
        subject_params.append(subject)
    candidates: list[dict] = []
    try:
        fts_query = " OR ".join(re.findall(r"[0-9A-Za-z_]+|[\u4e00-\u9fff]{2,}", query)) or query
        rows = conn.execute(
            f"""
            select c.*, d.source_type, bm25(rag_chunks_fts) as rank
            from rag_chunks_fts
            join rag_chunks c on c.id = rag_chunks_fts.chunk_id
            join rag_documents d on d.id = c.document_id
            where rag_chunks_fts match ? and {base_where}{subject_sql}{source_sql}
            order by rank limit ?
            """,
            [fts_query, *params, *subject_params, *source_params, limit * 3],
        ).fetchall()
        for row in rows:
            item = row_to_dict(row)
            item["score"] = round(100 / (1 + abs(float(row["rank"] or 0))), 2)
            candidates.append(item)
    except sqlite3.OperationalError:
        candidates = []
    if len(candidates) < limit:
        rows = conn.execute(
            f"""
            select c.*, d.source_type
            from rag_chunks c
            join rag_documents d on d.id = c.document_id
            where {base_where}{subject_sql}{source_sql}
            order by c.created_at desc limit 200
            """,
            [*params, *subject_params, *source_params],
        ).fetchall()
        seen = {item.get("id") for item in candidates}
        fallback = []
        for row in rows:
            item = row_to_dict(row)
            if item.get("id") in seen:
                continue
            score = fallback_rag_score(query, f"{item.get('title')}\n{item.get('subject')}\n{item.get('content')}", subject)
            if score > 0:
                item["score"] = round(score, 2)
                fallback.append(item)
        fallback.sort(key=lambda item: item.get("score", 0), reverse=True)
        candidates.extend(fallback[:limit * 2])
    candidates.sort(key=lambda item: item.get("score", 0), reverse=True)
    results = []
    for item in candidates[:limit]:
        results.append({
            "id": item.get("id"),
            "document_id": item.get("document_id"),
            "chunk_index": item.get("chunk_index"),
            "title": item.get("title"),
            "subject": clean_subject_name(item.get("subject")),
            "content": item.get("content"),
            "score": item.get("score", 0),
            "created_at": item.get("created_at"),
        })
    return results


def build_rag_context_for_prompt(chunks: list[dict]) -> str:
    if not chunks:
        return ""
    blocks = []
    for index, chunk in enumerate(chunks, start=1):
        blocks.append(
            f"【资料{index}｜{chunk.get('title') or '知识片段'}｜{chunk.get('subject') or '未标注'}｜相关度{chunk.get('score', 0)}】\n{compact_text(chunk.get('content'), 900)}"
        )
    return "\n\n".join(blocks)


def retrieve_gaokao_evidence(
    conn: sqlite3.Connection,
    query: str,
    subject: str = "自动识别",
    *,
    limit: int = 6,
) -> tuple[list[dict], str]:
    if not gaokao_rag.should_use_gaokao_rag(subject, query):
        return [], ""
    hits = gaokao_rag.search_gaokao_rag(conn, query, limit=limit)
    return hits, gaokao_rag.build_gaokao_rag_context(hits)


def merge_gaokao_search_results(
    conn: sqlite3.Connection,
    query: str,
    *,
    zone: str = ZONE_GAOKAO_MATH,
    limit: int = 8,
) -> dict:
    bank_hits = search_questions(conn, query, zone=zone, limit=limit)
    bank_hits = rerank_hits(query, bank_hits, text_key="stem")
    rag_hits = gaokao_rag.search_gaokao_rag(conn, query, limit=min(6, limit))
    method_packages = search_question_method_packages(conn, query, limit=min(6, limit))
    rag_context = gaokao_rag.build_gaokao_rag_context(rag_hits)
    package_context = method_packages_context(method_packages)
    return {
        "bank_hits": bank_hits,
        "rag_hits": rag_hits,
        "method_packages": method_packages,
        "rag_context": "\n\n".join(part for part in (rag_context, package_context) if part),
        "citations": gaokao_rag.rag_citations_from_hits(rag_hits),
    }


def generate_rag_quiz_from_hits(topic: str, hits: list[dict], count: int = 5) -> dict:
    topic = compact_text(topic or "综合复习", 80)
    questions = []
    for index, hit in enumerate(hits[:count], start=1):
        title = hit.get("title") or "知识资料"
        content = compact_text(hit.get("content"), 180)
        short = index % 2 == 0
        questions.append({
            "type": "short" if short else "single",
            "question": f"根据资料《{title}》，请判断或说明：{topic} 中最容易出错的关键点是什么？",
            "options": [] if short else [
                compact_text(content, 48),
                "只背题干，不需要理解条件",
                "不需要引用资料证据",
                "可以忽略错因复盘",
            ],
            "answer": content,
            "analysis": "答案应紧扣资料证据，并能说出这个知识点对应的错因、题型入口和训练动作。",
            "source": title,
        })
    if not questions:
        questions.append({
            "type": "short",
            "question": f"围绕“{topic}”，请先上传教材、错题总结或课堂笔记，再生成基于证据的自测题。",
            "options": [],
            "answer": "当前 RAG 未命中资料。",
            "analysis": "系统需要先有资料片段，才能可靠出题。",
            "source": "RAG 知识库",
        })
    return {"topic": topic, "questions": questions, "citations": hits, "ai_ready": False}


def generate_rag_study_os(topic: str, hits: list[dict]) -> dict:
    topic = compact_text(topic or "综合自学", 80)
    plan = []
    cards = []
    hints = []
    review = []
    for index, hit in enumerate(hits[:6], start=1):
        title = hit.get("title") or f"{topic}资料{index}"
        content = compact_text(hit.get("content"), 220)
        plan.append({
            "title": f"第{index}步：吃透《{title}》",
            "goal": f"先复述资料，再标出题型入口、易错条件和标准动作：{content}",
            "evidence": title,
        })
        cards.append({
            "front": title,
            "back": content,
            "source": title,
        })
        hints.append({
            "question": f"《{title}》对应的错题突破入口是什么？",
            "try_first": "先遮住解析，用 2 句话写出自己的判断。",
            "hint": "从题型、条件、目标、常见错因四个角度找。",
            "answer": content,
        })
        review.append({
            "mistake": f"把《{title}》当成孤立知识点背诵",
            "fix": "订正时写清：我错在哪里、正确入口是什么、下一道同类题先看什么。",
        })
    if not hits:
        plan.append({
            "title": "先建设错题资料库",
            "goal": "上传错题总结、教材、讲义或评分标准，系统会切片进入 RAG，再生成自学方案。",
            "evidence": "暂无 RAG 命中",
        })
    return {
        "topic": topic,
        "plan": plan,
        "flashcards": cards,
        "hints": hints,
        "review": review,
        "citations": hits,
        "ai_ready": False,
    }


def normalize_agent_result(result: dict, question_text: str, subject: str, student_answer: str) -> dict:
    if not isinstance(result, dict):
        result = {}
    result = repair_text_tree(result)
    subject_name = result.get("subject") or subject or "自动识别"
    structured = result.get("structured_question") or {}
    quick = result.get("quick_answer") or {}
    result.setdefault("title", compact_text(structured.get("target") or result.get("question_type") or question_text, 42))
    result.setdefault("subject", subject_name)
    result.setdefault("question_type", result.get("question_type") or "待归纳题型")
    result.setdefault("difficulty", 3)
    result.setdefault("confidence", 0.72)
    result["quick_answer"] = {
        "how_to_decompose": quick.get("how_to_decompose") or "先识别学科与题型，再拆已知条件、设问目标、可用模型、执行步骤和校验动作。",
        "make_it_easier": quick.get("make_it_easier") or "先做一个更小的原型题，把核心步骤练熟，再迁移回原题。",
        "first_entry": quick.get("first_entry") or result.get("question_type") or "题型入口",
    }
    structured.setdefault("cleaned_question", question_text)
    structured.setdefault("known_conditions", [])
    structured.setdefault("target", result.get("title") or "题目目标")
    structured.setdefault("hidden_constraints", [])
    structured.setdefault("student_work", student_answer or "")
    result["structured_question"] = structured
    answer_analysis = result.get("student_answer_analysis") or {}
    answer_analysis.setdefault("answer_presence", "已提供" if student_answer else "未提供")
    answer_analysis.setdefault("answer_status", "待模型进一步判断" if student_answer else "未提供作答")
    answer_analysis.setdefault("likely_issue", "根据学生作答定位错因" if student_answer else "未提供作答，暂不判断错因")
    answer_analysis.setdefault("evidence", [])
    answer_analysis.setdefault("next_action", "对照分层步骤复盘" if student_answer else "可补充学生作答后再诊断")
    result["student_answer_analysis"] = answer_analysis
    result.setdefault("solution_model", {
        "model_name": "识别题型-拆条件-选模型-执行-校验-复盘模型",
        "applies_when": "适用于高考全科题目的通用拆解",
        "step_formula": "识别题型→拆条件→选模型→执行→校验→复盘",
        "steps": ["识别入口", "拆分条件", "选择模型", "规范作答", "反向校验"],
        "checkpoints": ["答案是否回答设问", "步骤是否有依据"],
    })
    result.setdefault("multiple_solutions", [])
    result.setdefault("standard_solution", "模型未返回标准解析，请重新生成或补充题干。")
    result.setdefault("final_answer", "模型未返回最终答案。")
    result.setdefault("score_points", [])
    result.setdefault("common_mistakes", [])
    result.setdefault("training_tasks", [])
    result.setdefault("mother_question_reserved", {
        "status": "prompt_reserved",
        "name": result.get("question_type") or "题型原型待沉淀",
        "abstract_pattern": "后续接入母题库后沉淀",
        "future_interface_hint": "后续接入 /api/mother-questions 或 RAG",
    })
    result.setdefault("fun_analogy", {"theme": "拆题路线图", "overview": "把题目当成一条任务流水线逐层拆开。", "steps": []})
    result.setdefault("poem", {"title": "解题复盘诀", "lines": [], "line_reviews": []})
    result.setdefault("archive_payload", {
        "subject": subject_name,
        "question": question_text,
        "answer": result.get("final_answer"),
        "analysis": result.get("standard_solution"),
        "detailed_thinking": result.get("quick_answer", {}).get("how_to_decompose"),
        "similar_questions": [task.get("stem") for task in result.get("training_tasks", []) if isinstance(task, dict) and task.get("stem")],
        "tags": [result.get("question_type") or "题型待归纳"],
    })

    existing_layers = {}
    for layer in result.get("layers") or []:
        if isinstance(layer, dict) and layer.get("key"):
            existing_layers[layer["key"]] = layer
    normalized_layers = []
    for template in AGENT_LAYERS:
        layer = existing_layers.get(template["key"], {})
        normalized_layers.append({
            "key": template["key"],
            "name": layer.get("name") or template["name"],
            "status": layer.get("status") or "done",
            "summary": layer.get("summary") or template["role"],
            "input": layer.get("input") or compact_text(question_text, 120),
            "output": layer.get("output") or layer.get("summary") or template["role"],
            "quality_gate": layer.get("quality_gate") or template["quality_gate"],
            "next_action": layer.get("next_action") or "进入下一层处理",
        })
    result["layers"] = normalized_layers
    return _replace_placeholder_tree(repair_text_tree(result))


def agent_prompt_messages(subject: str, question_text: str, student_answer: str, rag_context: str = "") -> list[dict]:
    layer_spec = "\n".join(f"- {layer['key']}: {layer['name']}｜{layer['role']}｜质检：{layer['quality_gate']}" for layer in AGENT_LAYERS)
    user = f"""学科：{subject or '自动识别'}

题目：
{question_text}

学生作答/批注/卡点：
{student_answer or '未提供'}

RAG 检索资料：
{rag_context or '未启用或未检索到相关资料。'}

要求：如果 RAG 检索资料不为空，请优先把资料作为依据，但必须自行校验资料是否适用于本题；若资料与题目冲突，要说明“资料不适用/需谨慎”。输出 JSON 中必须写入 rag_context 字段，包含 used、evidence_count、citations、how_used。

智能体层次清单：
{layer_spec}

请严格按 AGENT_SOLVE_PROMPT 的 JSON 格式输出。"""
    return [{"role": "system", "content": AGENT_SOLVE_PROMPT}, {"role": "user", "content": user}]


def solve_with_agent(conn: sqlite3.Connection, data: dict, user: dict | None) -> dict:
    question_text = (data.get("question_text") or "").strip()
    if not question_text:
        raise ValueError("请先输入题目内容")
    subject = (data.get("subject") or "自动识别").strip()
    student_answer = (data.get("student_answer") or data.get("student_wrong_answer") or "").strip()
    model = get_model(conn, data.get("model_id"))
    use_rag = data.get("use_rag", True) is not False
    rag_hits: list[dict] = []
    if use_rag:
        if gaokao_rag.should_use_gaokao_rag(subject, question_text):
            rag_hits = gaokao_rag.search_gaokao_rag(conn, "\n".join([subject, question_text, student_answer]), int(data.get("rag_limit") or 6))
        if not rag_hits:
            rag_hits = search_rag(conn, "\n".join([subject, question_text, student_answer]), subject, user, int(data.get("rag_limit") or 5))
    rag_context = gaokao_rag.build_gaokao_rag_context(rag_hits) if rag_hits and rag_hits[0].get("source") == "gaokao_rag" else build_rag_context_for_prompt(rag_hits)
    raw = minimax_chat(
        model,
        agent_prompt_messages(subject, question_text, student_answer, rag_context),
        max_tokens=int(model.get("max_tokens") or 8000),
        temperature=float(model.get("temperature") or 0.28),
    )
    structured = normalize_agent_result(extract_json(raw), question_text, subject, student_answer)
    structured["rag_context"] = {
        "used": bool(rag_hits),
        "evidence_count": len(rag_hits),
        "citations": [
            {
                "chunk_id": item.get("id"),
                "document_id": item.get("document_id"),
                "title": item.get("title"),
                "subject": item.get("subject"),
                "score": item.get("score"),
                "excerpt": compact_text(item.get("content"), 180),
            }
            for item in rag_hits
        ],
        "how_used": structured.get("rag_context", {}).get("how_used") if isinstance(structured.get("rag_context"), dict) else "已作为解题依据注入提示词" if rag_hits else "未检索到可用资料",
    }
    run_id = str(uuid.uuid4())
    created_at = now_iso()
    user_id = user["id"] if user else None
    conn.execute(
        """
        insert into agent_runs
        (id, user_id, subject, question_text, student_answer, model_id, result, status, created_at)
        values (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            run_id,
            user_id,
            clean_subject_name(structured.get("subject") or subject),
            question_text,
            student_answer,
            model.get("id"),
            json.dumps(structured, ensure_ascii=False),
            "completed",
            created_at,
        ),
    )
    return {
        "id": run_id,
        "user_id": user_id,
        "subject": clean_subject_name(structured.get("subject") or subject),
        "question_text": question_text,
        "student_answer": student_answer,
        "model_id": model.get("id"),
        "result": structured,
        "status": "completed",
        "created_at": created_at,
    }


def user_owns_agent_run(user: dict | None, run: dict | None) -> bool:
    if not run:
        return False
    if not run.get("user_id"):
        return True
    if not user:
        return False
    if user_is_admin(user):
        return True
    return run.get("user_id") == user.get("id")


def history_item(item_id: str, item_type: str, label: str, title: str, summary: str, created_at: str, **extra) -> dict:
    row = {
        "id": item_id,
        "type": item_type,
        "label": label,
        "title": title or label,
        "summary": compact_text(summary, 220),
        "created_at": created_at or "",
    }
    row.update(extra)
    return row


def build_generation_history(conn: sqlite3.Connection, user: dict | None) -> dict:
    if not user:
        return {"items": [], "counts": {}, "total": 0}

    wq_where, wq_params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])
    join_where, join_params = ("", []) if user_is_admin(user) else (" where w.user_id = ?", [user["id"]])
    tool_where, tool_params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])
    export_where, export_params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])
    agent_where, agent_params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])

    items: list[dict] = []
    for row in conn.execute(
        f"select * from wrong_questions{wq_where} order by created_at desc",
        wq_params,
    ).fetchall():
        item = row_to_dict(row)
        diagnosis = item.get("diagnosis") or {}
        items.append(history_item(
            f"diagnosis:{item['id']}", "diagnosis", "AI拆题",
            diagnosis.get("core_pattern") or diagnosis.get("topic") or "题目拆解",
            item.get("corrected_text") or "", item.get("created_at"),
            wrong_question_id=item["id"], subject=clean_subject_name(diagnosis.get("subject")), status=item.get("status"),
            institution_id=item.get("institution_id"), institution_name=item.get("institution_name"), institution_badge=item.get("institution_badge"),
        ))
    for row in conn.execute(f"""
        select v.*, w.diagnosis from exercise_variants v
        join wrong_questions w on w.id = v.wrong_question_id
        {join_where}
        order by v.created_at desc
    """, join_params).fetchall():
        variant = row_to_dict(row)
        diagnosis = variant.get("diagnosis") or {}
        items.append(history_item(
            f"variant:{variant['id']}", "variant", "同类变式",
            f"第 {variant.get('level')} 题：{variant.get('title')}", variant.get("stem") or "", variant.get("created_at"),
            wrong_question_id=variant.get("wrong_question_id"), subject=clean_subject_name(diagnosis.get("subject")),
        ))
    for row in conn.execute(f"""
        select a.*, v.title, v.wrong_question_id, w.diagnosis
        from student_answers a
        join exercise_variants v on v.id = a.exercise_variant_id
        join wrong_questions w on w.id = v.wrong_question_id
        {join_where}
        order by a.submitted_at desc
    """, join_params).fetchall():
        answer = row_to_dict(row)
        grading = answer.get("grading_result") or {}
        diagnosis = answer.get("diagnosis") or {}
        items.append(history_item(
            f"grading:{answer['id']}", "grading", "AI批改", answer.get("title") or "变式批改",
            grading.get("comment") or grading.get("detected_issue") or answer.get("answer_text") or "", answer.get("submitted_at"),
            wrong_question_id=answer.get("wrong_question_id"), subject=clean_subject_name(diagnosis.get("subject")),
            score=grading.get("score"), is_correct=bool(answer.get("is_correct")),
        ))
    for row in conn.execute(f"""
        select c.*, w.diagnosis from study_cards c
        join wrong_questions w on w.id = c.wrong_question_id
        {join_where}
        order by c.created_at desc
    """, join_params).fetchall():
        card = row_to_dict(row)
        diagnosis = card.get("diagnosis") or {}
        items.append(history_item(
            f"card:{card['id']}", "card", "学习卡片",
            diagnosis.get("core_pattern") or diagnosis.get("topic") or "错题训练卡",
            f"{card.get('model')} · {card.get('size')} · {card.get('quality')}", card.get("created_at"),
            wrong_question_id=card.get("wrong_question_id"), subject=clean_subject_name(diagnosis.get("subject")), image_url=card.get("image_url"),
        ))
    for row in conn.execute(
        f"select * from tool_runs{tool_where} order by created_at desc",
        tool_params,
    ).fetchall():
        run = row_to_dict(row)
        items.append(history_item(
            f"tool:{run['id']}", "tool", run.get("tool_label") or "AI工具", run.get("tool_label") or "工具生成",
            run.get("output_text") or run.get("input_text") or "", run.get("created_at"),
            tool_run_id=run.get("id"), tool_id=run.get("tool_id"), subject=clean_subject_name(run.get("subject")),
        ))
    for row in conn.execute(
        f"select * from agent_runs{agent_where} order by created_at desc",
        agent_params,
    ).fetchall():
        run = row_to_dict(row)
        result = run.get("result") or {}
        quick = result.get("quick_answer") or {}
        items.append(history_item(
            f"agent:{run['id']}", "agent", "解题智能体", result.get("title") or result.get("question_type") or "分层解题",
            quick.get("how_to_decompose") or result.get("standard_solution") or run.get("question_text") or "",
            run.get("created_at"),
            agent_run_id=run.get("id"), subject=clean_subject_name(result.get("subject") or run.get("subject")),
            status=run.get("status"),
        ))
    for row in conn.execute(
        f"select * from profile_exports{export_where} order by created_at desc",
        export_params,
    ).fetchall():
        export = row_to_dict(row)
        items.append(history_item(
            f"profile_export:{export['id']}", "profile_export", "档案导出", export.get("filename") or "个人学习档案",
            f"{export.get('subject')} · {export.get('count')} 道题", export.get("created_at"),
            export_id=export.get("id"), subject=clean_subject_name(export.get("subject")),
        ))
    items.sort(key=lambda item: item.get("created_at") or "", reverse=True)
    counts: dict[str, int] = {}
    for item in items:
        counts[item["type"]] = counts.get(item["type"], 0) + 1
    return {"items": items, "counts": counts, "total": len(items)}


def institution_context_for_app_user(user: dict | None) -> dict:
    if not user:
        return {}
    store = org_inst_store.load_store()
    email = str(user.get("email") or "").strip()
    member = org_inst_store.find_member_by_username(store, email)
    if not member:
        return {}
    inst = org_inst_store.get_institution_by_id(store, member.get("institution_id"))
    return {
        "institution_id": member.get("institution_id") or "",
        "institution_name": member.get("institution_name") or (inst or {}).get("name") or "",
        "institution_badge": member.get("institution_badge") or (inst or {}).get("name") or "",
    }


def masked_email(value: str | None) -> str:
    email = (value or "").strip()
    if "@" not in email:
        return "学习者"
    name, domain = email.split("@", 1)
    if len(name) <= 2:
        shown = name[:1] + "*"
    else:
        shown = name[:2] + "***" + name[-1:]
    return f"{shown}@{domain}"


def public_profile_share(row: sqlite3.Row | dict) -> dict:
    item = dict(row)
    try:
        permissions = json.loads(item.get("permissions") or "{}")
    except json.JSONDecodeError:
        permissions = {}
    return {
        "id": item.get("id"),
        "token": item.get("token"),
        "title": item.get("title"),
        "audience": item.get("audience") or "家教/家长/教务",
        "note": item.get("note") or "",
        "status": item.get("status"),
        "permissions": permissions,
        "created_at": item.get("created_at"),
        "last_viewed_at": item.get("last_viewed_at"),
        "expires_at": item.get("expires_at"),
    }


def create_profile_share(conn: sqlite3.Connection, data: dict, user: dict | None) -> dict:
    if not user:
        raise PermissionError("请先登录后再创建共享链接")
    title = (data.get("title") or "学习状态共享").strip()[:80]
    audience = (data.get("audience") or "家教/家长/教务").strip()[:80]
    note = (data.get("note") or "").strip()[:300]
    permissions = data.get("permissions") or {
        "report": True,
        "profile": True,
        "history": True,
        "questions": True,
        "auth_required": False,
    }
    share_id = str(uuid.uuid4())
    token = secrets.token_urlsafe(24)
    conn.execute(
        """
        insert into profile_shares
        (id, token, user_id, title, audience, note, status, permissions, created_at, expires_at)
        values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            share_id,
            token,
            user["id"],
            title,
            audience,
            note,
            "active",
            json.dumps(permissions, ensure_ascii=False),
            now_iso(),
            (data.get("expires_at") or "").strip(),
        ),
    )
    row = conn.execute("select * from profile_shares where id = ?", (share_id,)).fetchone()
    return public_profile_share(row)


def list_profile_shares(conn: sqlite3.Connection, user: dict | None) -> list[dict]:
    if not user:
        return []
    where, params = ("", []) if user_is_admin(user) else (" where user_id = ?", [user["id"]])
    rows = conn.execute(f"select * from profile_shares{where} order by created_at desc", params).fetchall()
    return [public_profile_share(row) for row in rows]


def revoke_profile_share(conn: sqlite3.Connection, share_id: str, user: dict | None) -> dict:
    if not user:
        raise PermissionError("unauthorized")
    row = conn.execute("select * from profile_shares where id = ?", (share_id,)).fetchone()
    if not row:
        raise ValueError("共享链接不存在")
    item = dict(row)
    if not user_is_admin(user) and item.get("user_id") != user.get("id"):
        raise PermissionError("无权操作该共享链接")
    conn.execute("update profile_shares set status = ? where id = ?", ("revoked", share_id))
    row = conn.execute("select * from profile_shares where id = ?", (share_id,)).fetchone()
    return public_profile_share(row)


def build_public_share_payload(conn: sqlite3.Connection, token: str, viewer: dict | None = None) -> dict:
    row = conn.execute("select * from profile_shares where token = ?", (token,)).fetchone()
    if not row:
        raise ValueError("共享链接不存在")
    share = public_profile_share(row)
    if share.get("status") != "active":
        raise PermissionError("共享链接已关闭")
    expires_at = (share.get("expires_at") or "").strip()
    if expires_at and expires_at < now_iso():
        raise PermissionError("共享链接已过期")
    owner = conn.execute("select * from app_users where id = ?", (dict(row).get("user_id"),)).fetchone()
    if not owner:
        raise ValueError("共享用户不存在")
    user = public_user(owner)
    conn.execute("update profile_shares set last_viewed_at = ? where token = ?", (now_iso(), token))
    permissions = share.get("permissions") or {}
    if permissions.get("auth_required") and not viewer:
        raise PermissionError("该共享链接需要统一账号登录后查看")
    report = build_report(conn, user) if permissions.get("report", True) else {}
    profile = profile_items(conn, user) if permissions.get("profile", True) else {"subjects": [], "items": []}
    history = build_generation_history(conn, user) if permissions.get("history", True) else {"items": [], "counts": {}, "total": 0}
    items = profile.get("items") or []
    recent_items = sorted(items, key=lambda item: item.get("created_at") or "", reverse=True)[:12]
    subject_cards = []
    for subject in profile.get("subjects") or []:
        subject_items = [item for item in items if clean_subject_name((item.get("diagnosis") or {}).get("subject")) == subject.get("name")]
        passed = len([item for item in subject_items if item.get("status") == "passed"])
        subject_cards.append({
            "name": subject.get("name"),
            "count": subject.get("count"),
            "passed": passed,
            "pass_rate": round(passed / len(subject_items) * 100, 1) if subject_items else 0,
        })
    return {
        "share": share,
        "student": {
            "display_name": masked_email(user.get("email")),
            "created_at": user.get("created_at"),
        },
        "summary": {
            "total_wrong_questions": report.get("total_wrong_questions", len(items)),
            "passed_questions": report.get("passed_questions", 0),
            "pass_rate": report.get("pass_rate", 0),
            "history_total": history.get("total", 0),
            "subject_count": len(profile.get("subjects") or []),
        },
        "report": report,
        "subjects": subject_cards,
        "recent_questions": [
            {
                "id": item.get("id"),
                "subject": clean_subject_name((item.get("diagnosis") or {}).get("subject")),
                "title": (item.get("diagnosis") or {}).get("core_pattern") or (item.get("diagnosis") or {}).get("topic") or "题目档案",
                "question": compact_text(item.get("corrected_text"), 180),
                "status": item.get("status"),
                "created_at": item.get("created_at"),
                "variants": len(item.get("variants") or []),
            }
            for item in recent_items
        ],
        "history": (history.get("items") or [])[:16],
        "generated_at": now_iso(),
    }


def build_admin_overview(conn: sqlite3.Connection) -> dict:
    chat_rows = conn.execute("select * from model_configs order by is_default desc, updated_at desc").fetchall()
    image_rows = conn.execute("select * from image_model_configs order by is_default desc, updated_at desc").fetchall()
    chat_models = [public_model(row) for row in chat_rows]
    image_models = [public_image_model(row) for row in image_rows]
    default_chat = next((item for item in chat_models if item.get("is_default")), chat_models[0] if chat_models else None)
    default_image = next((item for item in image_models if item.get("is_default")), image_models[0] if image_models else None)
    return {
        "chat_models": chat_models,
        "image_models": image_models,
        "default_chat_model_id": default_chat.get("id") if default_chat else "",
        "default_image_model_id": default_image.get("id") if default_image else "",
    }


def test_model_connection(conn: sqlite3.Connection, model_id: str) -> dict:
    model = get_model(conn, model_id)
    try:
        sample = minimax_chat(
            model,
            [{"role": "user", "content": "请只回复 OK。"}],
            max_tokens=16,
            temperature=0,
        )
        return {
            "ok": True,
            "model_id": model_id,
            "message": "连接成功",
            "sample": compact_text(sample, 80),
        }
    except Exception as exc:
        return {
            "ok": False,
            "model_id": model_id,
            "message": str(exc),
        }


def update_model_api_key(conn: sqlite3.Connection, model_id: str, api_key: str) -> dict:
    row = conn.execute("select * from model_configs where id = ?", (model_id,)).fetchone()
    if not row:
        raise ValueError("模型不存在")
    if not api_key.strip():
        raise ValueError("API Key 不能为空")
    conn.execute(
        "update model_configs set api_key = ?, updated_at = ? where id = ?",
        (api_key.strip(), now_iso(), model_id),
    )
    updated = conn.execute("select * from model_configs where id = ?", (model_id,)).fetchone()
    return public_model(updated)


def update_image_model_api_key(conn: sqlite3.Connection, model_id: str, api_key: str) -> dict:
    row = conn.execute("select * from image_model_configs where id = ?", (model_id,)).fetchone()
    if not row:
        raise ValueError("图片模型不存在")
    if not api_key.strip():
        raise ValueError("API Key 不能为空")
    conn.execute(
        "update image_model_configs set api_key = ?, updated_at = ? where id = ?",
        (api_key.strip(), now_iso(), model_id),
    )
    updated = conn.execute("select * from image_model_configs where id = ?", (model_id,)).fetchone()
    return public_image_model(updated)


def public_app_config() -> dict:
    return {
        "use_unified_auth": unified_auth_enabled(),
        "unified_auth_url": public_auth_base_url(),
        "platform_id": UNIFIED_PLATFORM_ID,
        "local_auth_enabled": True,
    }


def normalize_ocr_image_data_url(image_data_url: str, max_long_side: int = 2200) -> tuple[str, dict]:
    """Apply EXIF orientation and bounded resizing; keep the original on any failure."""
    metadata = {"processed": False, "reason": "unchanged"}
    try:
        from PIL import Image, ImageOps

        header, encoded = image_data_url.split(",", 1)
        if not header.lower().startswith("data:image/") or ";base64" not in header.lower():
            return image_data_url, {**metadata, "reason": "unsupported_data_url"}
        source_bytes = base64.b64decode(encoded, validate=True)
        with Image.open(BytesIO(source_bytes)) as source:
            orientation = int(source.getexif().get(274, 1) or 1)
            image = ImageOps.exif_transpose(source)
            original_size = list(image.size)
            needs_resize = max(image.size) > max_long_side
            if needs_resize:
                image.thumbnail((max_long_side, max_long_side), Image.Resampling.LANCZOS)
            if orientation == 1 and not needs_resize:
                return image_data_url, {
                    **metadata, "reason": "already_optimized", "original_size": original_size,
                    "output_size": original_size, "original_bytes": len(source_bytes), "output_bytes": len(source_bytes),
                }
            if image.mode not in {"RGB", "L"}:
                image = image.convert("RGB")
            output = BytesIO()
            image.save(output, format="JPEG", quality=88, optimize=True)
            output_bytes = output.getvalue()
        return "data:image/jpeg;base64," + base64.b64encode(output_bytes).decode("ascii"), {
            "processed": True,
            "reason": "oriented_and_resized" if orientation != 1 and needs_resize else "exif_oriented" if orientation != 1 else "resized",
            "original_size": original_size,
            "output_size": list(image.size),
            "original_bytes": len(source_bytes),
            "output_bytes": len(output_bytes),
        }
    except Exception as exc:
        return image_data_url, {**metadata, "reason": "preprocess_failed", "error": compact_text(exc, 160)}


def run_paper_ocr(image_data_url: str, model_id: str | None = None, *, handwriting: bool = True) -> dict:
    with db() as conn:
        model = get_vision_model(conn, model_id)
        ocr_prompt = merge_handwriting_prompt(PAPER_OCR_PROMPT, handwriting)
    normalized_image, image_metadata = normalize_ocr_image_data_url(image_data_url, max_long_side=2400)
    if handwriting:
        enhanced_image, enhanced_meta = preprocess_for_handwriting(normalized_image, max_long_side=2600)
        image_metadata = {"base": image_metadata, "enhanced": enhanced_meta}
        candidates = [normalized_image, enhanced_image]
    else:
        candidates = [normalized_image]
    best_result = None
    best_score = -1.0
    best_raw = ""
    for index, candidate in enumerate(candidates):
        content = [
            {"type": "text", "text": ocr_prompt},
            {"type": "image_url", "image_url": {"url": candidate, "detail": "high", "max_long_side_pixel": 2600 if handwriting else 2200}},
        ]
        raw = minimax_chat(model, [{"role": "user", "content": content}], max_tokens=7000, temperature=0.08)
        result = parse_paper_ocr_response(raw)
        score = ocr_quality_score(result)
        if score > best_score:
            best_score = score
            best_result = result
            best_raw = raw
            best_result["recognition_variant"] = "enhanced" if index else "base"
    best_result = best_result or {"questions": [], "page_text": "", "page_confidence": 0}
    best_result["raw_response_preview"] = compact_text(best_raw, 600)
    best_result["image_preprocessing"] = image_metadata
    best_result["handwriting_mode"] = handwriting
    return best_result


def ocr_quality_score(result: dict) -> float:
    questions = result.get("questions") or []
    confidence = float(result.get("page_confidence") or 0)
    usable = sum(1 for q in questions if len(str(q.get("printed_text") or "").strip()) >= 6)
    return confidence * 0.65 + min(1.0, usable / max(1, len(questions))) * 0.2 + min(1.0, usable / 8) * 0.15


def paper_ocr_needs_retry(result: dict) -> bool:
    questions = result.get("questions") if isinstance(result.get("questions"), list) else []
    page_text = str(result.get("page_text") or "").strip()
    usable = [q for q in questions if len(str(q.get("printed_text") or "").strip()) >= 6]
    if not page_text or not usable:
        return True
    if result.get("parse_mode") in {"text_fallback", "truncated_json_fallback"}:
        return True
    confidence = _clamp_confidence(result.get("page_confidence"), 0.0)
    # A complete, segmented page is more valuable than a self-reported confidence score.
    return confidence < 0.42 and len(usable) < 2


def best_of_two_paper_ocr(image_data_url: str, model_id: str | None = None) -> dict:
    """Retry only unusable pages, then keep the structurally stronger result."""
    first = run_paper_ocr(image_data_url, model_id)
    if not paper_ocr_needs_retry(first):
        first["recognition_attempts"] = 1
        first["quality_score"] = round(ocr_quality_score(first), 3)
        return first
    second = run_paper_ocr(image_data_url, model_id)
    chosen = max((first, second), key=ocr_quality_score)
    chosen["recognition_attempts"] = 2
    chosen["quality_score"] = round(ocr_quality_score(chosen), 3)
    return chosen


def image_url_to_data_url(source_url: str) -> str:
    relative = str(source_url or "").split("?", 1)[0].lstrip("/")
    path = PUBLIC_DIR / relative
    if not path.exists():
        raise ValueError("试卷页面原图不存在")
    mime = "image/png" if path.suffix.lower() == ".png" else "image/jpeg"
    return f"data:{mime};base64,{base64.b64encode(path.read_bytes()).decode('ascii')}"


def build_eight_steps(diagnosis: dict, question: dict) -> list[dict]:
    decomposition = diagnosis.get("decomposition") or {}
    analysis = diagnosis.get("student_answer_analysis") or {}
    strategy = diagnosis.get("learning_strategy") or {}
    standard = diagnosis.get("standard_answer") or {}
    supplied = [
        {"key": "understand", "content": diagnosis.get("problem_goal") or question.get("printed_text")},
        {"key": "conditions", "content": decomposition.get("total_formula") or diagnosis.get("cleaned_question")},
        {"key": "knowledge", "content": "；".join(diagnosis.get("knowledge_points") or [])},
        {"key": "diagnose", "content": analysis.get("likely_issue"), "evidence": analysis.get("evidence") or []},
        {"key": "model", "content": strategy.get("entry_point") or diagnosis.get("core_pattern")},
        {"key": "solve", "content": standard.get("concise_solution") or "\n".join(str(x) for x in decomposition.get("step_formulas") or [])},
        {"key": "verify", "content": standard.get("final_answer") or "代入、单位、边界与逻辑校核"},
        {"key": "transfer", "content": strategy.get("make_it_easier") or "完成同型巩固、轻微变式与综合迁移"},
    ]
    return normalize_eight_steps({"eight_steps": supplied})


def ensure_gaokao_mother_catalog(conn: sqlite3.Connection) -> None:
    """幂等写入高考题型模型；后续教研导入不会被覆盖。"""
    columns = {row["name"] for row in conn.execute("pragma table_info(mother_questions)").fetchall()}
    for item in SEED_MOTHER_QUESTIONS:
        if conn.execute("select 1 from mother_questions where code=? limit 1", (item["code"],)).fetchone():
            continue
        metadata = {key: value for key, value in item.items() if key not in {"code", "name"}}
        values = {
            "id": str(uuid.uuid4()), "code": item["code"], "name": item["name"],
            "topic": item["name"], "difficulty": 3,
            "recognition_signals": json.dumps(item.get("keywords") or [], ensure_ascii=False),
            "knowledge_points": json.dumps(item.get("keywords") or [], ensure_ascii=False),
            "solution_steps": json.dumps([item.get("formula") or ""], ensure_ascii=False),
            "common_error_causes": json.dumps(item.get("reminders") or [], ensure_ascii=False),
            "mnemonic": item.get("formula") or "", "status": "seed_verified",
            "metadata": json.dumps(metadata, ensure_ascii=False), "created_at": now_iso(),
        }
        names = [name for name in values if name in columns]
        conn.execute(
            f"insert into mother_questions ({','.join(names)}) values ({','.join('?' for _ in names)})",
            [values[name] for name in names],
        )


def insert_reserved_mother_question(
    conn: sqlite3.Connection,
    wrong_id: str,
    reserved: dict,
    diagnosis: dict,
) -> None:
    """Best-effort reserved mother insert; never block diagnosis save."""
    if not reserved.get("name"):
        return
    if reserved.get("status") in {"search_matched", "search_fast_path"}:
        return
    columns = {row["name"] for row in conn.execute("pragma table_info(mother_questions)").fetchall()}
    if not columns:
        return
    code = str(reserved.get("code") or reserved.get("name") or f"RES-{wrong_id[:8]}")[:120]
    if conn.execute("select 1 from mother_questions where code = ? limit 1", (code,)).fetchone():
        return
    topic = str(
        diagnosis.get("core_pattern") or diagnosis.get("topic") or reserved.get("name") or "待归纳题型"
    ).strip()
    values = {
        "id": str(uuid.uuid4()),
        "code": code,
        "name": str(reserved.get("name") or "预留母题"),
        "topic": topic,
        "difficulty": int(reserved.get("difficulty") or 3),
        "recognition_signals": json.dumps(reserved.get("recognition_signals") or [], ensure_ascii=False),
        "knowledge_points": json.dumps(reserved.get("knowledge_points") or [], ensure_ascii=False),
        "solution_steps": json.dumps(reserved.get("solution_steps") or [], ensure_ascii=False),
        "common_error_causes": json.dumps(reserved.get("common_error_causes") or [], ensure_ascii=False),
        "mnemonic": str(reserved.get("mnemonic") or reserved.get("abstract_pattern") or ""),
        "status": str(reserved.get("status") or "prompt_reserved"),
        "metadata": json.dumps(reserved, ensure_ascii=False),
        "created_at": now_iso(),
    }
    names = [name for name in values if name in columns]
    if not names:
        return
    conn.execute(
        f"insert into mother_questions ({','.join(names)}) values ({','.join('?' for _ in names)})",
        [values[name] for name in names],
    )


def build_answer_verdict(question_text: str, wrong_answer: str, diagnosis: dict | None = None) -> dict:
    """First-screen verdict: did the student get it right?"""
    diagnosis = diagnosis or {}
    analysis = diagnosis.get("student_answer_analysis") if isinstance(diagnosis.get("student_answer_analysis"), dict) else {}
    standard = diagnosis.get("standard_answer") if isinstance(diagnosis.get("standard_answer"), dict) else {}
    student = (wrong_answer or analysis.get("extracted_work") or "").strip()
    final = str(standard.get("final_answer") or "").strip()
    status_text = " ".join(
        str(x or "")
        for x in (
            analysis.get("answer_status"),
            analysis.get("likely_issue"),
            analysis.get("error_type"),
            diagnosis.get("answer_state"),
        )
    )
    if not student:
        return {
            "status": "blank",
            "label": "未作答 / 未识别到手写",
            "score_hint": "没有读到学生作答。请确认手写区域清晰，或在「学生作答」里粘贴答案；无论对错都会给出讲解。",
            "student_answer": "",
            "standard_answer_preview": final[:240],
            "skip_error_step": True,
        }
    lowered = status_text.lower()
    if any(k in status_text for k in ("正确", "做对", "全对", "答案正确", "完全正确")) and "不正" not in status_text and "错误" not in status_text:
        status, label, hint = "correct", "做对了", "作答与标准结论一致。仍建议过一遍思路与同类巩固，防止碰巧蒙对。"
    elif any(k in status_text for k in ("部分正确", "半对", "部分得分", "漏步")):
        status, label, hint = "partial", "部分正确", "结论或步骤有得分点，但仍有断点。下一步会专门拆错因。"
    elif any(k in status_text for k in ("错误", "做错", "不对", "失分", "未做对", "概念混淆", "knowledge_gap", "concept_confusion")):
        status, label, hint = "wrong", "做错了", "已对照标准路径判定为未完全做对。下一步拆清错因，再讲正确思路。"
    elif final and student and final[:40] in student:
        status, label, hint = "correct", "做对了（初步）", "学生作答包含标准结论片段；仍建议核对完整步骤。"
    else:
        status, label, hint = "unknown", "待核对", "已读到学生作答，但置信度不足以自动判分。请结合标准答案与讲解自行确认。"
    return {
        "status": status,
        "label": label,
        "score_hint": hint,
        "student_answer": student[:800],
        "standard_answer_preview": final[:240],
        "skip_error_step": status == "correct",
    }


def enrich_gaokao_diagnosis(question: dict, diagnosis: dict, subject: str) -> dict:
    mother = match_mother_question(question.get("printed_text") or "", subject)
    enriched = dict(diagnosis or {})
    enriched["gaokao_card"] = build_gaokao_card(question, enriched, mother)
    enriched["poem"] = enriched["gaokao_card"]["memory_poem"]
    enriched["mother_question"] = mother
    standard = enriched.setdefault("standard_answer", {})
    if not standard.get("concise_solution"):
        standard["concise_solution"] = "待教师复核后补充规范分步解答"
        enriched["needs_review"] = True
    if not standard.get("scoring_points") and not enriched.get("scoring_points"):
        standard["scoring_points"] = ["识别条件与设问", "写出关键公式或依据", "完成推导并给出结论"]
        enriched["quality_warning"] = "评分点为结构化建议，需结合正式评分标准复核"
    enriched["quality_gate"] = {
        "has_answer_trace": bool((enriched.get("student_answer_analysis") or {}).get("extracted_work") or question.get("student_work")),
        "has_standard_solution": bool(standard.get("concise_solution")),
        "has_scoring_points": bool(standard.get("scoring_points") or enriched.get("scoring_points")),
        "has_mother_evidence": bool(mother),
    }
    return enriched


def fallback_paper_diagnosis(question: dict, error: Exception) -> dict:
    text = question.get("printed_text") or ""
    work = question.get("student_work") or ""
    return {
        "cleaned_question": text,
        "subject": "待确认",
        "topic": "模型结果待复核",
        "difficulty": 1,
        "confidence": 0.5,
        "core_pattern": "待教师复核题型",
        "knowledge_points": [],
        "problem_goal": "根据题干确定问题目标",
        "student_answer_analysis": {
            "answer_presence": "已提供" if work else "未提供",
            "extracted_work": work,
            "answer_status": "待复核",
            "likely_issue": "模型输出格式异常，已保留原题与作答，请人工确认",
            "evidence": [str(question.get("teacher_marks") or "").strip()] if question.get("teacher_marks") else [],
            "next_action": "确认题干、作答与批改痕迹后重新分析",
        },
        "learning_strategy": {
            "decomposition_answer": "先确认题意，再提取条件、选择方法、分步求解并验算",
            "make_it_easier": "先完成同知识点基础题，再回到原题",
            "entry_point": "确认题目条件和设问",
        },
        "decomposition": {"total_formula": "读题→条件→知识点→方法→求解→验算", "step_formulas": []},
        "standard_answer": {"final_answer": "待复核", "concise_solution": "待重新分析"},
        "practice_variants": [],
        "needs_review": True,
        "model_error": str(error),
    }


def paper_owner_clause(user: dict | None) -> tuple[str, list]:
    if user and user_is_admin(user):
        return "", []
    if user:
        return " and user_id = ?", [user["id"]]
    return " and user_id is null", []


def get_paper(conn: sqlite3.Connection, paper_id: str, user: dict | None) -> dict | None:
    owner_sql, params = paper_owner_clause(user)
    row = conn.execute(f"select * from exam_papers where id = ?{owner_sql}", [paper_id, *params]).fetchone()
    if not row:
        return None
    item = row_to_dict(row)
    item["summary"] = item.get("summary") if isinstance(item.get("summary"), dict) else json.loads(item.get("summary") or "{}")
    questions = []
    for qrow in conn.execute("select * from paper_questions where paper_id = ? order by cast(question_no as integer), question_no", (paper_id,)).fetchall():
        q = row_to_dict(qrow)
        q["bbox"] = q.get("bbox") if isinstance(q.get("bbox"), (list, dict)) else json.loads(q.get("bbox") or "null")
        q["eight_steps"] = q.get("eight_steps") if isinstance(q.get("eight_steps"), list) else json.loads(q.get("eight_steps") or "[]")
        q["diagnosis"] = q.get("diagnosis") if isinstance(q.get("diagnosis"), dict) else json.loads(q.get("diagnosis") or "{}")
        questions.append(q)
    item["questions"] = questions
    item["pages"] = [row_to_dict(x) for x in conn.execute("select * from paper_pages where paper_id = ? order by page_no", (paper_id,)).fetchall()]
    return item


def persist_wrong_from_paper(conn: sqlite3.Connection, paper_id: str, question_id: str, user_id: str | None, question: dict, diagnosis: dict) -> str:
    wrong_id = str(uuid.uuid4())
    confidence = float(diagnosis.get("confidence") or question.get("confidence") or 0.7)
    error_type = normalize_error_type((diagnosis.get("student_answer_analysis") or {}).get("likely_issue"))
    model = get_model(conn, None)
    conn.execute(
        """insert into wrong_questions
        (id,image_url,ocr_text,corrected_text,student_wrong_answer,model_id,diagnosis,status,confidence,user_id,error_type,workflow_state,created_at)
        values (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (wrong_id, None, question.get("printed_text"), question.get("printed_text") or "",
         question.get("student_work") or "", model["id"], json.dumps(diagnosis, ensure_ascii=False),
         "review_needed" if question.get("review_required") else "diagnosed", confidence, user_id,
         error_type, "diagnosed", now_iso()),
    )
    save_variants(conn, wrong_id, diagnosis)
    return wrong_id


def get_fast_diagnose_model(conn: sqlite3.Connection, model_id: str | None = None) -> dict:
    if model_id:
        return get_model(conn, model_id)
    for preferred in ("fenno-gpt-default", "minimax-m3-default"):
        row = conn.execute("select * from model_configs where id = ? and api_key != ''", (preferred,)).fetchone()
        if row:
            return row_to_dict(row)
    return get_model(conn, None)


def quick_diagnose_with_llm(
    question_text: str,
    wrong_answer: str = "",
    ocr_text: str = "",
    model_id: str | None = None,
    subject: str = "自动识别",
    rag_context: str = "",
) -> dict:
    with db() as conn:
        model = get_fast_diagnose_model(conn, model_id)
    user_text = f"""学科：{subject}
题干：{question_text}
OCR：{ocr_text or "(无)"}
学生作答/错解：{wrong_answer or "(未提供)"}
参考资料：{rag_context[:1200] if rag_context else "(无)"}"""
    try:
        raw = minimax_chat(
            model,
            [{"role": "system", "content": QUICK_DIAGNOSE_PROMPT}, {"role": "user", "content": user_text}],
            max_tokens=1800,
            temperature=0.2,
            timeout=35,
        )
        result = extract_json(raw)
        result.setdefault("subject", subject)
        result.setdefault("confidence", 0.78)
        result.setdefault("core_pattern", "待归纳题型")
        result["gaokao_rag"] = {"used": bool(rag_context), "mode": "quick_llm"}
        return result
    except Exception:
        return build_skeleton_diagnosis(question_text, wrong_answer, subject)


def generate_eliminate_variants(
    question_text: str,
    wrong_answer: str,
    diagnosis: dict,
    model_id: str | None = None,
    subject: str = "自动识别",
    *,
    use_llm: bool = False,
) -> list[dict]:
    core = diagnosis.get("core_pattern") or diagnosis.get("topic") or "同类母题"
    fallback = fallback_eliminate_variants(question_text, core, wrong_answer)
    if not use_llm:
        return fallback
    try:
        with db() as conn:
            model = get_model(conn, model_id)
        user_text = f"""学科：{subject}
原题：{question_text}
学生错答/卡点：{wrong_answer or "未提供"}
核心题型：{core}
错因摘要：{(diagnosis.get("student_answer_analysis") or {}).get("likely_issue") or "待核对"}

请生成 3 道由易到难的「消灭训练题」（同型巩固 / 轻微变式 / 综合迁移），JSON 格式：
{{"practice_variants":[{{"level":1,"title":"同型巩固","stem":"...","answer":"...","analysis":"..."}}, ...]}}"""
        raw = minimax_chat(model, [{"role": "user", "content": user_text}], max_tokens=2200, temperature=0.25)
        parsed = extract_json(raw)
        variants = parsed.get("practice_variants") or []
        if len(variants) >= 2:
            return variants[:3]
    except Exception:
        pass
    return fallback


def process_paper_job(paper_id: str, model_id: str | None = None) -> None:
    import paper_speed

    paper_speed.process_paper_job(sys.modules[__name__], paper_id, model_id)


def build_paper_docx(paper: dict) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading(paper.get("title") or "全卷分析报告", 0)
    summary = paper.get("summary") or {}
    doc.add_paragraph(f"题目总数：{summary.get('total_questions',0)}　错题/待复核：{summary.get('wrong_count',0)}　得分率：{summary.get('score_rate','-')}%")
    for q in paper.get("questions") or []:
        doc.add_heading(f"第 {q.get('question_no')} 题｜{q.get('answer_state')}", level=1)
        doc.add_paragraph(q.get("printed_text") or "")
        if q.get("student_work"): doc.add_paragraph("学生作答：" + q["student_work"])
        if q.get("teacher_marks"): doc.add_paragraph("批改痕迹：" + q["teacher_marks"])
        for step in q.get("eight_steps") or []:
            doc.add_heading(f"{step.get('number')}. {step.get('label')}", level=2)
            doc.add_paragraph(step.get("content") or "待复核")
        poem = (q.get("diagnosis") or {}).get("poem") or {}
        if poem.get("lines"):
            doc.add_heading(poem.get("title") or "题后记忆诗", level=2)
            for line in poem.get("lines") or []:
                doc.add_paragraph(line)
            doc.add_paragraph("模型解码：" + "；".join(x.get("model_hint") or x.get("review") or "" for x in poem.get("line_reviews") or []))
    stream = BytesIO(); doc.save(stream); return stream.getvalue()


def build_paper_pdf(paper: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    stream = BytesIO(); styles = getSampleStyleSheet()
    for style in styles.byName.values(): style.fontName = 'STSong-Light'
    story = [Paragraph(paper.get("title") or "全卷分析报告", styles['Title']), Spacer(1, 12)]
    summary = paper.get("summary") or {}
    story.append(Paragraph(f"题目总数：{summary.get('total_questions',0)}　错题/待复核：{summary.get('wrong_count',0)}　得分率：{summary.get('score_rate','-')}%", styles['BodyText']))
    for q in paper.get("questions") or []:
        story += [Spacer(1, 12), Paragraph(f"第 {q.get('question_no')} 题｜{q.get('answer_state')}", styles['Heading2']), Paragraph((q.get("printed_text") or "").replace('\n','<br/>'), styles['BodyText'])]
        for step in q.get("eight_steps") or []:
            story.append(Paragraph(f"{step.get('number')}. {step.get('label')}：{step.get('content') or '待复核'}", styles['BodyText']))
        poem = (q.get("diagnosis") or {}).get("poem") or {}
        if poem.get("lines"):
            story.append(Paragraph(poem.get("title") or "题后记忆诗", styles['Heading2']))
            for line in poem.get("lines") or []:
                story.append(Paragraph(line, styles['BodyText']))
    SimpleDocTemplate(stream, pagesize=A4).build(story); return stream.getvalue()


class AppHandler(BaseHTTPRequestHandler):
    server_version = "GaokaoMVP/0.2"

    def require_admin(self, conn: sqlite3.Connection) -> bool:
        user = current_user_from_request(conn, self.headers)
        if user_is_admin(user):
            return True
        self.send_json({"error": "需要管理员账号登录后访问后台配置。"}, 403)
        return False

    def end_headers(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type, Authorization")
        super().end_headers()

    def do_OPTIONS(self) -> None:
        self.send_response(204)
        self.end_headers()

    def do_GET(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path == "/healthz":
            self.send_json({"ok": True, "service": "gaokao-ai"})
            return
        if parsed.path.startswith("/api/"):
            self.handle_api_get(parsed.path, parse_qs(parsed.query))
            return
        self.serve_static(parsed.path)

    def do_POST(self) -> None:
        parsed = urlparse(self.path)
        if parsed.path.startswith("/api/"):
            self.handle_api_post(parsed.path)
            return
        self.not_found()

    def read_json(self) -> dict:
        length = int(self.headers.get("Content-Length", "0"))
        body = self.rfile.read(length) if length else b"{}"
        raw = None
        for encoding in ("utf-8-sig", "utf-16", "gb18030"):
            try:
                raw = body.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        if raw is None:
            raise ValueError("请求内容编码无法识别，请使用 UTF-8")
        return json.loads(raw or "{}")

    def send_json(self, payload: dict | list, status: int = 200) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_sse(self, event: str, payload: dict) -> None:
        chunk = f"event: {event}\ndata: {json.dumps(payload, ensure_ascii=False)}\n\n".encode("utf-8")
        self.wfile.write(chunk)
        self.wfile.flush()

    def begin_sse(self) -> None:
        self.send_response(200)
        self.send_header("Content-Type", "text/event-stream; charset=utf-8")
        self.send_header("Cache-Control", "no-cache")
        self.send_header("Connection", "close")
        self.send_header("X-Accel-Buffering", "no")
        self.end_headers()

    def send_bytes(self, body: bytes, content_type: str, filename: str) -> None:
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        self.send_header("Content-Disposition", f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def handle_api_get(self, path: str, query: dict) -> None:
        try:
            def platform_admin_ok() -> bool:
                with db() as conn:
                    return user_is_admin(current_user_from_request(conn, self.headers))

            org_resp = org_inst_store.handle_org_api("GET", path, self.headers, None, platform_admin_ok)
            if org_resp:
                self.send_json(org_resp[1], org_resp[0])
                return
            if path == "/api/aippt/code":
                uid = (query.get("uid") or [""])[0].strip() or None
                self.send_json(aippt_auth.fetch_grant_code(uid=uid))
                return
            if path == "/api/config":
                self.send_json(public_app_config())
                return
            with db() as conn:
                user = current_user_from_request(conn, self.headers)
                if path == "/api/portal-tools":
                    self.send_json(PORTAL_TOOLS)
                    return
                if path == "/api/gaokao-2026":
                    if not user:
                        self.send_json({"error": "请先通过统一账号登录"}, 401); return
                    manifest_path = DATA_DIR / "gaokao_2026" / "manifest.json"
                    if not manifest_path.exists():
                        self.send_json({"error": "真题库尚未导入"}, 404); return
                    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                    self.send_json(manifest)
                    return
                if path == "/api/gaokao-zone/stats":
                    self.send_json(zone_stats(conn, ZONE_GAOKAO_MATH))
                    return
                if path == "/api/gaokao-zone/documents":
                    self.send_json({"documents": list_documents(conn, ZONE_GAOKAO_MATH)})
                    return
                if path == "/api/agent/layers":
                    self.send_json(AGENT_LAYERS)
                    return
                if path == "/api/me":
                    self.send_json(user if user else {"error": "unauthorized"}, 200 if user else 401)
                    return
                if path == "/api/history":
                    self.send_json(build_generation_history(conn, user))
                    return
                if path == "/api/papers":
                    owner_sql, params = paper_owner_clause(user)
                    rows = conn.execute(f"select * from exam_papers where 1=1{owner_sql} order by created_at desc", params).fetchall()
                    payload = []
                    for row in rows:
                        item = row_to_dict(row); item["summary"] = item.get("summary") if isinstance(item.get("summary"), dict) else json.loads(item.get("summary") or "{}"); payload.append(item)
                    self.send_json(payload)
                    return
                paper_match = re.fullmatch(r"/api/papers/([^/]+)", path)
                if paper_match:
                    item = get_paper(conn, paper_match.group(1), user)
                    self.send_json(item if item else {"error":"not found"}, 200 if item else 404)
                    return
                export_match = re.fullmatch(r"/api/papers/([^/]+)/export/(docx|pdf)", path)
                if export_match:
                    item = get_paper(conn, export_match.group(1), user)
                    if not item:
                        self.send_json({"error":"not found"},404); return
                    fmt = export_match.group(2)
                    if fmt == "docx":
                        self.send_bytes(build_paper_docx(item), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{item['title']}-全卷分析.docx")
                    else:
                        self.send_bytes(build_paper_pdf(item), "application/pdf", f"{item['title']}-全卷分析.pdf")
                    return
                if path == "/api/rag/documents":
                    where, params = rag_user_filter(user)
                    rows = conn.execute(f"select * from rag_documents{where} order by created_at desc", params).fetchall()
                    self.send_json([public_rag_document(row) for row in rows])
                    return
                if path == "/api/rag/search":
                    q = (query.get("q") or [""])[0]
                    subject = (query.get("subject") or ["自动识别"])[0]
                    limit = int((query.get("limit") or [5])[0])
                    self.send_json(search_rag(conn, q, subject, user, limit))
                    return
                match = re.fullmatch(r"/api/agent/runs/([^/]+)", path)
                if match:
                    row = conn.execute("select * from agent_runs where id = ?", (match.group(1),)).fetchone()
                    item = row_to_dict(row)
                    if user_owns_agent_run(user, item):
                        self.send_json(item)
                        return
                    self.send_json({"error": "not found"}, 404)
                    return
                match = re.fullmatch(r"/api/tool-runs/([^/]+)", path)
                if match:
                    row = conn.execute("select * from tool_runs where id = ?", (match.group(1),)).fetchone()
                    item = row_to_dict(row)
                    if item and user_owns_tool_run(user, item):
                        self.send_json(item)
                        return
                    self.send_json({"error": "not found"}, 404)
                    return
                match = re.fullmatch(r"/api/profile/exports/([^/]+)", path)
                if match:
                    row = conn.execute("select * from profile_exports where id = ?", (match.group(1),)).fetchone()
                    item = row_to_dict(row)
                    if item and user and (user_is_admin(user) or item.get("user_id") == user["id"]):
                        self.send_json(item)
                        return
                    self.send_json({"error": "not found"}, 404)
                    return
                if path == "/api/models":
                    rows = conn.execute("select * from model_configs where provider != 'deepseek' order by is_default desc, updated_at desc").fetchall()
                    self.send_json([public_model(row) for row in rows])
                    return
                if path == "/api/image-models":
                    rows = conn.execute("select * from image_model_configs order by is_default desc, updated_at desc").fetchall()
                    self.send_json([public_image_model(row) for row in rows])
                    return
                if path == "/api/admin/overview":
                    if not self.require_admin(conn):
                        return
                    self.send_json(build_admin_overview(conn))
                    return
                if path == "/api/admin/models":
                    if not self.require_admin(conn):
                        return
                    rows = conn.execute("select * from model_configs order by is_default desc, updated_at desc").fetchall()
                    self.send_json([public_model(row) for row in rows])
                    return
                if path == "/api/admin/image-models":
                    if not self.require_admin(conn):
                        return
                    rows = conn.execute("select * from image_model_configs order by is_default desc, updated_at desc").fetchall()
                    self.send_json([public_image_model(row) for row in rows])
                    return
                if path == "/api/admin/prompts":
                    if not self.require_admin(conn):
                        return
                    rows = conn.execute("select * from prompt_templates order by key").fetchall()
                    self.send_json([row_to_dict(row) for row in rows])
                    return
                if path == "/api/mother-questions":
                    ensure_gaokao_mother_catalog(conn)
                    rows = conn.execute("select * from mother_questions order by created_at desc").fetchall()
                    self.send_json([row_to_dict(row) for row in rows])
                    return
                if path == "/api/wrong-questions/due":
                    if not user:
                        self.send_json([])
                        return
                    owner_sql, params = ("", []) if user_is_admin(user) else (" and user_id = ?", [user["id"]])
                    rows = conn.execute(
                        f"""
                        select id from wrong_questions
                        where workflow_state = 'review_scheduled'
                          and next_review_at is not null and next_review_at <= ?{owner_sql}
                        order by next_review_at asc
                        """,
                        [now_iso(), *params],
                    ).fetchall()
                    self.send_json([get_wrong_question(conn, row["id"]) for row in rows])
                    return
                if path == "/api/wrong-questions":
                    ids_param = (query.get("ids") or [""])[0]
                    if user and user_is_admin(user):
                        wq_where, wq_params = "", []
                    elif user:
                        wq_where, wq_params = " where user_id = ?", [user["id"]]
                    elif ids_param.strip():
                        id_list = [item.strip() for item in ids_param.split(",") if item.strip()]
                        if not id_list:
                            self.send_json([])
                            return
                        placeholders = ",".join("?" * len(id_list))
                        wq_where = f" where id in ({placeholders}) and (user_id is null or trim(user_id) = '')"
                        wq_params = id_list
                    else:
                        self.send_json([])
                        return
                    rows = conn.execute(
                        f"select * from wrong_questions{wq_where} order by created_at desc",
                        wq_params,
                    ).fetchall()
                    self.send_json([row_to_dict(row) for row in rows])
                    return
                if path == "/api/checkins/today":
                    if not user:
                        self.send_json({"error": "请先登录后选择身份"}, 401)
                        return
                    self.send_json(checkin_dashboard(conn, user))
                    return
                if path == "/api/question-method-packages/stats":
                    total = conn.execute("select count(*) c from question_method_packages").fetchone()["c"]
                    source_total = conn.execute("select count(*) c from gaokao_questions").fetchone()["c"]
                    self.send_json({"packages": int(total or 0), "source_questions": int(source_total or 0)})
                    return
                if path == "/api/report":
                    self.send_json(build_report(conn, user))
                    return
                if path == "/api/teacher/classes":
                    user_id = user["id"] if user else None
                    self.send_json({"classes": teacher_portal.list_classes(conn, user_id)})
                    return
                if path == "/api/teacher/review-queue":
                    user_id = user["id"] if user else None
                    limit = int((query.get("limit") or ["50"])[0])
                    items = teacher_portal.review_queue(conn, user_id, limit=limit)
                    self.send_json({"items": items, "count": len(items)})
                    return
                if path == "/api/teacher/review-queue/export":
                    user_id = user["id"] if user else None
                    items = teacher_portal.review_queue(conn, user_id, limit=100)
                    text = teacher_portal.review_export_text(items)
                    self.send_json({"markdown": text, "count": len(items)})
                    return
                if path == "/api/parent/weekly-report":
                    user_id = user["id"] if user else None
                    self.send_json(parent_report.build_weekly_report(conn, user_id))
                    return
                if path == "/api/profile":
                    self.send_json(profile_items(conn, user))
                    return
                if path == "/api/profile/shares":
                    self.send_json(list_profile_shares(conn, user))
                    return
                match = re.fullmatch(r"/api/share/([^/]+)", path)
                if match:
                    payload = build_public_share_payload(conn, match.group(1), user)
                    self.send_json(payload)
                    return
                match = re.fullmatch(r"/api/wrong-questions/([^/]+)", path)
                if match:
                    wrong_id = match.group(1)
                    if not user_owns_wrong_question(conn, user, wrong_id):
                        self.send_json({"error": "not found"}, 404)
                        return
                    item = get_wrong_question(conn, wrong_id)
                    self.send_json(item if item else {"error": "not found"}, 200 if item else 404)
                    return
                match = re.fullmatch(r"/api/wrong-questions/([^/]+)/variants", path)
                if match:
                    wrong_id = match.group(1)
                    if not user_owns_wrong_question(conn, user, wrong_id):
                        self.send_json({"error": "not found"}, 404)
                        return
                    rows = conn.execute(
                        "select * from exercise_variants where wrong_question_id = ? order by level",
                        (wrong_id,),
                    ).fetchall()
                    self.send_json([row_to_dict(row) for row in rows])
                    return
                match = re.fullmatch(r"/api/wrong-questions/([^/]+)/cards", path)
                if match:
                    wrong_id = match.group(1)
                    if not user_owns_wrong_question(conn, user, wrong_id):
                        self.send_json({"error": "not found"}, 404)
                        return
                    rows = conn.execute(
                        """
                        select * from study_cards
                        where wrong_question_id = ?
                        order by created_at desc
                        """,
                        (wrong_id,),
                    ).fetchall()
                    self.send_json([row_to_dict(row) for row in rows])
                    return
        except PermissionError as exc:
            self.send_json({"error": str(exc)}, 401)
            return
        except Exception as exc:
            self.send_json({"error": str(exc)}, 500)
            return
        self.not_found()

    def handle_api_post(self, path: str) -> None:
        try:
            def platform_admin_ok() -> bool:
                with db() as conn:
                    return user_is_admin(current_user_from_request(conn, self.headers))

            org_resp = org_inst_store.handle_org_api("POST", path, self.headers, self.read_json, platform_admin_ok)
            if org_resp:
                self.send_json(org_resp[1], org_resp[0])
                return
            if path == "/api/latex/convert":
                data = self.read_json()
                filename = str(data.get("filename") or "魔法粘贴.txt")[:240]
                paste = str(data.get("text") or "")
                data_url = str(data.get("data_url") or "")
                if len(paste) > 2_000_000 or len(data_url) > 22_000_000:
                    self.send_json({"error": "文件过大，单次转换上限约15MB"}, 413); return
                if data_url:
                    header, content = decode_data_url(data_url)
                    if header.startswith("data:image/"):
                        ocr = run_ocr(data_url, data.get("model_id"))
                        extracted = str(ocr.get("text") or ocr.get("ocr_text") or "")
                        result = convert_bytes(filename, content, extracted)
                        result["ocr_confidence"] = ocr.get("confidence")
                    else:
                        extracted = extract_document_text(filename, content)
                        result = convert_bytes(filename, content, extracted)
                else:
                    result = {"filename": f"{Path(filename).stem}.tex", "latex": text_to_latex(paste, Path(filename).stem), "engine": "magic-paste-web", "detected_format": detect_paste_format(paste), "requires_formula_review": False}
                self.send_json(result, 201); return
            if path == "/api/papers":
                data = self.read_json()
                def request_text(value, default=""):
                    if value is None: return default
                    if isinstance(value, str): return value
                    return json.dumps(value, ensure_ascii=False)
                title = request_text(data.get("title") or data.get("source_name"), "未命名试卷").strip()
                pages = data.get("pages") or []
                if data.get("paper_text"):
                    pages = [{"text": data.get("paper_text")}]
                if not pages:
                    self.send_json({"error":"请上传至少一页试卷或提供试卷文本"},400); return
                prepared_pages = []
                try:
                    for page in pages:
                        if not isinstance(page, dict):
                            page = {"text": request_text(page)}
                        prepared_pages.append(prepare_paper_page(page))
                except (ValueError, RuntimeError) as exc:
                    self.send_json({"error": str(exc)}, 400); return
                paper_id, job_id, ts = str(uuid.uuid4()), str(uuid.uuid4()), now_iso()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    conn.execute("insert into exam_papers (id,user_id,title,subject,status,source_name,summary,progress,created_at,updated_at) values (?,?,?,?,?,?,?,?,?,?)",
                                 (paper_id,request_text(user.get("id")) if user else None,title,request_text(data.get("subject"),"自动识别"),"queued",request_text(data.get("source_name"),title),"{}",0,ts,ts))
                    for page_no, (source_url, source_text) in enumerate(prepared_pages, start=1):
                        conn.execute("insert into paper_pages (id,paper_id,page_no,source_url,source_text,ocr_result,confidence) values (?,?,?,?,?,?,?)",
                                     (str(uuid.uuid4()),paper_id,page_no,source_url,source_text,"{}",0))
                    conn.execute("insert into paper_jobs (id,paper_id,status,progress,message,attempts,created_at,updated_at) values (?,?,?,?,?,?,?,?)",
                                 (job_id,paper_id,"queued",0,"等待分析",0,ts,ts))
                threading.Thread(target=process_paper_job,args=(paper_id,data.get("model_id")),daemon=True,name=f"paper-{paper_id[:8]}").start()
                self.send_json({"id":paper_id,"job_id":job_id,"status":"queued","progress":0},202)
                return
            retry_match = re.fullmatch(r"/api/papers/([^/]+)/retry", path)
            if retry_match:
                with db() as conn:
                    user = current_user_from_request(conn,self.headers)
                    paper = get_paper(conn,retry_match.group(1),user)
                    if not paper: self.send_json({"error":"not found"},404); return
                    conn.execute("delete from paper_questions where paper_id=?",(paper["id"],))
                    conn.execute("update exam_papers set status='queued',progress=0,error=null,updated_at=? where id=?",(now_iso(),paper["id"]))
                    conn.execute("update paper_jobs set status='queued',progress=0,message='重新分析',updated_at=? where paper_id=?",(now_iso(),paper["id"]))
                threading.Thread(target=process_paper_job,args=(paper["id"],None),daemon=True).start()
                self.send_json({"id":paper["id"],"status":"queued"},202); return
            if path == "/api/register":
                self.register_user()
                return
            if path == "/api/login":
                self.login_user()
                return
            if path == "/api/redeem":
                self.redeem_code()
                return
            if path == "/api/agent/solve":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = solve_with_agent(conn, data, user)
                self.send_json(result, 201)
                return
            if path == "/api/rag/documents":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = create_rag_document(conn, data, user)
                self.send_json(result, 201)
                return
            if path == "/api/rag/search":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = search_rag(conn, data.get("query") or "", data.get("subject") or "自动识别", user, int(data.get("limit") or 5))
                self.send_json(result)
                return
            if path == "/api/rag/generate-quiz":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    topic = data.get("topic") or data.get("query") or "错题复习"
                    subject = data.get("subject") or "自动识别"
                    hits = search_rag(conn, str(topic), subject, user, max(6, int(data.get("count") or 5)))
                    result = generate_rag_quiz_from_hits(str(topic), hits, int(data.get("count") or 5))
                self.send_json(result)
                return
            if path == "/api/rag/study-os":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    topic = data.get("topic") or data.get("query") or "错题自学"
                    subject = data.get("subject") or "自动识别"
                    hits = search_rag(conn, str(topic), subject, user, 8)
                    result = generate_rag_study_os(str(topic), hits)
                self.send_json(result)
                return
            if path == "/api/tool-runs":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = run_portal_tool(conn, data, user)
                self.send_json(result, 201)
                return
            match = re.fullmatch(r"/api/admin/models/([^/]+)/test", path)
            if match:
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    result = test_model_connection(conn, match.group(1))
                status = 200 if result.get("ok") else 502
                self.send_json(result, status)
                return
            match = re.fullmatch(r"/api/admin/models/([^/]+)/api-key", path)
            if match:
                data = self.read_json()
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    result = update_model_api_key(conn, match.group(1), data.get("api_key") or "")
                self.send_json(result)
                return
            match = re.fullmatch(r"/api/admin/image-models/([^/]+)/api-key", path)
            if match:
                data = self.read_json()
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    result = update_image_model_api_key(conn, match.group(1), data.get("api_key") or "")
                self.send_json(result)
                return
            if path == "/api/admin/models":
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                self.save_model()
                return
            if path == "/api/admin/default-model":
                data = self.read_json()
                model_id = data.get("id")
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    conn.execute("update model_configs set is_default = 0")
                    conn.execute("update model_configs set is_default = 1, updated_at = ? where id = ?", (now_iso(), model_id))
                self.send_json({"ok": True})
                return
            if path == "/api/admin/image-models":
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                self.save_image_model()
                return
            if path == "/api/admin/default-image-model":
                data = self.read_json()
                model_id = data.get("id")
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    conn.execute("update image_model_configs set is_default = 0")
                    conn.execute("update image_model_configs set is_default = 1, updated_at = ? where id = ?", (now_iso(), model_id))
                self.send_json({"ok": True})
                return
            if path == "/api/admin/prompts":
                data = self.read_json()
                key = data.get("key")
                content = data.get("content", "")
                if key not in {"ocr", "diagnosis", "grading"}:
                    self.send_json({"error": "invalid prompt key"}, 400)
                    return
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    conn.execute(
                        "update prompt_templates set content = ?, updated_at = ? where key = ?",
                        (content, now_iso(), key),
                    )
                self.send_json({"ok": True})
                return
            if path == "/api/ocr":
                data = self.read_json()
                image_data_url = data.get("image_data_url")
                handwriting = bool(data.get("handwriting")) or bool(data.get("handwriting_mode"))
                image_url = save_data_url(image_data_url)
                result = run_ocr(image_data_url, data.get("model_id"), handwriting=handwriting)
                cleaned_question = deduplicate_repeated_ocr_text(result.get("printed_question") or result.get("ocr_text") or "")
                result["printed_question"] = cleaned_question
                result["ocr_text"] = deduplicate_repeated_ocr_text(result.get("ocr_text") or cleaned_question)
                questions = split_numbered_questions(cleaned_question)
                result["questions"] = questions if len(questions) > 1 else []
                result["question_count"] = max(1, len(questions))
                result["image_url"] = image_url
                self.send_json(result, 201)
                return
            if path == "/api/extract-document":
                data = self.read_json()
                result = extract_document_from_data_url(
                    data.get("filename") or "upload.txt",
                    data.get("file_data_url") or "",
                )
                self.send_json(result, 201)
                return
            if path == "/api/diagnose":
                self.create_diagnosis()
                return
            if path == "/api/diagnose/stream":
                self.create_diagnosis_stream()
                return
            if path == "/api/diagnose/enrich":
                data = self.read_json()
                text = (data.get("question_text") or "").strip()
                wrong_answer = (data.get("student_wrong_answer") or "").strip()
                if not text:
                    self.send_json({"error": "question_text is required"}, 400)
                    return
                subject = data.get("subject") or "自动识别"
                with db() as conn:
                    rag_hits, rag_context = retrieve_gaokao_evidence(conn, text, subject)
                diagnosis = quick_diagnose_with_llm(
                    text,
                    wrong_answer,
                    (data.get("ocr_text") or "").strip(),
                    data.get("model_id"),
                    subject,
                    rag_context,
                )
                diagnosis["practice_variants"] = generate_eliminate_variants(
                    text, wrong_answer, diagnosis, data.get("model_id"), subject, use_llm=False
                )
                self.send_json({"diagnosis": diagnosis, "mode": "enriched_llm"}, 201)
                return
            if path == "/api/teacher/classes":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    user_id = user["id"] if user else None
                    item = teacher_portal.create_class(conn, user_id, data)
                self.send_json(item, 201)
                return
            if path in {"/api/checkins/profile", "/api/checkins/tasks", "/api/checkins/submit"}:
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    if not user:
                        self.send_json({"error": "请先登录后操作打卡"}, 401)
                        return
                    if path == "/api/checkins/profile":
                        result = save_learning_profile(conn, user, data)
                    elif path == "/api/checkins/tasks":
                        result = create_checkin_task(conn, user, data)
                    else:
                        result = submit_checkin_task(conn, user, data)
                self.send_json(result, 201)
                return
            if path == "/api/search/questions":
                data = self.read_json()
                query = (data.get("query") or data.get("question_text") or "").strip()
                if not query:
                    self.send_json({"error": "请提供需要深拆的题目文本"}, 400)
                    return
                with db() as conn:
                    merged = merge_gaokao_search_results(
                        conn,
                        query,
                        zone=data.get("zone") or ZONE_GAOKAO_MATH,
                        limit=int(data.get("limit") or 8),
                    )
                self.send_json(
                    {
                        "mode": "search",
                        "hits": merged["bank_hits"],
                        "count": len(merged["bank_hits"]),
                        "rag_hits": merged["rag_hits"],
                        "rag_count": len(merged["rag_hits"]),
                        "citations": merged["citations"],
                        "rag_context_preview": compact_text(merged["rag_context"], 600),
                    }
                )
                return
            if path == "/api/gaokao-zone/sync-rag":
                data = self.read_json()
                with db() as conn:
                    if not self.require_admin(conn):
                        return
                    result = gaokao_rag.sync_all_gaokao_to_rag(conn, rebuild=data.get("rebuild", True) is not False)
                self.send_json(result, 201)
                return
            if path == "/api/gaokao-zone/upload":
                data = self.read_json()
                filename = (data.get("filename") or "upload.pdf").strip()
                data_url = data.get("data_url") or data.get("file_data_url") or ""
                if not data_url:
                    self.send_json({"error": "请上传 PDF 文件"}, 400)
                    return
                header, content = decode_data_url(data_url)
                if "pdf" not in header.lower() and not filename.lower().endswith(".pdf"):
                    self.send_json({"error": "目前专区导入仅支持 PDF"}, 400)
                    return
                upload_path = UPLOAD_DIR / f"gaokao-{uuid.uuid4().hex[:10]}-{Path(filename).name}"
                upload_path.write_bytes(content)
                doc_type = gaokao_import.guess_doc_type(filename)
                try:
                    with db() as conn:
                        result = gaokao_import.import_pdf_file(
                            conn,
                            upload_path,
                            public_dir=PUBLIC_DIR,
                            doc_type=doc_type,
                            title=data.get("title") or Path(filename).stem,
                            also_rag=True,
                        )
                    self.send_json(result, 201)
                except Exception as exc:
                    self.send_json({"error": str(exc)}, 500)
                return
            match = re.fullmatch(r"/api/wrong-questions/([^/]+)/variants", path)
            if match:
                wrong_id = match.group(1)
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    if not user_owns_wrong_question(conn, user, wrong_id):
                        self.send_json({"error": "wrong question not found"}, 404)
                        return
                    wrong = get_wrong_question(conn, wrong_id)
                    if not wrong:
                        self.send_json({"error": "wrong question not found"}, 404)
                        return
                    variants = save_variants(conn, wrong_id, wrong["diagnosis"])
                    conn.execute("update wrong_questions set status = ? where id = ?", ("training", wrong_id))
                self.send_json(variants, 201)
                return
            match = re.fullmatch(r"/api/wrong-questions/([^/]+)/cards", path)
            if match:
                data = self.read_json()
                card = generate_study_card(
                    match.group(1),
                    data.get("image_model_id"),
                    data.get("style") or "",
                )
                self.send_json(card, 201)
                return
            match = re.fullmatch(r"/api/exercise-variants/([^/]+)/answers", path)
            if match:
                self.grade_answer(match.group(1))
                return
            if path == "/api/mother-questions":
                self.create_reserved_mother()
                return
            if path == "/api/profile/share":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = create_profile_share(conn, data, user)
                self.send_json(result, 201)
                return
            match = re.fullmatch(r"/api/profile/shares/([^/]+)/revoke", path)
            if match:
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = revoke_profile_share(conn, match.group(1), user)
                self.send_json(result)
                return
            if path == "/api/profile/export":
                data = self.read_json()
                with db() as conn:
                    user = current_user_from_request(conn, self.headers)
                    result = build_profile_markdown(conn, data, user)
                    export_id = str(uuid.uuid4())
                    user_id = user["id"] if user else None
                    conn.execute(
                        """
                        insert into profile_exports (id, filename, subject, markdown, count, user_id, created_at)
                        values (?, ?, ?, ?, ?, ?, ?)
                        """,
                        (
                            export_id,
                            result.get("filename") or "个人学习档案.md",
                            data.get("subject") or "全部学科",
                            result.get("markdown") or "",
                            int(result.get("count") or 0),
                            user_id,
                            now_iso(),
                        ),
                    )
                    result["id"] = export_id
                    self.send_json(result, 201)
                return
        except Exception as exc:
            self.send_json({"error": str(exc)}, 500)
            return
        self.not_found()

    def register_user(self) -> None:
        data = self.read_json()
        email = (data.get("email") or "").strip().lower()
        password = data.get("password") or ""
        if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email):
            self.send_json({"error": "请输入有效邮箱"}, 400)
            return
        if email == ADMIN_EMAIL:
            self.send_json({"error": "该邮箱为系统管理员账号，请直接登录"}, 409)
            return
        if len(password) < 6:
            self.send_json({"error": "密码至少 6 位"}, 400)
            return
        user_id = str(uuid.uuid4())
        token = secrets.token_urlsafe(32)
        with db() as conn:
            exists = conn.execute("select id from app_users where email = ?", (email,)).fetchone()
            if exists:
                self.send_json({"error": "邮箱已注册，请直接登录"}, 409)
                return
            conn.execute(
                "insert into app_users (id, email, password_hash, credits, created_at) values (?, ?, ?, ?, ?)",
                (user_id, email, hash_password(password), 9, now_iso()),
            )
            conn.execute(
                "insert into app_sessions (token, user_id, created_at) values (?, ?, ?)",
                (token, user_id, now_iso()),
            )
            row = conn.execute("select * from app_users where id = ?", (user_id,)).fetchone()
        self.send_json({"token": token, **public_user(row)}, 201)

    def login_user(self) -> None:
        data = self.read_json()
        email = (data.get("email") or "").strip().lower()
        password = data.get("password") or ""
        with db() as conn:
            row = conn.execute("select * from app_users where email = ?", (email,)).fetchone()
            if not row or not verify_password(password, row["password_hash"]):
                self.send_json({"error": "邮箱或密码错误"}, 401)
                return
            token = secrets.token_urlsafe(32)
            conn.execute(
                "insert into app_sessions (token, user_id, created_at) values (?, ?, ?)",
                (token, row["id"], now_iso()),
            )
        self.send_json({"token": token, **public_user(row)})

    def redeem_code(self) -> None:
        data = self.read_json()
        code = (data.get("code") or "").strip().upper()
        with db() as conn:
            user = current_user_from_request(conn, self.headers)
            if not user:
                self.send_json({"error": "请先登录后再兑换积分"}, 401)
                return
            row = conn.execute("select * from redeem_codes where code = ?", (code,)).fetchone()
            if not row:
                self.send_json({"error": "兑换码不存在"}, 404)
                return
            reusable_demo = code == "DEMO2026"
            if row["is_used"] and not reusable_demo:
                self.send_json({"error": "兑换码已使用"}, 409)
                return
            if not reusable_demo:
                conn.execute(
                    "update redeem_codes set is_used = 1, used_by = ?, used_at = ? where code = ?",
                    (user["id"], now_iso(), code),
                )
            conn.execute("update app_users set credits = credits + ? where id = ?", (int(row["credits"]), user["id"]))
            updated = conn.execute("select * from app_users where id = ?", (user["id"],)).fetchone()
        self.send_json({"ok": True, "added": int(row["credits"]), "credits": int(updated["credits"])})

    def save_model(self) -> None:
        data = self.read_json()
        model_id = data.get("id") or str(uuid.uuid4())
        name = data.get("name") or "MiniMax M3"
        provider = data.get("provider") or "minimax"
        endpoint = data.get("endpoint") or MINIMAX_ENDPOINT
        model = data.get("model") or "MiniMax-M3"
        api_key = data.get("api_key")
        with db() as conn:
            existing = conn.execute("select * from model_configs where id = ?", (model_id,)).fetchone()
            if existing:
                old = dict(existing)
                conn.execute(
                    """
                    update model_configs
                    set name=?, provider=?, endpoint=?, model=?,
                        api_key=?, supports_vision=?, temperature=?, max_tokens=?, updated_at=?
                    where id=?
                    """,
                    (
                        name,
                        provider,
                        endpoint,
                        model,
                        api_key if api_key is not None and api_key != "" else old.get("api_key"),
                        1 if data.get("supports_vision", True) else 0,
                        float(data.get("temperature", old.get("temperature", 0.45))),
                        int(data.get("max_tokens", old.get("max_tokens", 7000))),
                        now_iso(),
                        model_id,
                    ),
                )
            else:
                conn.execute(
                    """
                    insert into model_configs
                    (id, name, provider, endpoint, model, api_key, supports_vision,
                     temperature, max_tokens, is_default, created_at, updated_at)
                    values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        model_id,
                        name,
                        provider,
                        endpoint,
                        model,
                        api_key or "",
                        1 if data.get("supports_vision", True) else 0,
                        float(data.get("temperature", 0.45)),
                        int(data.get("max_tokens", 7000)),
                        0,
                        now_iso(),
                        now_iso(),
                    ),
                )
            row = conn.execute("select * from model_configs where id = ?", (model_id,)).fetchone()
        self.send_json(public_model(row), 201)

    def save_image_model(self) -> None:
        data = self.read_json()
        model_id = data.get("id") or str(uuid.uuid4())
        name = data.get("name") or "OpenAI GPT Image 卡片生成"
        provider = data.get("provider") or "openai"
        endpoint = data.get("endpoint") or OPENAI_IMAGE_ENDPOINT
        model = data.get("model") or DEFAULT_OPENAI_IMAGE_MODEL
        api_key = data.get("api_key")
        size = data.get("size") or DEFAULT_OPENAI_IMAGE_SIZE
        quality = data.get("quality") or DEFAULT_OPENAI_IMAGE_QUALITY
        with db() as conn:
            existing = conn.execute("select * from image_model_configs where id = ?", (model_id,)).fetchone()
            if existing:
                old = dict(existing)
                conn.execute(
                    """
                    update image_model_configs
                    set name=?, provider=?, endpoint=?, model=?, api_key=?,
                        size=?, quality=?, updated_at=?
                    where id=?
                    """,
                    (
                        name,
                        provider,
                        endpoint,
                        model,
                        api_key if api_key is not None and api_key != "" else old.get("api_key"),
                        size,
                        quality,
                        now_iso(),
                        model_id,
                    ),
                )
            else:
                conn.execute(
                    """
                    insert into image_model_configs
                    (id, name, provider, endpoint, model, api_key, size, quality,
                     is_default, created_at, updated_at)
                    values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """,
                    (
                        model_id,
                        name,
                        provider,
                        endpoint,
                        model,
                        api_key or "",
                        size,
                        quality,
                        0,
                        now_iso(),
                        now_iso(),
                    ),
                )
            row = conn.execute("select * from image_model_configs where id = ?", (model_id,)).fetchone()
        self.send_json(public_image_model(row), 201)

    def create_diagnosis_stream(self) -> None:
        data = self.read_json()
        text = (data.get("question_text") or "").strip()
        if not text:
            self.send_json({"error": "question_text is required"}, 400)
            return
        subject = data.get("subject") or "自动识别"
        wrong_answer = (data.get("student_wrong_answer") or "").strip()
        model_id = data.get("model_id")
        self.begin_sse()
        try:
            question_parts = split_numbered_questions(deduplicate_repeated_ocr_text(text))
            if len(question_parts) > 1:
                question_parts = question_parts[:10]
                self.send_sse("batch_detected", {"count": len(question_parts), "label": f"识别到 {len(question_parts)} 道题，将逐题进行 RAG 深拆"})
                items = []
                for index, question in enumerate(question_parts, start=1):
                    question_text = str(question.get("printed_text") or "").strip()
                    self.send_sse("question_start", {"index": index, "count": len(question_parts), "question_no": question.get("question_no"), "label": f"正在拆解第 {index}/{len(question_parts)} 题"})
                    instant = build_instant_diagnosis(question_text, wrong_answer, subject, [], [])
                    self.send_sse("partial", {"step": "instant", "index": index, "count": len(question_parts), "core_pattern": instant.get("core_pattern"), "subject": instant.get("subject"), "steps": (instant.get("decomposition") or {}).get("step_formulas") or [], "show_result": True})
                    with db() as conn:
                        merged = merge_gaokao_search_results(conn, question_text, limit=int(data.get("limit") or 6))
                    rag_hits = merged["rag_hits"]
                    self.send_sse("rag_hit", {"index": index, "count": len(rag_hits), "hits": rag_hits[:5], "bank_hits": merged["bank_hits"][:5], "method_packages": (merged.get("method_packages") or [])[:5]})
                    deep_data = {**data, "question_text": question_text, "ocr_text": question_text, "mode": "deep", "skip_search": False}
                    item = self._build_diagnosis_record(deep_data, question_text, rag_hits=rag_hits, rag_context=merged.get("rag_context") or "")
                    items.append(item)
                    self.send_sse("question_done", {"index": index, "count": len(question_parts), "question_no": question.get("question_no"), "id": item.get("id"), "core_pattern": (item.get("diagnosis") or {}).get("core_pattern")})
                self.send_sse("done", {"mode": "deep-batch", "item": items[0], "items": items, "question_count": len(items), "fast_path": False})
                return
            # 首段正文绝不等待数据库或模型；客户端应在 1–2 秒内看到可读分析，而不只是进度提示。
            self.send_sse("phase", {"step": "accepted", "label": "题目已接收，正在识别题型并检索知识库…"})
            instant = build_instant_diagnosis(text, wrong_answer, subject, [], [])
            self.send_sse(
                "partial",
                {
                    "step": "instant",
                    "core_pattern": instant.get("core_pattern"),
                    "subject": instant.get("subject"),
                    "steps": (instant.get("decomposition") or {}).get("step_formulas") or [],
                    "show_result": True,
                },
            )
            with db() as conn:
                merged = merge_gaokao_search_results(conn, text, limit=int(data.get("limit") or 6))
            hits = merged["bank_hits"]
            rag_hits = merged["rag_hits"]
            method_packages = merged.get("method_packages") or []
            rag_context = merged.get("rag_context") or ""
            self.send_sse("phase", {"step": "rag", "label": f"已匹配 {len(hits)} 道母题、{len(rag_hits)} 条知识证据和 {len(method_packages)} 个方法包…"})
            if rag_hits or method_packages or hits:
                self.send_sse(
                    "rag_hit",
                    {
                        "hits": rag_hits[:5],
                        "count": len(rag_hits),
                        "bank_hits": hits[:5],
                        "method_packages": method_packages[:5],
                        "citations": gaokao_rag.rag_citations_from_hits(rag_hits[:5]),
                    },
                )
            skeleton = build_instant_diagnosis(text, wrong_answer, subject, hits, rag_hits)
            self.send_sse("partial", {"step": "pattern", "core_pattern": skeleton.get("core_pattern"), "subject": skeleton.get("subject"), "steps": (skeleton.get("decomposition") or {}).get("step_formulas") or [], "show_result": True})
            self.send_sse("phase", {"step": "deep", "label": "正在结合母题、方法包与知识证据完成八步深拆…"})
            deep_data = {**data, "mode": "deep", "skip_search": False}
            item = self._build_diagnosis_record(deep_data, text, rag_hits=rag_hits, rag_context=rag_context)
            self.send_sse(
                "partial",
                {
                    "step": "crush",
                    "core_pattern": (item.get("diagnosis") or {}).get("core_pattern"),
                    "subject": (item.get("diagnosis") or {}).get("subject"),
                },
            )
            self.send_sse("done", {"mode": "deep", "item": item, "fast_path": False, "hits": hits[:5], "rag_hits": rag_hits[:5], "method_packages": method_packages[:5]})
        except (BrokenPipeError, ConnectionAbortedError, ConnectionResetError):
            return
        except Exception as exc:
            self.send_sse("error", {"error": str(exc)})

    def _build_diagnosis_record(
        self,
        data: dict,
        text: str,
        *,
        rag_hits: list[dict] | None = None,
        rag_context: str = "",
        prebuilt_diagnosis: dict | None = None,
    ) -> dict:
        image_url = save_data_url(data.get("image_data_url"))
        wrong_answer = (data.get("student_wrong_answer") or "").strip()
        ocr_text = (data.get("ocr_text") or "").strip()
        model_id = data.get("model_id")
        subject = data.get("subject") or "自动识别"
        if rag_hits is None or not rag_context:
            with db() as conn:
                rag_hits, rag_context = retrieve_gaokao_evidence(conn, text, subject)
        if prebuilt_diagnosis is not None:
            diagnosis = dict(prebuilt_diagnosis)
        elif (data.get("mode") or "auto").strip().lower() == "deep":
            diagnosis = diagnose_with_llm(text, wrong_answer, ocr_text, model_id, subject, rag_context=rag_context)
        else:
            with db() as conn:
                merged = merge_gaokao_search_results(conn, text, limit=6)
            diagnosis = build_instant_diagnosis(
                text, wrong_answer, subject, merged["bank_hits"], merged["rag_hits"]
            )
        diagnosis.setdefault(
            "gaokao_rag",
            {
                "used": bool(rag_hits),
                "evidence_count": len(rag_hits or []),
                "citations": gaokao_rag.rag_citations_from_hits(rag_hits or []),
                "how_used": "已作为拆题优先依据注入提示词" if rag_hits else "未命中母题库 RAG",
            },
        )
        answer_analysis = diagnosis.setdefault("student_answer_analysis", {})
        if wrong_answer and not str(answer_analysis.get("extracted_work") or "").strip():
            answer_analysis["extracted_work"] = wrong_answer
            answer_analysis["answer_presence"] = "已提供"
        error_type = normalize_error_type(answer_analysis.get("error_type") or answer_analysis.get("likely_issue"))
        answer_analysis["error_type"] = error_type
        diagnosis["answer_verdict"] = build_answer_verdict(text, wrong_answer, diagnosis)
        if not (diagnosis.get("practice_variants") or []):
            use_llm = (data.get("mode") or "auto").strip().lower() == "deep"
            diagnosis["practice_variants"] = generate_eliminate_variants(
                text, wrong_answer, diagnosis, model_id, subject, use_llm=use_llm
            )
        # Prefer 5 consolidate items for the guided eighth step.
        while len(diagnosis.get("practice_variants") or []) < 5:
            base = list(diagnosis.get("practice_variants") or [])
            extras = generate_eliminate_variants(
                text, wrong_answer, diagnosis, model_id, subject, use_llm=False
            )
            for extra in extras:
                if len(diagnosis["practice_variants"]) >= 5:
                    break
                diagnosis["practice_variants"].append(extra)
            if len(diagnosis["practice_variants"]) == len(base):
                break
        diagnosis["practice_variants"] = (diagnosis.get("practice_variants") or [])[:5]
        for index, variant in enumerate(diagnosis.get("practice_variants") or [], start=1):
            variant["tier"] = practice_tier(index)
        wrong_id = str(uuid.uuid4())
        confidence = float(diagnosis.get("confidence") or 0.75)
        status = "review_needed" if diagnosis.get("needs_review") else "diagnosed"
        saved_variants: list[dict] = []
        with db() as conn:
            user = current_user_from_request(conn, self.headers)
            user_id = user["id"] if user else None
            org_ctx = institution_context_for_app_user(user)
            actual_model = get_model(conn, model_id)
            conn.execute(
                """
                insert into wrong_questions
                (id, image_url, ocr_text, corrected_text, student_wrong_answer, model_id,
                 diagnosis, status, confidence, user_id, institution_id, institution_name,
                 institution_badge, error_type, workflow_state, created_at)
                values (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    wrong_id,
                    image_url,
                    ocr_text,
                    text,
                    wrong_answer,
                    actual_model["id"],
                    json.dumps(diagnosis, ensure_ascii=False),
                    status,
                    confidence,
                    user_id,
                    org_ctx.get("institution_id") or None,
                    org_ctx.get("institution_name") or None,
                    org_ctx.get("institution_badge") or None,
                    error_type,
                    "diagnosed",
                    now_iso(),
                ),
            )
            saved_variants = save_variants(conn, wrong_id, diagnosis)
            reserved = diagnosis.get("mother_question_reserved") or {}
            try:
                insert_reserved_mother_question(conn, wrong_id, reserved, diagnosis)
            except Exception:
                pass
            row = conn.execute("select * from wrong_questions where id = ?", (wrong_id,)).fetchone()
        item = row_to_dict(row)
        item["diagnosis"] = diagnosis
        item["variants"] = saved_variants
        return item

    def create_diagnosis(self) -> None:
        data = self.read_json()
        text = (data.get("question_text") or "").strip()
        if not text:
            self.send_json({"error": "question_text is required"}, 400)
            return
        with db() as conn:
            merged = merge_gaokao_search_results(conn, text, limit=6)
        deep_data = {**data, "mode": "deep", "skip_search": False}
        item = self._build_diagnosis_record(deep_data, text, rag_hits=merged["rag_hits"], rag_context=merged.get("rag_context") or "")
        self.send_json(item, 201)

    def grade_answer(self, variant_id: str) -> None:
        data = self.read_json()
        answer_text = (data.get("answer_text") or "").strip()
        model_id = data.get("model_id")
        hint_count = max(0, min(int(data.get("hint_count") or 0), 3))
        with db() as conn:
            variant = row_to_dict(conn.execute("select * from exercise_variants where id = ?", (variant_id,)).fetchone())
            if not variant:
                self.send_json({"error": "variant not found"}, 404)
                return
            wrong = get_wrong_question(conn, variant["wrong_question_id"])
        result = grade_with_llm(variant, answer_text, wrong.get("diagnosis") if wrong else None, model_id)
        answer_id = str(uuid.uuid4())
        with db() as conn:
            conn.execute(
                """
                insert into student_answers
                (id, exercise_variant_id, answer_text, grading_result, is_correct, hint_count, submitted_at)
                values (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    answer_id,
                    variant_id,
                    answer_text,
                    json.dumps(result, ensure_ascii=False),
                    1 if result["is_correct"] else 0,
                    hint_count,
                    now_iso(),
                ),
            )
            update_pass_status(conn, variant["wrong_question_id"])
        self.send_json({"id": answer_id, **result}, 201)

    def create_reserved_mother(self) -> None:
        data = self.read_json()
        item_id = str(uuid.uuid4())
        with db() as conn:
            insert_reserved_mother_question(
                conn,
                item_id,
                {
                    "code": data.get("code") or f"RES-{item_id[:8]}",
                    "name": data.get("name") or "预留母题",
                    "status": "reserved",
                    **(data.get("metadata") or {}),
                },
                {"core_pattern": data.get("topic") or data.get("name") or "预留母题"},
            )
            row = conn.execute(
                "select * from mother_questions where code = ? order by created_at desc limit 1",
                (data.get("code") or f"RES-{item_id[:8]}",),
            ).fetchone()
            if not row:
                self.send_json({"error": "创建母题失败"}, 500)
                return
        self.send_json(row_to_dict(row), 201)

    def serve_static(self, path: str) -> None:
        if path == "/":
            path = "/index.html"
        if path in {"/investor", "/investors", "/investor/", "/investors/"}:
            path = "/investors.html"
        if path == "/org":
            path = "/org.html"
        if path == "/org-admin":
            path = "/org-admin.html"
        # SPA routes for nav + 22 tool workbenches
        spa_paths = {
            "/paper-analysis", "/wrong-analysis", "/agent", "/gaokao-math",
            "/wrongbook", "/learning-report", "/history", "/workbenches",
            "/profile", "/mother-questions", "/admin",
        }
        normalized = path.rstrip("/") or "/"
        if normalized in spa_paths or re.fullmatch(r"/tools/[a-z0-9-]+/?", path):
            path = "/index.html"
        safe_path = path.lstrip("/").replace("..", "")
        file_path = PUBLIC_DIR / safe_path
        if not file_path.exists() or not file_path.is_file():
            self.not_found()
            return
        content_type = "text/plain; charset=utf-8"
        if file_path.suffix == ".html":
            content_type = "text/html; charset=utf-8"
        elif file_path.suffix == ".css":
            content_type = "text/css; charset=utf-8"
        elif file_path.suffix == ".js":
            content_type = "text/javascript; charset=utf-8"
        elif file_path.suffix in (".png", ".jpg", ".jpeg", ".webp"):
            kind = "jpeg" if file_path.suffix in (".jpg", ".jpeg") else file_path.suffix.lstrip(".")
            content_type = f"image/{kind}"
        elif file_path.suffix == ".docx":
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif file_path.suffix == ".pptx":
            content_type = "application/vnd.openxmlformats-officedocument.presentationml.presentation"
        elif file_path.suffix == ".md":
            content_type = "text/markdown; charset=utf-8"
        body = file_path.read_bytes()
        self.send_response(200)
        self.send_header("Content-Type", content_type)
        if "exports" in file_path.parts and file_path.suffix in (".docx", ".pptx", ".md", ".png"):
            self.send_header("Content-Disposition", f'attachment; filename="{file_path.name}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def not_found(self) -> None:
        self.send_json({"error": "not found"}, 404)

    def log_message(self, fmt: str, *args) -> None:
        print(f"[{now_iso()}] {self.address_string()} {fmt % args}")


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        try:
            sys.stdout.reconfigure(encoding="utf-8")
            sys.stderr.reconfigure(encoding="utf-8")
        except Exception:
            pass
    init_db()
    host = os.environ.get("HOST", "127.0.0.1")
    port = int(os.environ.get("PORT", "8021"))
    server = ThreadingHTTPServer((host, port), AppHandler)
    print(f"AI错题拆博士已启动: http://{host}:{port}")
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        pass
    finally:
        server.server_close()


if __name__ == "__main__":
    main()
